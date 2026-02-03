"""
LLM-Driven Document Injection Analyzer

Instead of relying on regex patterns to detect injection points,
this module uses Claude to intelligently analyze document content
and identify all areas that need Kahua token injection.

The LLM understands:
- Context and semantics of the document
- Complex patterns (checkboxes, inline blanks, conditional text)
- Field type inference from surrounding context
- Domain-specific knowledge (construction, contracts, change orders)

Now schema-aware: Uses SchemaService to provide actual field information
to the LLM, enabling accurate field path mapping.
"""

import os
import io
import json
import logging
import asyncio
from dataclasses import dataclass, field
from typing import Optional, List, Dict, Any, Tuple
from enum import Enum

from anthropic import Anthropic
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

# Import Kahua token builders
try:
    from template_gen.kahua_tokens import (
        build_attribute_token,
        build_boolean_token,
        build_currency_token,
        build_date_token,
        build_number_token,
        build_field_token,
        to_kahua_path,
    )
except ImportError:
    from kahua_tokens import (
        build_attribute_token,
        build_boolean_token,
        build_currency_token,
        build_date_token,
        build_number_token,
        build_field_token,
        to_kahua_path,
    )

# Schema service for field context
try:
    from src.report_genius.schema_service import (
        get_schema_service,
        get_entity_schema,
        resolve_entity,
        EntitySchema,
    )
except ImportError:
    # Running from different directory
    try:
        from report_genius.schema_service import (
            get_schema_service,
            get_entity_schema,
            resolve_entity,
            EntitySchema,
        )
    except ImportError:
        get_schema_service = None
        get_entity_schema = None
        resolve_entity = None
        EntitySchema = None

log = logging.getLogger("llm_injection_analyzer")


# ============== Data Models ==============

class InjectionType(str, Enum):
    """Type of injection needed."""
    CURRENCY = "currency"
    DATE = "date"
    NUMBER = "number"
    TEXT = "text"
    BOOLEAN = "boolean"
    CHECKBOX = "checkbox"
    RICH_TEXT = "rich_text"
    CONDITIONAL = "conditional"


@dataclass
class InjectionPoint:
    """A single point in the document needing token injection."""
    # Location
    location_type: str  # "paragraph" or "table_cell"
    paragraph_index: Optional[int] = None
    table_index: Optional[int] = None
    row_index: Optional[int] = None
    cell_index: Optional[int] = None
    
    # What to inject
    original_text: str = ""
    text_to_replace: str = ""  # The specific text/pattern to replace
    kahua_field_path: str = ""  # e.g., "OriginalContractAmount"
    injection_type: InjectionType = InjectionType.TEXT
    
    # For conditional/checkbox fields
    condition_context: Optional[Dict[str, Any]] = None
    
    # Token to inject (generated)
    token: str = ""
    
    # LLM reasoning
    reasoning: str = ""
    confidence: float = 0.9


@dataclass
class DocumentContent:
    """Structured content extracted from a DOCX."""
    paragraphs: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    raw_text: str = ""


@dataclass
class LLMAnalysisResult:
    """Result from LLM analysis."""
    success: bool
    injection_points: List[InjectionPoint] = field(default_factory=list)
    document_summary: str = ""
    entity_type_detected: str = ""
    warnings: List[str] = field(default_factory=list)
    suggestions: List[str] = field(default_factory=list)
    error: Optional[str] = None


# ============== Document Content Extraction ==============

def extract_document_content(doc_bytes: bytes) -> DocumentContent:
    """
    Extract structured content from a DOCX file for LLM analysis.
    Preserves location information for later injection.
    """
    doc = Document(io.BytesIO(doc_bytes))
    content = DocumentContent()
    all_text_parts = []
    
    # Extract paragraphs with indices
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            content.paragraphs.append({
                "index": idx,
                "text": text,
                "style": para.style.name if para.style else "Normal",
            })
            all_text_parts.append(f"[P{idx}] {text}")
    
    # Extract tables with cell indices
    for table_idx, table in enumerate(doc.tables):
        table_data = {
            "index": table_idx,
            "rows": []
        }
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                row_data.append({
                    "cell_index": cell_idx,
                    "text": cell_text
                })
            table_data["rows"].append({
                "row_index": row_idx,
                "cells": row_data
            })
        content.tables.append(table_data)
        
        # Add to text representation - IMPROVED: show cell indices explicitly
        all_text_parts.append(f"\n[TABLE {table_idx}] ({len(table.rows)} rows x {len(table.columns)} cols)")
        for row in table_data["rows"]:
            # Show each cell with its index for precise injection targeting
            cell_strs = []
            for c in row["cells"]:
                cell_text = c["text"][:150] if c["text"] else "(empty)"  # Increased limit to capture full text
                cell_strs.append(f"[C{c['cell_index']}]{cell_text}")
            all_text_parts.append(f"  [R{row['row_index']}] {' | '.join(cell_strs)}")
    
    content.raw_text = "\n".join(all_text_parts)
    return content


# ============== LLM Analysis ==============

INJECTION_ANALYSIS_SYSTEM = """You are an expert document analyzer specializing in construction project management templates.

Your task is to analyze a document and identify ALL locations that need dynamic data injection (Kahua tokens).

## FIELD TYPES

1. **Currency fields**: Amounts, costs, sums, prices - anything financial
2. **Date fields**: Due dates, start/end dates, completion dates  
3. **Number fields**: Quantities, days, counts, percentages
4. **Text fields**: Names, descriptions, IDs, project numbers
5. **Boolean/Checkbox fields**: Yes/No decisions, checkboxes
6. **Conditional fields**: "increased/decreased/unchanged" selections

## CRITICAL PATTERNS TO DETECT

### 1. LABEL + DOLLAR SIGN IN SEPARATE CELLS
When you see a row like:
  [R10] [C0]Original Contract Sum | [C3]$ | [C4]$ | [C5]$

The `$` cells (C3, C4, C5) need currency tokens. The label "Original Contract Sum" tells you which field.
→ Inject in cell C3 (or first $ cell): `[Currency(Source=Attribute,Path=OriginalContractAmount,Format="C2")]`

### 2. CHECKBOX/OPTION PATTERNS  
When you see text like "increased   decreased" or "be increased   be decreased   unchanged":
- These are checkbox options where ONE will be checked
- Each option needs a Boolean token
→ Replace "increased" with `[Boolean(Source=Attribute,Path=IsContractSumIncreased,TrueValue="☒ increased",FalseValue="☐ increased")]`
→ Replace "decreased" with `[Boolean(Source=Attribute,Path=IsContractSumDecreased,TrueValue="☒ decreased",FalseValue="☐ decreased")]`

### 3. INLINE NUMBER BLANKS
When you see patterns like "(     ) days" or "(    ) calendar days":
- The blank space in parentheses is for a number
→ Replace "(     )" with `([Number(Source=Attribute,Path=ContractTimeChange,Format="N0")])`

### 4. LABEL: BLANK PATTERNS
When you see "PROJECT:" followed by empty or placeholder text:
→ Inject `[Attribute(DomainPartition.Name)]` after the colon

### 5. DATE PLACEHOLDERS
When you see "DATE:" or "CONTRACT DATE:" with blank space:
→ Inject `[Date(Source=Attribute,Path=Date,Format="d")]`

## TABLE CELL TARGETING

Document shows cells as [C0], [C1], etc. within each row [R0], [R1], etc.
When injecting into a table:
- Specify the exact table_index, row_index, AND cell_index
- The injection goes into that specific cell

## OUTPUT FORMAT

Output your analysis as structured JSON with precise cell targeting."""


INJECTION_ANALYSIS_PROMPT = """Analyze this document content and identify ALL injection points.

Target Entity Type: {entity_def}

Document Content (cell indices shown as [C0], [C1], etc.):
```
{document_content}
```

{schema_context}

## COMMON CHANGE ORDER FIELD MAPPINGS

| Label in Document | Kahua Field Path | Notes |
|-------------------|------------------|-------|
| Original Contract Sum | OriginalContractAmount | Currency - replace "$" |
| Net change by previously authorized Change Orders | PreviouslyApprovedChangesAmount | Currency - replace "$" |
| Contract Sum prior to this Change Order | CurrentContractAmount | Currency - replace "$" |
| Change Order will alter/change | CostItemsTotalTotalValue or NetAmount | Currency - replace "$" |
| New Contract Sum including this Change Order | NewContractAmount | Currency - replace "$" |
| Contract Time (days), "(     ) days" | ContractTimeChange | Number - replace blank in parentheses |
| Date of Substantial Completion | SubstantialCompletionDate | Date |
| Sum increased | IsContractSumIncreased | Boolean checkbox |
| Sum decreased | IsContractSumDecreased | Boolean checkbox |
| Time increased | IsContractTimeIncreased | Boolean checkbox |
| Time decreased | IsContractTimeDecreased | Boolean checkbox |
| Time unchanged | IsContractTimeUnchanged | Boolean checkbox |

## CRITICAL PATTERNS TO DETECT

1. **Parentheses with blanks** like `(     ) days` - This is a number field. Replace the blank spaces inside parentheses with a number token.
   - Example: `by (     ) days` → `by ([Number(Source=Attribute,Path=ContractTimeChange)]) days`

2. **Checkbox option text** like "increased", "decreased", "unchanged" - Each option gets its own boolean token.
   - Example: `be increased   be decreased` → each word becomes a boolean token

3. **Dollar sign cells** - Cells containing just "$" need the full currency token.
   - Example: `$` → `[Currency(Source=Attribute,Path=OriginalContractAmount)]`

Return JSON with this structure:
{{
  "document_summary": "Brief description of what this document is",
  "entity_type_detected": "Change Order / Contract / RFI / etc.",
  "injection_points": [
    {{
      "location_type": "paragraph" or "table_cell",
      "paragraph_index": <number if paragraph>,
      "table_index": <number if table>,
      "row_index": <number if table>,
      "cell_index": <number if table - CRITICAL: specify which cell gets the token>,
      "original_text": "The full text of the paragraph/cell being modified",
      "text_to_replace": "The specific text to replace (e.g., '$' or '(     )' or blank)",
      "kahua_field_path": "FieldPath from schema or mapping above",
      "injection_type": "currency|date|number|text|boolean|checkbox",
      "reasoning": "Why this needs injection and what field it maps to",
      "confidence": 0.0-1.0
    }}
  ],
  "warnings": ["Any concerns or ambiguities"],
  "suggestions": ["Recommendations for improving the template"]
}}

## RULES

1. For "$" cells: Replace the "$" with the full currency token
2. For checkbox options: Create separate injection points for each option (increased, decreased, unchanged)
3. For inline blanks "(     )": Replace with number token in parentheses
4. For "LABEL:" patterns: Inject after the colon
5. Always specify exact cell_index when targeting table cells
6. Use field paths from schema when available, otherwise use the mapping table above"""


def get_anthropic_client() -> Anthropic:
    """Get Anthropic client configured for Azure or direct."""
    endpoint = os.environ.get("AZURE_ENDPOINT", "").rstrip("/")
    api_key = os.environ.get("AZURE_KEY", "")
    
    if endpoint and api_key:
        # Azure-hosted Anthropic - use /anthropic path
        return Anthropic(
            base_url=f"{endpoint}/anthropic",
            api_key=api_key,
        )
    else:
        # Direct Anthropic API
        return Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))


async def _fetch_schema_for_entity(entity_def: str) -> Optional[Any]:
    """Fetch schema from schema service if available."""
    if get_entity_schema is None:
        return None
    try:
        return await get_entity_schema(entity_def)
    except Exception as e:
        log.warning(f"Could not fetch schema for {entity_def}: {e}")
        return None


def _build_schema_context(schema: Optional[Any], schema_fields: Optional[List[Dict[str, str]]]) -> str:
    """Build schema context string for LLM prompt."""
    if schema is not None and hasattr(schema, 'to_llm_context'):
        return f"\n=== ENTITY SCHEMA (USE THESE EXACT FIELD PATHS) ===\n{schema.to_llm_context()}\n"
    
    if schema_fields:
        field_list = "\n".join([
            f"- {f.get('path', f.get('name', ''))}: {f.get('format_hint', f.get('format', 'text'))} - {f.get('label', '')}"
            for f in schema_fields[:60]  # Limit to avoid token overflow
        ])
        return f"\nAvailable fields from schema (use these exact paths):\n{field_list}\n"
    
    return ""


def analyze_document_with_llm(
    doc_bytes: bytes,
    entity_def: str = "",
    schema_fields: Optional[List[Dict[str, str]]] = None,
) -> LLMAnalysisResult:
    """
    Use LLM to analyze document and identify all injection points.
    
    This is the primary analysis function - the LLM drives the process
    rather than relying on regex patterns.
    
    Args:
        doc_bytes: Raw DOCX file bytes
        entity_def: Target entity type (e.g., "kahua_AEC_ChangeOrder.ChangeOrder")
        schema_fields: Optional list of available fields from schema
        
    Returns:
        LLMAnalysisResult with identified injection points
    """
    # Run async version synchronously
    loop = asyncio.new_event_loop()
    try:
        return loop.run_until_complete(
            analyze_document_with_llm_async(doc_bytes, entity_def, schema_fields)
        )
    finally:
        loop.close()


async def analyze_document_with_llm_async(
    doc_bytes: bytes,
    entity_def: str = "",
    schema_fields: Optional[List[Dict[str, str]]] = None,
) -> LLMAnalysisResult:
    """
    Async version: Use LLM to analyze document with schema context.
    
    This function automatically fetches the entity schema to provide
    accurate field paths to the LLM.
    """
    try:
        # Extract document content
        content = extract_document_content(doc_bytes)
        
        # Resolve entity alias if needed
        if resolve_entity and entity_def:
            entity_def = resolve_entity(entity_def)
        
        # Fetch schema from service
        schema = None
        if entity_def:
            schema = await _fetch_schema_for_entity(entity_def)
        
        # Build schema context
        schema_context = _build_schema_context(schema, schema_fields)
        
        # Call LLM
        client = get_anthropic_client()
        model = os.environ.get("AZURE_DEPLOYMENT", "claude-sonnet-4-5")
        
        response = client.messages.create(
            model=model,
            system=INJECTION_ANALYSIS_SYSTEM,
            messages=[
                {
                    "role": "user",
                    "content": INJECTION_ANALYSIS_PROMPT.format(
                        entity_def=entity_def or "Unknown - please infer from content",
                        document_content=content.raw_text,
                        schema_context=schema_context,
                    )
                }
            ],
            max_tokens=4000,
            temperature=0.2,  # Low temperature for consistency
        )
        
        # Parse response
        response_text = response.content[0].text
        
        # Extract JSON from response
        json_str = response_text
        if "```json" in response_text:
            json_str = response_text.split("```json")[1].split("```")[0]
        elif "```" in response_text:
            json_str = response_text.split("```")[1].split("```")[0]
        
        analysis = json.loads(json_str)
        
        # Convert to InjectionPoints
        injection_points = []
        for point in analysis.get("injection_points", []):
            ip = InjectionPoint(
                location_type=point.get("location_type", "paragraph"),
                paragraph_index=point.get("paragraph_index"),
                table_index=point.get("table_index"),
                row_index=point.get("row_index"),
                cell_index=point.get("cell_index"),
                original_text=point.get("original_text", ""),
                text_to_replace=point.get("text_to_replace", ""),
                kahua_field_path=point.get("kahua_field_path", ""),
                injection_type=InjectionType(point.get("injection_type", "text")),
                condition_context=point.get("condition_context"),
                reasoning=point.get("reasoning", ""),
                confidence=point.get("confidence", 0.8),
            )
            
            # Generate the actual Kahua token
            ip.token = _generate_token(ip)
            injection_points.append(ip)
        
        return LLMAnalysisResult(
            success=True,
            injection_points=injection_points,
            document_summary=analysis.get("document_summary", ""),
            entity_type_detected=analysis.get("entity_type_detected", ""),
            warnings=analysis.get("warnings", []),
            suggestions=analysis.get("suggestions", []),
        )
        
    except json.JSONDecodeError as e:
        log.error(f"Failed to parse LLM response: {e}")
        return LLMAnalysisResult(
            success=False,
            error=f"Failed to parse LLM analysis: {e}"
        )
    except Exception as e:
        log.error(f"LLM analysis failed: {e}", exc_info=True)
        return LLMAnalysisResult(
            success=False,
            error=str(e)
        )


def _generate_token(point: InjectionPoint) -> str:
    """Generate the appropriate Kahua token for an injection point."""
    field_path = to_kahua_path(point.kahua_field_path)
    
    if point.injection_type == InjectionType.CURRENCY:
        return build_currency_token(field_path)
    elif point.injection_type == InjectionType.DATE:
        return build_date_token(field_path)
    elif point.injection_type == InjectionType.NUMBER:
        return build_number_token(field_path)
    elif point.injection_type in (InjectionType.BOOLEAN, InjectionType.CHECKBOX):
        # For checkboxes, use checkbox symbols
        return build_boolean_token(field_path, true_value="☒", false_value="☐")
    elif point.injection_type == InjectionType.CONDITIONAL:
        # For conditional selections with multiple options
        # This might need special handling based on condition_context
        return build_boolean_token(field_path, true_value="☒", false_value="☐")
    else:
        return build_attribute_token(field_path)


# ============== Token Injection ==============

def inject_tokens_from_analysis(
    doc_bytes: bytes,
    analysis: LLMAnalysisResult,
) -> Tuple[bytes, List[str]]:
    """
    Inject tokens into document based on LLM analysis results.
    
    Args:
        doc_bytes: Original document bytes
        analysis: LLM analysis result with injection points
        
    Returns:
        Tuple of (modified document bytes, list of changes made)
    """
    doc = Document(io.BytesIO(doc_bytes))
    changes_made = []
    
    # Group injection points by location for efficient processing
    para_points = [p for p in analysis.injection_points if p.location_type == "paragraph"]
    table_points = [p for p in analysis.injection_points if p.location_type == "table_cell"]
    
    # Process paragraphs (in reverse order to preserve indices)
    for point in sorted(para_points, key=lambda p: p.paragraph_index or 0, reverse=True):
        if point.paragraph_index is not None and point.paragraph_index < len(doc.paragraphs):
            para = doc.paragraphs[point.paragraph_index]
            success = _inject_in_paragraph(para, point)
            if success:
                changes_made.append(f"Injected {point.token} for '{point.kahua_field_path}'")
    
    # Process table cells
    for point in table_points:
        if (point.table_index is not None and 
            point.table_index < len(doc.tables) and
            point.row_index is not None and
            point.cell_index is not None):
            
            table = doc.tables[point.table_index]
            if point.row_index < len(table.rows):
                row = table.rows[point.row_index]
                if point.cell_index < len(row.cells):
                    cell = row.cells[point.cell_index]
                    if cell.paragraphs:
                        success = _inject_in_paragraph(cell.paragraphs[0], point)
                        if success:
                            changes_made.append(f"Injected {point.token} for '{point.kahua_field_path}' in table")
    
    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output.getvalue(), changes_made


def _inject_in_paragraph(para: Paragraph, point: InjectionPoint) -> bool:
    """
    Inject a token into a paragraph based on the injection point spec.
    
    Strategy:
    1. Try to find exact text_to_replace in a single run (fast path)
    2. If not found in single run, search full paragraph text and replace across runs
    3. Handle special cases for currency, checkbox, parentheses
    4. Only use append fallback for certain injection types
    """
    import re
    
    try:
        full_text = para.text
        text_to_replace = point.text_to_replace
        
        # If we have specific text to replace
        if text_to_replace and text_to_replace in full_text:
            # Fast path: try to find in single run
            for run in para.runs:
                if text_to_replace in run.text:
                    run.text = run.text.replace(text_to_replace, point.token, 1)
                    return True
            
            # Slow path: text spans multiple runs - use the cross-run replace
            success = _replace_text_across_runs(para, text_to_replace, point.token)
            if success:
                return True
        
        # Handle special cases
        
        # Case: Currency - replace "$" or just "$" cell
        if point.injection_type == InjectionType.CURRENCY:
            # Check if cell is just "$"
            if full_text.strip() == '$':
                for run in para.runs:
                    if '$' in run.text:
                        run.text = run.text.replace('$', point.token, 1)
                        return True
            # Check if ends with "$"
            if full_text.rstrip().endswith('$'):
                for run in para.runs:
                    if '$' in run.text:
                        run.text = run.text.replace('$', point.token, 1)
                        return True
        
        # Case: Checkbox/Boolean - look for checkbox characters OR the option text
        if point.injection_type in (InjectionType.CHECKBOX, InjectionType.BOOLEAN):
            # Look for checkbox Unicode characters
            checkbox_chars = ['☐', '□', '■', '▢', '☑', '☒', '◯', '●']
            for run in para.runs:
                for char in checkbox_chars:
                    if char in run.text:
                        run.text = run.text.replace(char, point.token, 1)
                        return True
            
            # For boolean fields, if text_to_replace wasn't found but we have the field,
            # look for the option text in the paragraph (might have extra whitespace)
            if text_to_replace:
                # Try flexible matching - look for the word with possible surrounding spaces
                pattern = re.compile(r'\s*' + re.escape(text_to_replace) + r'\s*', re.IGNORECASE)
                for run in para.runs:
                    match = pattern.search(run.text)
                    if match:
                        # Replace just the word (preserve some spacing)
                        run.text = pattern.sub(f' {point.token} ', run.text, count=1)
                        return True
        
        # Case: Parentheses with blank "( )" or "(   )" or "(     )"
        paren_pattern = r'\(\s{2,}\)'  # At least 2 spaces
        if re.search(paren_pattern, full_text):
            for run in para.runs:
                if re.search(paren_pattern, run.text):
                    run.text = re.sub(paren_pattern, f'({point.token})', run.text, count=1)
                    return True
        
        # Case: Date fields - look for blank underscores like "____"
        if point.injection_type == InjectionType.DATE:
            underscore_pattern = r'_{3,}'  # 3+ underscores
            if re.search(underscore_pattern, full_text):
                for run in para.runs:
                    if re.search(underscore_pattern, run.text):
                        run.text = re.sub(underscore_pattern, point.token, run.text, count=1)
                        return True
        
        # Only fallback to append for text fields where we expect free-form text
        # Don't fallback for currency/boolean/checkbox - those should replace specific text
        if point.injection_type in (InjectionType.TEXT, InjectionType.DATE, InjectionType.NUMBER):
            if para.runs:
                para.runs[-1].text = para.runs[-1].text.rstrip() + " " + point.token
            else:
                para.add_run(" " + point.token)
            return True
        
        # For other types, log a warning but still report success if we tried
        log.warning(f"Could not find exact injection location for {point.kahua_field_path}")
        return False
        
    except Exception as e:
        log.error(f"Failed to inject token: {e}")
        return False


def _replace_text_across_runs(para: Paragraph, old_text: str, new_text: str) -> bool:
    """
    Replace text that may span multiple runs in a paragraph.
    This preserves formatting of the first run that contains part of the text.
    """
    # Build a map of character position to run
    full_text = para.text
    if old_text not in full_text:
        return False
    
    # Find the position of the text to replace
    start_pos = full_text.find(old_text)
    end_pos = start_pos + len(old_text)
    
    # Map character positions to runs
    char_to_run = []
    for run in para.runs:
        for _ in run.text:
            char_to_run.append(run)
    
    if start_pos >= len(char_to_run):
        return False
    
    # Find which runs contain parts of the text
    affected_runs = set()
    for pos in range(start_pos, min(end_pos, len(char_to_run))):
        affected_runs.add(char_to_run[pos])
    
    # Simple case: all in one run (shouldn't reach here but handle it)
    if len(affected_runs) == 1:
        run = list(affected_runs)[0]
        run.text = run.text.replace(old_text, new_text, 1)
        return True
    
    # Complex case: spans multiple runs
    # Strategy: Rebuild the full text with replacement, then redistribute to runs
    new_full_text = full_text[:start_pos] + new_text + full_text[end_pos:]
    _replace_paragraph_text(para, new_full_text)
    return True


def _replace_paragraph_text(para: Paragraph, new_text: str):
    """Replace all text in a paragraph while preserving first run's formatting."""
    if not para.runs:
        para.add_run(new_text)
        return
    
    # Store first run's format
    first_run = para.runs[0]
    
    # Clear all runs
    for i in range(len(para.runs) - 1, -1, -1):
        para.runs[i].text = ""
    
    # Set new text on first run
    first_run.text = new_text


# ============== High-Level API ==============

def analyze_and_inject_with_llm(
    doc_bytes: bytes,
    entity_def: str = "",
    schema_fields: Optional[List[Dict[str, str]]] = None,
    auto_inject: bool = True,
) -> Dict[str, Any]:
    """
    Main entry point: Analyze document with LLM and optionally inject tokens.
    
    Args:
        doc_bytes: The uploaded DOCX file bytes
        entity_def: Target Kahua entity definition
        schema_fields: Optional list of available schema fields
        auto_inject: If True, automatically inject tokens
        
    Returns:
        Dict with analysis results and optionally modified document
    """
    # Analyze with LLM
    analysis = analyze_document_with_llm(doc_bytes, entity_def, schema_fields)
    
    result = {
        "success": analysis.success,
        "analysis": {
            "document_summary": analysis.document_summary,
            "entity_type_detected": analysis.entity_type_detected,
            "injection_points": [
                {
                    "location_type": p.location_type,
                    "original_text": p.original_text,
                    "text_to_replace": p.text_to_replace,
                    "kahua_field_path": p.kahua_field_path,
                    "injection_type": p.injection_type.value,
                    "token": p.token,
                    "reasoning": p.reasoning,
                    "confidence": p.confidence,
                }
                for p in analysis.injection_points
            ],
            "warnings": analysis.warnings,
            "suggestions": analysis.suggestions,
        },
        "error": analysis.error,
    }
    
    # Inject if requested and analysis succeeded
    if auto_inject and analysis.success and analysis.injection_points:
        modified_doc, changes = inject_tokens_from_analysis(doc_bytes, analysis)
        result["injection"] = {
            "success": True,
            "tokens_injected": len(changes),
            "changes_made": changes,
        }
        result["modified_document"] = modified_doc
    
    return result


# ============== Async Version for API Use ==============

async def analyze_document_with_llm_async(
    doc_bytes: bytes,
    entity_def: str = "",
    schema_fields: Optional[List[Dict[str, str]]] = None,
) -> LLMAnalysisResult:
    """Async version of analyze_document_with_llm for use in FastAPI endpoints."""
    from anthropic import AsyncAnthropic
    
    try:
        content = extract_document_content(doc_bytes)
        
        schema_context = ""
        if schema_fields:
            field_list = "\n".join([
                f"- {f.get('path', f.get('name', ''))}: {f.get('format', 'text')} - {f.get('label', '')}"
                for f in schema_fields[:60]
            ])
            schema_context = f"\nAvailable fields from schema:\n{field_list}\n"
        
        # Get async client - use /anthropic path for Azure
        endpoint = os.environ.get("AZURE_ENDPOINT", "").rstrip("/")
        api_key = os.environ.get("AZURE_KEY", "")
        deployment = os.environ.get("AZURE_DEPLOYMENT", "claude-sonnet-4-5")
        
        if endpoint and api_key:
            client = AsyncAnthropic(
                base_url=f"{endpoint}/anthropic",
                api_key=api_key,
            )
        else:
            client = AsyncAnthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))
        
        response = await client.messages.create(
            model=deployment,
            system=INJECTION_ANALYSIS_SYSTEM,
            messages=[
                {
                    "role": "user",
                    "content": INJECTION_ANALYSIS_PROMPT.format(
                        entity_def=entity_def or "Unknown - please infer from content",
                        document_content=content.raw_text,
                        schema_context=schema_context,
                    )
                }
            ],
            max_tokens=8000,  # Increased for complex documents
            temperature=0.2,
        )
        
        response_text = response.content[0].text
        json_str = response_text
        if "```json" in response_text:
            json_str = response_text.split("```json")[1].split("```")[0]
        elif "```" in response_text:
            json_str = response_text.split("```")[1].split("```")[0]
        
        # Handle truncated JSON - try to fix common issues
        json_str = json_str.strip()
        if not json_str.endswith("}"):
            # Try to close the JSON properly
            log.warning("JSON appears truncated, attempting to fix")
            # Count open braces/brackets
            open_braces = json_str.count("{") - json_str.count("}")
            open_brackets = json_str.count("[") - json_str.count("]")
            json_str += "]" * open_brackets + "}" * open_braces
        
        analysis = json.loads(json_str)
        
        injection_points = []
        for point in analysis.get("injection_points", []):
            ip = InjectionPoint(
                location_type=point.get("location_type", "paragraph"),
                paragraph_index=point.get("paragraph_index"),
                table_index=point.get("table_index"),
                row_index=point.get("row_index"),
                cell_index=point.get("cell_index"),
                original_text=point.get("original_text", ""),
                text_to_replace=point.get("text_to_replace", ""),
                kahua_field_path=point.get("kahua_field_path", ""),
                injection_type=InjectionType(point.get("injection_type", "text")),
                condition_context=point.get("condition_context"),
                reasoning=point.get("reasoning", ""),
                confidence=point.get("confidence", 0.8),
            )
            ip.token = _generate_token(ip)
            injection_points.append(ip)
        
        return LLMAnalysisResult(
            success=True,
            injection_points=injection_points,
            document_summary=analysis.get("document_summary", ""),
            entity_type_detected=analysis.get("entity_type_detected", ""),
            warnings=analysis.get("warnings", []),
            suggestions=analysis.get("suggestions", []),
        )
        
    except Exception as e:
        log.error(f"Async LLM analysis failed: {e}", exc_info=True)
        return LLMAnalysisResult(success=False, error=str(e))


# ============== CLI for Testing ==============

if __name__ == "__main__":
    import sys
    from pathlib import Path
    
    logging.basicConfig(level=logging.INFO)
    
    if len(sys.argv) < 2:
        print("Usage: python llm_injection_analyzer.py <path_to_docx> [entity_def]")
        sys.exit(1)
    
    doc_path = Path(sys.argv[1])
    entity_def = sys.argv[2] if len(sys.argv) > 2 else ""
    
    if not doc_path.exists():
        print(f"File not found: {doc_path}")
        sys.exit(1)
    
    with open(doc_path, 'rb') as f:
        doc_bytes = f.read()
    
    print("\n=== Analyzing document with LLM ===\n")
    result = analyze_and_inject_with_llm(doc_bytes, entity_def, auto_inject=False)
    
    print(f"Document Summary: {result['analysis']['document_summary']}")
    print(f"Entity Type: {result['analysis']['entity_type_detected']}")
    print(f"\nIdentified {len(result['analysis']['injection_points'])} injection points:\n")
    
    for i, point in enumerate(result['analysis']['injection_points'], 1):
        print(f"{i}. {point['kahua_field_path']} ({point['injection_type']})")
        print(f"   Replace: '{point['text_to_replace']}'")
        print(f"   Token: {point['token']}")
        print(f"   Reason: {point['reasoning']}")
        print()
    
    if result['analysis']['warnings']:
        print("Warnings:")
        for w in result['analysis']['warnings']:
            print(f"  ⚠ {w}")
    
    if result['analysis']['suggestions']:
        print("\nSuggestions:")
        for s in result['analysis']['suggestions']:
            print(f"  • {s}")
