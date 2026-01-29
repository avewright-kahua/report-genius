"""
Markdown-based Portable View Renderer
Templates are Markdown with Jinja2 syntax, converted to DOCX.
"""

import os
import re
import uuid
import tempfile
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple
from datetime import datetime
from jinja2 import Environment, BaseLoader, DebugUndefined

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Directories
TEMPLATES_DIR = Path(__file__).parent / "pv_templates"
REPORTS_DIR = Path(__file__).parent / "reports"
REPORTS_DIR.mkdir(exist_ok=True)

# In-memory preview cache for agent workflow
_preview_cache: Dict[str, Dict[str, Any]] = {}


def list_md_templates() -> Dict[str, Any]:
    """
    List available markdown templates.
    
    Returns:
        Dict with list of template info
    """
    templates = []
    
    if TEMPLATES_DIR.exists():
        for f in TEMPLATES_DIR.glob("*.md"):
            templates.append({
                "id": f.stem,  # filename without extension
                "name": f.stem.replace("_", " ").title(),
                "path": str(f),
            })
    
    return {
        "status": "ok",
        "templates": templates,
        "count": len(templates)
    }


def preview_portable_view(
    template_id: str,
    data: Dict[str, Any]
) -> Dict[str, Any]:
    """
    Preview a portable view as rendered Markdown (no DOCX conversion).
    Caches the result for finalization.
    
    Args:
        template_id: Template name (e.g., "contract", "rfi")
        data: Entity data dictionary
    
    Returns:
        Dict with preview_id and rendered markdown
    """
    # Find template
    template_path = TEMPLATES_DIR / f"{template_id}.md"
    
    if not template_path.exists():
        # Try case-insensitive search
        for f in TEMPLATES_DIR.glob("*.md"):
            if f.stem.lower() == template_id.lower():
                template_path = f
                break
    
    if not template_path.exists():
        return {"status": "error", "message": f"Template '{template_id}' not found"}
    
    try:
        template_str = template_path.read_text(encoding='utf-8')
        rendered = render_md_template(template_str, data)
        
        # Cache for finalization
        preview_id = f"preview-{uuid.uuid4().hex[:8]}"
        _preview_cache[preview_id] = {
            "template_id": template_id,
            "rendered_markdown": rendered,
            "data": data,
            "created_at": datetime.now().isoformat(),
        }
        
        return {
            "status": "ok",
            "preview_id": preview_id,
            "rendered_markdown": rendered,
            "message": "Preview generated. Display the markdown above to the user."
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


def finalize_portable_view(preview_id: str, output_name: str = None) -> Dict[str, Any]:
    """
    Finalize a previewed portable view and generate DOCX.
    
    Args:
        preview_id: Preview ID from preview_portable_view
        output_name: Optional custom output name
    
    Returns:
        Dict with download URL
    """
    if preview_id not in _preview_cache:
        return {"status": "error", "message": f"Preview {preview_id} not found or expired"}
    
    preview = _preview_cache[preview_id]
    rendered_md = preview["rendered_markdown"]
    
    if not output_name:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_name = f"{preview['template_id']}_{timestamp}"
    
    output_path = REPORTS_DIR / f"{output_name}.docx"
    
    try:
        md_to_docx(rendered_md, output_path)
        
        base_url = os.getenv("REPORT_BASE_URL", "http://localhost:8000")
        download_url = f"{base_url}/reports/{output_name}.docx"
        
        # Clean up cache
        del _preview_cache[preview_id]
        
        return {
            "status": "ok",
            "filename": f"{output_name}.docx",
            "download_url": download_url,
            "message": f"Document generated! Download: {download_url}"
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


class SilentUndefined(DebugUndefined):
    """Returns empty string for missing variables instead of error."""
    def __str__(self):
        return ""
    
    def __iter__(self):
        return iter([])
    
    def __bool__(self):
        return False


def create_jinja_env() -> Environment:
    """Create Jinja2 environment with custom filters."""
    env = Environment(
        loader=BaseLoader(),
        undefined=SilentUndefined,
        autoescape=False
    )
    
    # Custom filters
    env.filters['currency'] = lambda v, symbol='$': f"{symbol}{v:,.2f}" if v else "-"
    env.filters['date'] = lambda v, fmt='%b %d, %Y': _format_date(v, fmt)
    env.filters['datetime'] = lambda v, fmt='%b %d, %Y %I:%M %p': _format_date(v, fmt)
    env.filters['default'] = lambda v, d='-': v if v else d
    env.filters['yesno'] = lambda v: 'Yes' if v else 'No'
    env.filters['status_badge'] = lambda v: f"**{v}**" if v else "-"
    
    return env


def _format_date(value: Any, fmt: str) -> str:
    """Format a date value."""
    if not value:
        return "-"
    if isinstance(value, str):
        try:
            from dateutil import parser
            dt = parser.parse(value)
            return dt.strftime(fmt)
        except:
            return value
    return str(value)


def get_nested(data: Dict[str, Any], path: str, default: Any = None) -> Any:
    """Get nested value from dict using dot notation."""
    if not path or not data:
        return default
    
    parts = path.split('.')
    current = data
    
    for part in parts:
        if current is None:
            return default
        if isinstance(current, dict):
            # Case-insensitive lookup
            found = None
            for k, v in current.items():
                if k.lower() == part.lower():
                    found = v
                    break
            if found is None:
                return default
            current = found
        else:
            return default
    
    return current if current is not None else default


def prepare_data(data: Dict[str, Any]) -> Dict[str, Any]:
    """Prepare data dict with helper access methods."""
    class DataWrapper(dict):
        def __getattr__(self, key):
            val = self.get(key)
            if isinstance(val, dict):
                return DataWrapper(val)
            if isinstance(val, list):
                return [DataWrapper(v) if isinstance(v, dict) else v for v in val]
            return val if val is not None else ""
        
        def __getitem__(self, key):
            val = super().get(key)
            if isinstance(val, dict):
                return DataWrapper(val)
            return val
    
    return DataWrapper(data)


def render_md_template(template_str: str, data: Dict[str, Any]) -> str:
    """Render a Markdown template with data."""
    env = create_jinja_env()
    
    # Wrap data for easy access
    wrapped_data = prepare_data(data)
    
    # Add utility values
    wrapped_data['_now'] = datetime.now()
    wrapped_data['_today'] = datetime.now().strftime('%B %d, %Y')
    
    # Compile and render
    template = env.from_string(template_str)
    return template.render(**wrapped_data)


def md_to_docx(md_content: str, output_path: Path, style_config: Optional[Dict] = None) -> Path:
    """Convert Markdown to DOCX using python-docx (no pandoc needed)."""
    doc = Document()
    
    # Default style config
    style = style_config or {
        'primary_color': '#0f172a',
        'font': 'Calibri',
        'title_size': 18,
        'heading_size': 12,
        'body_size': 10,
    }
    
    # Parse and render markdown
    lines = md_content.strip().split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Heading 1: # Title
        if line.startswith('# '):
            _add_heading(doc, line[2:].strip(), 1, style)
            i += 1
            continue
        
        # Heading 2: ## Section
        if line.startswith('## '):
            _add_heading(doc, line[3:].strip(), 2, style)
            i += 1
            continue
        
        # Heading 3: ### Subsection
        if line.startswith('### '):
            _add_heading(doc, line[4:].strip(), 3, style)
            i += 1
            continue
        
        # Horizontal rule: ---
        if line.strip() in ('---', '***', '___'):
            # Add subtle spacing instead of ugly line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        # Table: | ... |
        if line.strip().startswith('|'):
            # Collect all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            _add_table(doc, table_lines, style)
            continue
        
        # Empty line
        if not line.strip():
            i += 1
            continue
        
        # Regular paragraph (may contain **bold**, *italic*)
        _add_paragraph(doc, line, style)
        i += 1
    
    doc.save(output_path)
    return output_path


def _add_heading(doc: Document, text: str, level: int, style: Dict):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = style.get('font', 'Calibri')
        if level == 1:
            run.font.size = Pt(style.get('title_size', 18))
            run.font.color.rgb = RGBColor(*_hex_to_rgb(style.get('primary_color', '#0f172a')))
        else:
            run.font.size = Pt(style.get('heading_size', 12))
            run.font.color.rgb = RGBColor(*_hex_to_rgb(style.get('primary_color', '#0f172a')))


def _add_paragraph(doc: Document, text: str, style: Dict):
    """Add a paragraph with inline formatting."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    
    # Parse inline formatting: **bold**, *italic*
    _add_formatted_text(para, text, style)


def _add_formatted_text(para, text: str, style: Dict):
    """Add text with **bold** and *italic* formatting."""
    # Pattern for **bold** and *italic*
    pattern = r'(\*\*.*?\*\*|\*.*?\*|[^*]+)'
    parts = re.findall(pattern, text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            run = para.add_run(part)
        
        run.font.name = style.get('font', 'Calibri')
        run.font.size = Pt(style.get('body_size', 10))


def _add_table(doc: Document, lines: List[str], style: Dict):
    """Add a table from markdown table lines."""
    # Parse rows
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        # Skip separator row (|---|---|)
        if cells and not all(re.match(r'^[-:]+$', c) for c in cells):
            rows.append(cells)
    
    if not rows:
        return
    
    # Detect alignment from separator row
    alignments = []
    for line in lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        if cells and all(re.match(r'^[-:]+$', c) for c in cells):
            for c in cells:
                if c.startswith(':') and c.endswith(':'):
                    alignments.append('center')
                elif c.endswith(':'):
                    alignments.append('right')
                else:
                    alignments.append('left')
            break
    
    # Create table
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    
    # Fill cells
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= num_cols:
                break
            cell = row.cells[col_idx]
            
            # Clear and add formatted text
            cell.text = ''
            para = cell.paragraphs[0]
            _add_formatted_text(para, cell_text, style)
            
            # Apply alignment
            if alignments and col_idx < len(alignments):
                if alignments[col_idx] == 'right':
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alignments[col_idx] == 'center':
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Style first row as header
            if row_idx == 0:
                for run in para.runs:
                    run.bold = True
                _set_cell_shading(cell, style.get('primary_color', '#0f172a'))
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
    
    doc.add_paragraph()  # Spacing after table


def _set_cell_shading(cell, color: str):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color.lstrip('#'))
    cell._tc.get_or_add_tcPr().append(shading)


def _hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
    """Convert hex to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def render_portable_view(
    template_path: Path,
    data: Dict[str, Any],
    output_name: Optional[str] = None,
    style_config: Optional[Dict] = None
) -> Path:
    """
    Full render pipeline: Load MD template -> Render with data -> Convert to DOCX.
    
    Args:
        template_path: Path to .md template file
        data: Entity data dictionary
        output_name: Optional output filename (without extension)
        style_config: Optional style overrides
    
    Returns:
        Path to generated .docx file
    """
    # Load template
    template_str = template_path.read_text(encoding='utf-8')
    
    # Render markdown with data
    rendered_md = render_md_template(template_str, data)
    
    # Generate output path
    if not output_name:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_name = f"{template_path.stem}_{timestamp}"
    
    output_path = REPORTS_DIR / f"{output_name}.docx"
    
    # Convert to DOCX
    md_to_docx(rendered_md, output_path, style_config)
    
    return output_path


# Quick test function
if __name__ == "__main__":
    # Test with sample data
    sample_template = """
# {{ Number }} - {{ Description }}

| Field | Value |
|-------|-------|
| **Contractor** | {{ ContractorCompany.ShortLabel | default }} |
| **Status** | {{ WorkflowStatus | default }} |
| **Type** | {{ Type | default }} |

## Financials

| Category | Amount |
|----------|-------:|
| Original Value | {{ OriginalContractValue | currency }} |
| Approved | {{ ContractApprovedTotalValue | currency }} |
| Pending | {{ ContractPendingTotalValue | currency }} |

{% if ScopeOfWork %}
## Scope of Work

{{ ScopeOfWork }}
{% endif %}

{% if ChangeOrders %}
## Change Orders

| # | Description | Status | Amount |
|---|-------------|--------|-------:|
{% for co in ChangeOrders %}
| {{ co.Number }} | {{ co.Description }} | {{ co.Status.Name | default }} | {{ co.TotalValue | currency }} |
{% endfor %}
{% endif %}

---
*Generated {{ _today }}*
"""
    
    test_data = {
        'Number': '0001',
        'Description': 'Example Contract',
        'ContractorCompany': {'ShortLabel': 'ACME Contractors'},
        'WorkflowStatus': 'Approved',
        'Type': 'Fixed Price',
        'OriginalContractValue': 50000,
        'ContractApprovedTotalValue': 52000,
        'ContractPendingTotalValue': 55000,
        'ScopeOfWork': 'Complete renovation of Building A.',
        'ChangeOrders': [
            {'Number': 'CO-001', 'Description': 'Extra outlets', 'Status': {'Name': 'Approved'}, 'TotalValue': 1500},
            {'Number': 'CO-002', 'Description': 'HVAC upgrade', 'Status': {'Name': 'Pending'}, 'TotalValue': 3500},
        ]
    }
    
    rendered = render_md_template(sample_template, test_data)
    print(rendered)
    print("\n--- Converting to DOCX ---")
    
    # Save template and render
    template_path = Path("pv_templates/contract_simple.md")
    template_path.parent.mkdir(exist_ok=True)
    template_path.write_text(sample_template, encoding='utf-8')
    
    output = render_portable_view(template_path, test_data, "contract_md_test")
    print(f"Generated: {output}")
