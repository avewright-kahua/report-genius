"""
Word Template Builder
Direct Word document creation with full formatting control.
No markdown intermediary - full access to OOXML capabilities.
"""

import io
import re
import json
import uuid
import logging
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple, Union, Callable
from dataclasses import dataclass, field, asdict
from datetime import datetime
from enum import Enum

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

log = logging.getLogger("word_template_builder")

# Storage paths
TEMPLATES_DIR = Path(__file__).parent / "word_templates"
TEMPLATES_DIR.mkdir(exist_ok=True)
OUTPUTS_DIR = Path(__file__).parent / "reports"
OUTPUTS_DIR.mkdir(exist_ok=True)


# ============== Enums ==============

class Orientation(str, Enum):
    PORTRAIT = "portrait"
    LANDSCAPE = "landscape"


class Alignment(str, Enum):
    LEFT = "left"
    CENTER = "center"
    RIGHT = "right"
    JUSTIFY = "justify"


class VerticalAlignment(str, Enum):
    TOP = "top"
    CENTER = "center"
    BOTTOM = "bottom"


class BorderStyle(str, Enum):
    NONE = "none"
    SINGLE = "single"
    DOUBLE = "double"
    THICK = "thick"
    DOTTED = "dotted"
    DASHED = "dashed"


class FieldFormat(str, Enum):
    TEXT = "text"
    NUMBER = "number"
    CURRENCY = "currency"
    DATE = "date"
    DATETIME = "datetime"
    PERCENT = "percent"
    BOOLEAN = "boolean"
    RICH_TEXT = "rich_text"


# ============== Style Dataclasses ==============

@dataclass
class FontStyle:
    """Font styling options."""
    name: str = "Calibri"
    size: float = 11  # points
    bold: bool = False
    italic: bool = False
    underline: bool = False
    color: str = "#000000"
    highlight: Optional[str] = None
    
    def apply(self, run):
        """Apply font style to a run."""
        run.font.name = self.name
        run.font.size = Pt(self.size)
        run.font.bold = self.bold
        run.font.italic = self.italic
        run.font.underline = self.underline
        run.font.color.rgb = RGBColor(*hex_to_rgb(self.color))


@dataclass
class ParagraphStyle:
    """Paragraph styling options."""
    alignment: Alignment = Alignment.LEFT
    line_spacing: float = 1.15
    space_before: float = 0  # points
    space_after: float = 6  # points
    first_line_indent: float = 0  # inches
    left_indent: float = 0  # inches
    right_indent: float = 0  # inches
    
    def apply(self, paragraph):
        """Apply paragraph style to a paragraph."""
        alignment_map = {
            Alignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
            Alignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
            Alignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
            Alignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        paragraph.alignment = alignment_map.get(self.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        paragraph.paragraph_format.line_spacing = self.line_spacing
        paragraph.paragraph_format.space_before = Pt(self.space_before)
        paragraph.paragraph_format.space_after = Pt(self.space_after)
        if self.first_line_indent:
            paragraph.paragraph_format.first_line_indent = Inches(self.first_line_indent)
        if self.left_indent:
            paragraph.paragraph_format.left_indent = Inches(self.left_indent)
        if self.right_indent:
            paragraph.paragraph_format.right_indent = Inches(self.right_indent)


@dataclass
class BorderDef:
    """Border definition."""
    style: BorderStyle = BorderStyle.SINGLE
    width: float = 0.5  # points
    color: str = "#000000"
    space: float = 0  # points


@dataclass
class CellStyle:
    """Table cell styling."""
    background: Optional[str] = None
    vertical_align: VerticalAlignment = VerticalAlignment.TOP
    padding_top: float = 2  # points
    padding_bottom: float = 2
    padding_left: float = 4
    padding_right: float = 4
    border_top: Optional[BorderDef] = None
    border_bottom: Optional[BorderDef] = None
    border_left: Optional[BorderDef] = None
    border_right: Optional[BorderDef] = None


@dataclass
class PageMargins:
    """Page margin settings."""
    top: float = 1.0  # inches
    bottom: float = 1.0
    left: float = 1.0
    right: float = 1.0
    header: float = 0.5
    footer: float = 0.5


@dataclass
class DocumentTheme:
    """Document-wide styling theme."""
    name: str = "Default"
    
    # Colors
    primary_color: str = "#1a365d"
    secondary_color: str = "#3182ce"
    accent_color: str = "#38b2ac"
    text_color: str = "#1a202c"
    muted_color: str = "#718096"
    
    # Typography
    heading_font: str = "Calibri Light"
    body_font: str = "Calibri"
    title_size: float = 24
    heading1_size: float = 18
    heading2_size: float = 14
    heading3_size: float = 12
    body_size: float = 11
    small_size: float = 9
    
    # Table styling
    table_header_bg: str = "#1a365d"
    table_header_fg: str = "#ffffff"
    table_stripe_bg: str = "#f7fafc"
    table_border_color: str = "#e2e8f0"
    
    # Page setup
    orientation: Orientation = Orientation.PORTRAIT
    margins: PageMargins = field(default_factory=PageMargins)


# ============== Utility Functions ==============

def hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def resolve_data_path(data: Dict[str, Any], path: str) -> Any:
    """Resolve a dot-notation path in data dictionary (case-insensitive)."""
    if not path or not data:
        return None
    
    parts = path.split('.')
    current = data
    
    for part in parts:
        if current is None:
            return None
        
        # Handle array indexing: Items[0]
        match = re.match(r'(\w+)\[(\d+)\]', part)
        if match:
            key, idx = match.groups()
            current = get_case_insensitive(current, key)
            if isinstance(current, list) and int(idx) < len(current):
                current = current[int(idx)]
            else:
                return None
        elif isinstance(current, dict):
            current = get_case_insensitive(current, part)
        else:
            return None
    
    return current


def get_case_insensitive(data: Dict[str, Any], key: str) -> Any:
    """Get value from dict with case-insensitive key matching."""
    if not isinstance(data, dict):
        return None
    if key in data:
        return data[key]
    key_lower = key.lower()
    for k, v in data.items():
        if k.lower() == key_lower:
            return v
    return None


def format_value(value: Any, fmt: FieldFormat, options: Dict[str, Any] = None) -> str:
    """Format a value according to field format."""
    if value is None:
        return ""
    
    options = options or {}
    
    if fmt == FieldFormat.CURRENCY:
        try:
            num = float(value)
            prefix = options.get("prefix", "$")
            decimals = options.get("decimals", 2)
            return f"{prefix}{num:,.{decimals}f}"
        except (ValueError, TypeError):
            return str(value)
    
    elif fmt == FieldFormat.NUMBER:
        try:
            num = float(value)
            decimals = options.get("decimals", 0)
            return f"{num:,.{decimals}f}"
        except (ValueError, TypeError):
            return str(value)
    
    elif fmt == FieldFormat.PERCENT:
        try:
            num = float(value)
            decimals = options.get("decimals", 1)
            return f"{num:.{decimals}f}%"
        except (ValueError, TypeError):
            return str(value)
    
    elif fmt == FieldFormat.DATE:
        if isinstance(value, str):
            try:
                from dateutil import parser
                dt = parser.parse(value)
                return dt.strftime(options.get("format", "%m/%d/%Y"))
            except:
                return value[:10] if len(value) >= 10 else value
        return str(value)
    
    elif fmt == FieldFormat.DATETIME:
        if isinstance(value, str):
            try:
                from dateutil import parser
                dt = parser.parse(value)
                return dt.strftime(options.get("format", "%m/%d/%Y %I:%M %p"))
            except:
                return value
        return str(value)
    
    elif fmt == FieldFormat.BOOLEAN:
        return "Yes" if value else "No"
    
    return str(value)


def interpolate_template(template: str, data: Dict[str, Any]) -> str:
    """Replace {field.path} placeholders with values."""
    def replace_match(match):
        path = match.group(1)
        value = resolve_data_path(data, path)
        return str(value) if value is not None else ""
    
    return re.sub(r'\{([^}]+)\}', replace_match, template)


# ============== OOXML Helpers ==============

def set_cell_shading(cell, color: str):
    """Set background color for a cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color.lstrip('#'))
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set cell borders. Pass BorderDef or None."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for name, border in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if border:
            b = OxmlElement(f'w:{name}')
            b.set(qn('w:val'), border.style.value if border.style != BorderStyle.NONE else 'nil')
            b.set(qn('w:sz'), str(int(border.width * 8)))  # eighth-points
            b.set(qn('w:color'), border.color.lstrip('#'))
            tcBorders.append(b)
    
    tcPr.append(tcBorders)


def set_cell_padding(cell, top=0, bottom=0, left=0, right=0):
    """Set cell padding in points."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    
    for side, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin = OxmlElement(f'w:{side}')
        margin.set(qn('w:w'), str(int(value * 20)))  # twips
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    
    tcPr.append(tcMar)


def set_cell_vertical_alignment(cell, alignment: VerticalAlignment):
    """Set vertical alignment for cell content."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), alignment.value)
    tcPr.append(vAlign)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def add_paragraph_border(paragraph, position='bottom', color='#e2e8f0', width=0.5):
    """Add border to paragraph (for dividers)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    border = OxmlElement(f'w:{position}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), str(int(width * 8)))
    border.set(qn('w:color'), color.lstrip('#'))
    pBdr.append(border)
    pPr.append(pBdr)


# ============== Word Document Builder ==============

class WordDocumentBuilder:
    """
    Fluent API for building Word documents with full formatting control.
    
    Example:
        builder = WordDocumentBuilder(theme=DocumentTheme(primary_color="#003366"))
        builder.add_title("Contract Summary")
        builder.add_field_grid([
            {"label": "Number", "value": data["Number"]},
            {"label": "Status", "value": data["Status"]},
        ], columns=2)
        builder.add_section_header("Financial Details")
        builder.add_table(...)
        doc_bytes = builder.save_to_bytes()
    """
    
    def __init__(self, theme: DocumentTheme = None):
        self.theme = theme or DocumentTheme()
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """Apply theme settings to document."""
        section = self.doc.sections[0]
        
        # Orientation
        if self.theme.orientation == Orientation.LANDSCAPE:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        
        # Margins
        section.top_margin = Inches(self.theme.margins.top)
        section.bottom_margin = Inches(self.theme.margins.bottom)
        section.left_margin = Inches(self.theme.margins.left)
        section.right_margin = Inches(self.theme.margins.right)
        section.header_distance = Inches(self.theme.margins.header)
        section.footer_distance = Inches(self.theme.margins.footer)
        
        # Setup styles
        self._setup_styles()
    
    def _setup_styles(self):
        """Configure document styles based on theme."""
        styles = self.doc.styles
        
        # Heading 1
        h1 = styles['Heading 1']
        h1.font.name = self.theme.heading_font
        h1.font.size = Pt(self.theme.heading1_size)
        h1.font.color.rgb = RGBColor(*hex_to_rgb(self.theme.primary_color))
        h1.font.bold = True
        h1.paragraph_format.space_before = Pt(0)
        h1.paragraph_format.space_after = Pt(12)
        
        # Heading 2
        h2 = styles['Heading 2']
        h2.font.name = self.theme.heading_font
        h2.font.size = Pt(self.theme.heading2_size)
        h2.font.color.rgb = RGBColor(*hex_to_rgb(self.theme.secondary_color))
        h2.font.bold = True
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(6)
        
        # Heading 3
        h3 = styles['Heading 3']
        h3.font.name = self.theme.heading_font
        h3.font.size = Pt(self.theme.heading3_size)
        h3.font.color.rgb = RGBColor(*hex_to_rgb(self.theme.text_color))
        h3.font.bold = True
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after = Pt(4)
        
        # Normal
        normal = styles['Normal']
        normal.font.name = self.theme.body_font
        normal.font.size = Pt(self.theme.body_size)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15
    
    def add_title(self, text: str, subtitle: str = None) -> 'WordDocumentBuilder':
        """Add a styled title block with accent bar."""
        # Title with colored banner
        table = self.doc.add_table(rows=1, cols=1)
        table.autofit = False
        table.allow_autofit = False
        
        cell = table.rows[0].cells[0]
        cell.width = Inches(6.5)
        set_cell_shading(cell, self.theme.primary_color)
        set_cell_padding(cell, top=12, bottom=12, left=12, right=12)
        
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(self.theme.title_size)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = self.theme.heading_font
        
        if subtitle:
            para.add_run('\n')
            sub_run = para.add_run(subtitle)
            sub_run.font.size = Pt(self.theme.body_size)
            sub_run.font.color.rgb = RGBColor(230, 230, 230)
            sub_run.font.name = self.theme.body_font
        
        self.doc.add_paragraph()  # Spacing
        return self
    
    def add_section_header(self, text: str, level: int = 2) -> 'WordDocumentBuilder':
        """Add a section heading."""
        self.doc.add_heading(text, level=level)
        return self
    
    def add_paragraph(
        self,
        text: str,
        font: FontStyle = None,
        paragraph_style: ParagraphStyle = None
    ) -> 'WordDocumentBuilder':
        """Add a styled paragraph."""
        para = self.doc.add_paragraph()
        
        if paragraph_style:
            paragraph_style.apply(para)
        
        run = para.add_run(text)
        if font:
            font.apply(run)
        
        return self
    
    def add_text(self, text: str) -> 'WordDocumentBuilder':
        """Add simple text paragraph."""
        self.doc.add_paragraph(text)
        return self
    
    def add_divider(self, color: str = None) -> 'WordDocumentBuilder':
        """Add a horizontal divider line."""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        add_paragraph_border(para, 'bottom', color or self.theme.table_border_color, 0.5)
        return self
    
    def add_spacer(self, points: float = 12) -> 'WordDocumentBuilder':
        """Add vertical spacing."""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(points)
        para.paragraph_format.space_after = Pt(0)
        return self
    
    def add_field_grid(
        self,
        fields: List[Dict[str, Any]],
        columns: int = 2,
        show_labels: bool = True,
        striped: bool = True
    ) -> 'WordDocumentBuilder':
        """
        Add a grid of label-value pairs.
        
        Args:
            fields: List of {"label": str, "value": any, "format": FieldFormat, "format_options": dict}
            columns: Number of field columns
            show_labels: Whether to show labels
            striped: Whether to apply alternating row colors
        """
        if not fields:
            return self
        
        rows_needed = (len(fields) + columns - 1) // columns
        table = self.doc.add_table(rows=rows_needed, cols=columns * 2)
        table.autofit = False
        table.allow_autofit = False
        remove_table_borders(table)
        
        # Calculate widths
        total_width = 6.5
        label_width = total_width / (columns * 3)
        value_width = total_width / (columns * 3) * 2
        
        for row_idx, row in enumerate(table.rows):
            # Apply striping
            if striped and row_idx % 2 == 0:
                for cell in row.cells:
                    set_cell_shading(cell, self.theme.table_stripe_bg)
            
            for col_idx, cell in enumerate(row.cells):
                cell.width = Inches(label_width if col_idx % 2 == 0 else value_width)
        
        for idx, field_def in enumerate(fields):
            row_idx = idx // columns
            col_idx = (idx % columns) * 2
            row = table.rows[row_idx]
            
            # Label
            label_cell = row.cells[col_idx]
            label = field_def.get("label", "")
            if show_labels and label:
                label_para = label_cell.paragraphs[0]
                label_para.paragraph_format.space_before = Pt(4)
                label_para.paragraph_format.space_after = Pt(4)
                label_run = label_para.add_run(label)
                label_run.bold = True
                label_run.font.size = Pt(self.theme.body_size)
                label_run.font.color.rgb = RGBColor(*hex_to_rgb(self.theme.primary_color))
                label_run.font.name = self.theme.body_font
            
            # Value
            value_cell = row.cells[col_idx + 1]
            raw_value = field_def.get("value")
            fmt = field_def.get("format", FieldFormat.TEXT)
            if isinstance(fmt, str):
                fmt = FieldFormat(fmt)
            options = field_def.get("format_options", {})
            formatted_value = format_value(raw_value, fmt, options) if raw_value is not None else "—"
            
            value_para = value_cell.paragraphs[0]
            value_para.paragraph_format.space_before = Pt(4)
            value_para.paragraph_format.space_after = Pt(4)
            value_run = value_para.add_run(formatted_value)
            value_run.font.size = Pt(self.theme.body_size)
            value_run.font.name = self.theme.body_font
            if not raw_value:
                value_run.font.color.rgb = RGBColor(*hex_to_rgb(self.theme.muted_color))
        
        self.doc.add_paragraph()
        return self
    
    def add_data_table(
        self,
        columns: List[Dict[str, Any]],
        rows: List[Dict[str, Any]],
        show_header: bool = True,
        show_row_numbers: bool = False,
        striped: bool = True,
        totals: Dict[str, Any] = None
    ) -> 'WordDocumentBuilder':
        """
        Add a data table.
        
        Args:
            columns: List of {"key": str, "label": str, "format": FieldFormat, "width": float, "align": Alignment}
            rows: List of row data dictionaries
            show_header: Show header row
            show_row_numbers: Show row numbers
            striped: Alternating row colors
            totals: Dict of {column_key: total_value} for totals row
        """
        if not columns or not rows:
            return self
        
        num_cols = len(columns) + (1 if show_row_numbers else 0)
        num_rows = len(rows) + (1 if show_header else 0) + (1 if totals else 0)
        
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        row_offset = 0
        col_offset = 1 if show_row_numbers else 0
        
        # Header row
        if show_header:
            header_row = table.rows[0]
            
            if show_row_numbers:
                cell = header_row.cells[0]
                cell.text = "#"
                self._style_header_cell(cell)
            
            for i, col_def in enumerate(columns):
                cell = header_row.cells[i + col_offset]
                cell.text = col_def.get("label", col_def.get("key", ""))
                self._style_header_cell(cell)
                
                # Column alignment
                align = col_def.get("align", Alignment.LEFT)
                if isinstance(align, str):
                    align = Alignment(align)
                self._set_cell_alignment(cell, align)
            
            row_offset = 1
        
        # Data rows
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + row_offset]
            
            # Striping
            if striped and row_idx % 2 == 1:
                for cell in row.cells:
                    set_cell_shading(cell, self.theme.table_stripe_bg)
            
            if show_row_numbers:
                row.cells[0].text = str(row_idx + 1)
            
            for col_idx, col_def in enumerate(columns):
                cell = row.cells[col_idx + col_offset]
                key = col_def.get("key")
                raw_value = resolve_data_path(row_data, key) if key else None
                
                fmt = col_def.get("format", FieldFormat.TEXT)
                if isinstance(fmt, str):
                    fmt = FieldFormat(fmt)
                options = col_def.get("format_options", {})
                
                cell.text = format_value(raw_value, fmt, options)
                
                align = col_def.get("align", Alignment.LEFT)
                if isinstance(align, str):
                    align = Alignment(align)
                self._set_cell_alignment(cell, align)
        
        # Totals row
        if totals:
            totals_row = table.rows[-1]
            for cell in totals_row.cells:
                set_cell_shading(cell, self.theme.table_header_bg)
            
            first_cell = totals_row.cells[col_offset] if col_offset else totals_row.cells[0]
            self._add_styled_cell_text(first_cell, "Total", bold=True, color=self.theme.table_header_fg)
            
            for col_idx, col_def in enumerate(columns):
                key = col_def.get("key")
                if key and key in totals:
                    cell = totals_row.cells[col_idx + col_offset]
                    fmt = col_def.get("format", FieldFormat.TEXT)
                    if isinstance(fmt, str):
                        fmt = FieldFormat(fmt)
                    options = col_def.get("format_options", {})
                    self._add_styled_cell_text(
                        cell, 
                        format_value(totals[key], fmt, options),
                        bold=True,
                        color=self.theme.table_header_fg
                    )
        
        self.doc.add_paragraph()
        return self
    
    def add_key_value_table(
        self,
        items: List[Tuple[str, Any]],
        value_format: FieldFormat = FieldFormat.TEXT,
        format_options: Dict = None
    ) -> 'WordDocumentBuilder':
        """Add a simple two-column key-value table."""
        table = self.doc.add_table(rows=len(items), cols=2)
        table.style = 'Table Grid'
        
        for idx, (key, value) in enumerate(items):
            row = table.rows[idx]
            
            # Key cell
            key_cell = row.cells[0]
            key_cell.text = key
            para = key_cell.paragraphs[0]
            for run in para.runs:
                run.bold = True
            set_cell_shading(key_cell, self.theme.table_stripe_bg)
            
            # Value cell
            value_cell = row.cells[1]
            value_cell.text = format_value(value, value_format, format_options)
        
        self.doc.add_paragraph()
        return self
    
    def add_image(
        self,
        image_source: Union[str, bytes, io.BytesIO],
        width: float = None,
        height: float = None,
        alignment: Alignment = Alignment.LEFT,
        caption: str = None
    ) -> 'WordDocumentBuilder':
        """Add an image to the document."""
        try:
            if isinstance(image_source, str):
                if image_source.startswith(('http://', 'https://')):
                    import httpx
                    response = httpx.get(image_source)
                    image_source = io.BytesIO(response.content)
                # else treat as file path
            elif isinstance(image_source, bytes):
                image_source = io.BytesIO(image_source)
            
            width_arg = Inches(width) if width else None
            height_arg = Inches(height) if height else None
            
            para = self.doc.add_paragraph()
            alignment_map = {
                Alignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
                Alignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
                Alignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
            }
            para.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
            
            run = para.add_run()
            run.add_picture(image_source, width=width_arg, height=height_arg)
            
            if caption:
                caption_para = self.doc.add_paragraph(caption)
                caption_para.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
                caption_para.runs[0].font.size = Pt(self.theme.small_size)
                caption_para.runs[0].font.color.rgb = RGBColor(*hex_to_rgb(self.theme.muted_color))
        
        except Exception as e:
            log.error(f"Failed to add image: {e}")
        
        return self
    
    def add_page_break(self) -> 'WordDocumentBuilder':
        """Add a page break."""
        self.doc.add_page_break()
        return self
    
    def add_header_footer(
        self,
        header_text: str = None,
        footer_text: str = None,
        show_page_numbers: bool = True
    ) -> 'WordDocumentBuilder':
        """Configure document header and footer."""
        section = self.doc.sections[0]
        
        if header_text:
            header = section.header
            para = header.paragraphs[0]
            para.text = header_text
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if footer_text or show_page_numbers:
            footer = section.footer
            para = footer.paragraphs[0]
            
            if footer_text:
                para.add_run(footer_text)
            
            if show_page_numbers:
                if footer_text:
                    para.add_run("  |  Page ")
                else:
                    para.add_run("Page ")
                
                # Add page number field
                run = para.add_run()
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                run._r.append(fldChar1)
                
                instrText = OxmlElement('w:instrText')
                instrText.text = "PAGE"
                run._r.append(instrText)
                
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')
                run._r.append(fldChar2)
            
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return self
    
    # ============== Helper Methods ==============
    
    def _style_header_cell(self, cell):
        """Apply header cell styling."""
        set_cell_shading(cell, self.theme.table_header_bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(*hex_to_rgb(self.theme.table_header_fg))
    
    def _set_cell_alignment(self, cell, alignment: Alignment):
        """Set cell text alignment."""
        alignment_map = {
            Alignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
            Alignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
            Alignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
        }
        for para in cell.paragraphs:
            para.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
    
    def _add_styled_cell_text(self, cell, text: str, bold: bool = False, color: str = None):
        """Add styled text to a cell, clearing existing content."""
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*hex_to_rgb(color))
    
    # ============== Output Methods ==============
    
    def save(self, path: Union[str, Path]) -> Path:
        """Save document to file."""
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(path)
        return path
    
    def save_to_bytes(self) -> bytes:
        """Get document as bytes."""
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def save_to_stream(self) -> io.BytesIO:
        """Get document as BytesIO stream."""
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer
