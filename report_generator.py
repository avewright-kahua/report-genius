"""
Report Generator Module
Converts LLM-generated markdown + chart specs into professional Word documents.
Supports: charts, images, company logos, tables, and professional styling.
"""

import os
import io
import re
import base64
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple
from dataclasses import dataclass
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Chart libraries
import matplotlib
matplotlib.use('Agg')  # Non-interactive backend
import matplotlib.pyplot as plt

# Default reports directory
DEFAULT_REPORTS_DIR = Path(__file__).parent / "reports"


def hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


@dataclass
class ChartSpec:
    """Specification for generating a chart."""
    chart_type: str  # "bar", "line", "pie", "horizontal_bar", "stacked_bar"
    title: str
    data: Dict[str, Any]  # {"labels": [...], "values": [...]} or {"series": [...]}
    colors: Optional[List[str]] = None
    x_label: Optional[str] = None
    y_label: Optional[str] = None
    width: float = 6
    height: float = 4


@dataclass
class ReportConfig:
    """Configuration for report generation."""
    title: str
    subtitle: Optional[str] = None
    author: Optional[str] = None
    date: Optional[str] = None
    logo_path: Optional[str] = None
    logo_base64: Optional[str] = None
    header_color: str = "#1a365d"
    accent_color: str = "#3182ce"


class ReportGenerator:
    """Generates Word documents from markdown content with chart support."""
    
    def __init__(self, output_dir: Optional[str] = None):
        self.output_dir = Path(output_dir) if output_dir else DEFAULT_REPORTS_DIR
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_chart_image(self, spec: ChartSpec) -> io.BytesIO:
        """Generate a chart and return as BytesIO PNG image."""
        fig, ax = plt.subplots(figsize=(spec.width, spec.height))
        
        default_colors = ['#3182ce', '#63b3ed', '#4fd1c5', '#38b2ac', '#319795', 
                         '#2c7a7b', '#285e61', '#234e52', '#1d4044', '#1a365d']
        colors = spec.colors or default_colors
        
        data = spec.data
        labels = data.get("labels", [])
        values = data.get("values", [])
        series = data.get("series", [])
        
        if spec.chart_type == "bar":
            bars = ax.bar(labels, values, color=colors[:len(labels)], edgecolor='white', linewidth=0.7)
            ax.set_ylabel(spec.y_label or "")
            ax.set_xlabel(spec.x_label or "")
            for bar, val in zip(bars, values):
                height = bar.get_height()
                ax.annotate(f'{val:,.0f}' if isinstance(val, (int, float)) else str(val),
                           xy=(bar.get_x() + bar.get_width() / 2, height),
                           xytext=(0, 3), textcoords="offset points",
                           ha='center', va='bottom', fontsize=9)
        
        elif spec.chart_type == "horizontal_bar":
            bars = ax.barh(labels, values, color=colors[:len(labels)], edgecolor='white', linewidth=0.7)
            ax.set_xlabel(spec.y_label or "")
            ax.set_ylabel(spec.x_label or "")
            for bar, val in zip(bars, values):
                width = bar.get_width()
                ax.annotate(f'{val:,.0f}' if isinstance(val, (int, float)) else str(val),
                           xy=(width, bar.get_y() + bar.get_height() / 2),
                           xytext=(3, 0), textcoords="offset points",
                           ha='left', va='center', fontsize=9)
        
        elif spec.chart_type == "line":
            if series:
                for i, s in enumerate(series):
                    ax.plot(s.get("x", labels), s.get("y", []), 
                           label=s.get("name", f"Series {i+1}"),
                           color=colors[i % len(colors)], linewidth=2, marker='o', markersize=4)
                ax.legend(loc='best', framealpha=0.9)
            else:
                ax.plot(labels, values, color=colors[0], linewidth=2, marker='o', markersize=4)
            ax.set_ylabel(spec.y_label or "")
            ax.set_xlabel(spec.x_label or "")
            ax.grid(True, alpha=0.3)
        
        elif spec.chart_type == "pie":
            wedges, texts, autotexts = ax.pie(
                values, labels=labels, colors=colors[:len(labels)],
                autopct='%1.1f%%', startangle=90, 
                wedgeprops={'edgecolor': 'white', 'linewidth': 1}
            )
            for autotext in autotexts:
                autotext.set_fontsize(9)
                autotext.set_color('white')
                autotext.set_fontweight('bold')
        
        elif spec.chart_type == "stacked_bar":
            if series:
                bottom = [0] * len(labels)
                for i, s in enumerate(series):
                    ax.bar(labels, s.get("values", []), bottom=bottom,
                          label=s.get("name", f"Series {i+1}"),
                          color=colors[i % len(colors)], edgecolor='white', linewidth=0.5)
                    bottom = [b + v for b, v in zip(bottom, s.get("values", []))]
                ax.legend(loc='best', framealpha=0.9)
            ax.set_ylabel(spec.y_label or "")
            ax.set_xlabel(spec.x_label or "")
        
        ax.set_title(spec.title, fontsize=12, fontweight='bold', pad=10)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        plt.tight_layout()
        
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', 
                   facecolor='white', edgecolor='none')
        buf.seek(0)
        plt.close(fig)
        
        return buf
    
    def _set_cell_shading(self, cell, color_hex: str):
        """Set cell background color."""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex.lstrip('#'))
        cell._tc.get_or_add_tcPr().append(shading)
    
    def _parse_markdown_table(self, table_text: str) -> Tuple[List[str], List[List[str]]]:
        """Parse a markdown table into headers and rows."""
        lines = [l.strip() for l in table_text.strip().split('\n') if l.strip()]
        if len(lines) < 2:
            return [], []
        
        # Parse header
        header_line = lines[0]
        headers = [cell.strip() for cell in header_line.strip('|').split('|')]
        
        # Skip separator line (line with dashes)
        rows = []
        for line in lines[2:]:
            if line.startswith('|') or '|' in line:
                cells = [cell.strip() for cell in line.strip('|').split('|')]
                rows.append(cells)
        
        return headers, rows
    
    def _add_formatted_text(self, paragraph, text: str, config: ReportConfig):
        """Add text with inline formatting (bold, italic) to a paragraph."""
        # Pattern to match **bold**, *italic*, and regular text
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'
        parts = re.findall(pattern, text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)
    
    def generate_docx(self,
                      markdown_content: str,
                      config: ReportConfig,
                      charts: Optional[List[ChartSpec]] = None,
                      images: Optional[List[str]] = None,
                      filename: Optional[str] = None) -> Path:
        """Generate a Word document from markdown content."""
        
        doc = Document()
        
        # Set up styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Get colors
        header_rgb = hex_to_rgb(config.header_color)
        accent_rgb = hex_to_rgb(config.accent_color)
        
        # === HEADER SECTION ===
        # Create a table for the header background
        header_table = doc.add_table(rows=1, cols=2 if config.logo_path else 1)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_table.autofit = False
        
        # Set table width to full page
        for cell in header_table.rows[0].cells:
            self._set_cell_shading(cell, config.header_color)
        
        # Title cell
        title_cell = header_table.rows[0].cells[0]
        title_cell.width = Inches(5.5)
        
        # Title
        title_para = title_cell.paragraphs[0]
        title_run = title_para.add_run(config.title)
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Subtitle
        if config.subtitle:
            subtitle_para = title_cell.add_paragraph()
            subtitle_run = subtitle_para.add_run(config.subtitle)
            subtitle_run.font.size = Pt(12)
            subtitle_run.font.color.rgb = RGBColor(200, 213, 224)
        
        # Author and date
        date_str = config.date or datetime.now().strftime("%B %d, %Y")
        meta_para = title_cell.add_paragraph()
        meta_text = f"Prepared by: {config.author}  |  {date_str}" if config.author else date_str
        meta_run = meta_para.add_run(meta_text)
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(226, 232, 240)
        
        # Logo cell (if logo provided)
        if config.logo_path and Path(config.logo_path).exists():
            logo_cell = header_table.rows[0].cells[1]
            logo_cell.width = Inches(1.5)
            logo_para = logo_cell.paragraphs[0]
            logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            logo_para.add_run().add_picture(config.logo_path, width=Inches(1.2))
        
        # Add spacing after header
        doc.add_paragraph()
        
        # === PROCESS MARKDOWN CONTENT ===
        # Generate chart images for placeholders
        chart_images = {}
        if charts:
            for i, chart_spec in enumerate(charts):
                placeholder = f"{{{{CHART_{i}}}}}"
                chart_images[placeholder] = self.generate_chart_image(chart_spec)
        
        # Map image placeholders to file paths
        image_paths = {}
        if images:
            for i, img_path in enumerate(images):
                placeholder = f"{{{{IMAGE_{i}}}}}"
                if Path(img_path).exists():
                    image_paths[placeholder] = img_path
        
        # Split content into blocks
        lines = markdown_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # Check for chart placeholder
            chart_found = False
            for placeholder, chart_buf in chart_images.items():
                if placeholder in stripped:
                    chart_buf.seek(0)
                    doc.add_picture(chart_buf, width=Inches(5.5))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    chart_found = True
                    break
            if chart_found:
                i += 1
                continue
            
            # Check for image placeholder
            image_found = False
            for placeholder, img_path in image_paths.items():
                if placeholder in stripped:
                    # Add the image
                    doc.add_picture(img_path, width=Inches(5.0))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Add caption if there's text after the placeholder
                    caption_text = stripped.replace(placeholder, '').strip()
                    if caption_text:
                        caption = doc.add_paragraph()
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = caption.add_run(caption_text)
                        run.italic = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(100, 100, 100)
                    image_found = True
                    break
            if image_found:
                i += 1
                continue
            
            # Heading 1
            if stripped.startswith('# '):
                para = doc.add_heading(stripped[2:], level=1)
                para.runs[0].font.color.rgb = RGBColor(*header_rgb)
                i += 1
                continue
            
            # Heading 2
            if stripped.startswith('## '):
                para = doc.add_heading(stripped[3:], level=2)
                para.runs[0].font.color.rgb = RGBColor(*header_rgb)
                i += 1
                continue
            
            # Heading 3
            if stripped.startswith('### '):
                para = doc.add_heading(stripped[4:], level=3)
                para.runs[0].font.color.rgb = RGBColor(*header_rgb)
                i += 1
                continue
            
            # Table detection
            if stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
                # Collect all table lines
                table_lines = []
                while i < len(lines) and ('|' in lines[i] or not lines[i].strip()):
                    if lines[i].strip():
                        table_lines.append(lines[i])
                    i += 1
                    if i < len(lines) and not lines[i].strip():
                        break
                
                table_text = '\n'.join(table_lines)
                headers, rows = self._parse_markdown_table(table_text)
                
                if headers and rows:
                    # Create table
                    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Header row
                    header_row = table.rows[0]
                    for j, header in enumerate(headers):
                        cell = header_row.cells[j]
                        cell.text = header
                        self._set_cell_shading(cell, config.header_color)
                        para = cell.paragraphs[0]
                        para.runs[0].font.bold = True
                        para.runs[0].font.color.rgb = RGBColor(255, 255, 255)
                        para.runs[0].font.size = Pt(9)
                    
                    # Data rows
                    for row_idx, row_data in enumerate(rows):
                        row = table.rows[row_idx + 1]
                        for j, cell_text in enumerate(row_data):
                            if j < len(row.cells):
                                cell = row.cells[j]
                                cell.text = cell_text
                                cell.paragraphs[0].runs[0].font.size = Pt(10)
                                # Alternating row colors
                                if row_idx % 2 == 1:
                                    self._set_cell_shading(cell, 'f8fafc')
                    
                    doc.add_paragraph()  # Space after table
                continue
            
            # Bullet list
            if stripped.startswith('- ') or stripped.startswith('* '):
                para = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(para, stripped[2:], config)
                i += 1
                continue
            
            # Numbered list
            if re.match(r'^\d+\.\s', stripped):
                para = doc.add_paragraph(style='List Number')
                text = re.sub(r'^\d+\.\s', '', stripped)
                self._add_formatted_text(para, text, config)
                i += 1
                continue
            
            # Blockquote
            if stripped.startswith('> '):
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.3)
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                run = para.add_run(stripped[2:])
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
                i += 1
                continue
            
            # Regular paragraph
            if stripped:
                para = doc.add_paragraph()
                self._add_formatted_text(para, stripped, config)
            
            i += 1
        
        # === FOOTER ===
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(f"Generated on {date_str}")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(160, 174, 192)
        
        # Generate filename
        if not filename:
            safe_title = "".join(c if c.isalnum() or c in (' ', '-', '_') else '' for c in config.title)
            safe_title = safe_title.replace(' ', '_')[:50]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_title}_{timestamp}"
        
        output_path = self.output_dir / f"{filename}.docx"
        doc.save(str(output_path))
        
        return output_path


def create_report(
    title: str,
    markdown_content: str,
    charts: Optional[List[Dict[str, Any]]] = None,
    images: Optional[List[str]] = None,
    logo_path: Optional[str] = None,
    subtitle: Optional[str] = None,
    author: Optional[str] = None,
    output_dir: Optional[str] = None,
    header_color: str = "#1a365d",
    accent_color: str = "#3182ce"
) -> Dict[str, str]:
    """
    Create a Word document report.
    
    Args:
        title: Report title
        markdown_content: Markdown content with optional placeholders:
            - {{CHART_0}}, {{CHART_1}}, etc. for charts
            - {{IMAGE_0}}, {{IMAGE_1}}, etc. for photos/images
        charts: List of chart specifications
        images: List of image file paths (corresponding to IMAGE_0, IMAGE_1, etc.)
        logo_path: Path to company logo
        subtitle: Report subtitle
        author: Author name
        output_dir: Output directory
        header_color: Header background color (hex)
        accent_color: Accent color (hex)
    
    Returns a dict with:
        - filename: Just the filename (e.g., "Report_20240101_120000.docx")
        - file_path: Full path to the generated document
    """
    generator = ReportGenerator(output_dir)
    
    config = ReportConfig(
        title=title,
        subtitle=subtitle,
        author=author,
        logo_path=logo_path,
        header_color=header_color,
        accent_color=accent_color
    )
    
    # Convert chart dicts to ChartSpec objects
    chart_specs = None
    if charts:
        chart_specs = [ChartSpec(**c) for c in charts]
    
    output_path = generator.generate_docx(markdown_content, config, chart_specs, images)
    return {
        "filename": output_path.name,
        "file_path": str(output_path)
    }
