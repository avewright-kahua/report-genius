"""
SOTA DOCX Renderer - Professional Grade Word Document Generation

This is a state-of-the-art document renderer that produces visually appealing,
professionally styled Word documents with proper typography, tables, and layout.

Key improvements over basic renderer:
- Professional visual hierarchy with proper typography
- Styled tables with alternating rows, proper borders, and spacing
- Flexible header layouts with logo positioning (left/right/center)
- Better field groupings with visual separation
- Proper use of Word styles for consistency
- Modern color schemes and spacing

Token generation uses the kahua_tokens module for standards-compliant Kahua syntax.
"""

from __future__ import annotations

import re
from datetime import date, datetime
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor, Twips, Cm
from docx.table import Table, _Cell

# Import Kahua token builders for standards-compliant token generation
try:
    from .kahua_tokens import (
        build_attribute_token,
        build_boolean_token,
        build_currency_token,
        build_date_token,
        build_number_token,
        build_rich_text_token,
        build_start_table_token,
        build_end_table_token,
        build_if_token,
        build_end_if_token,
        build_company_logo_token,
        build_field_token,
        to_kahua_path,
        TokenSource,
        ConditionOperator,
        DateFormat,
        NumberFormat,
        CurrencyFormat,
        PROJECT_NAME_TOKEN,
        PROJECT_NUMBER_TOKEN,
        REPORT_MODIFIED_TIMESTAMP_TOKEN,
    )
except ImportError:
    from kahua_tokens import (
        build_attribute_token,
        build_boolean_token,
        build_currency_token,
        build_date_token,
        build_number_token,
        build_rich_text_token,
        build_start_table_token,
        build_end_table_token,
        build_if_token,
        build_end_if_token,
        build_company_logo_token,
        build_field_token,
        to_kahua_path,
        TokenSource,
        ConditionOperator,
        DateFormat,
        NumberFormat,
        CurrencyFormat,
        PROJECT_NAME_TOKEN,
        PROJECT_NUMBER_TOKEN,
        REPORT_MODIFIED_TIMESTAMP_TOKEN,
    )

try:
    from .template_schema import (
        Alignment,
        DetailConfig,
        FieldDef,
        FieldFormat,
        HeaderConfig,
        HyperlinkDef,
        LayoutConfig,
        ListConfig,
        PageHeaderFooterConfig,
        PortableViewTemplate,
        Section,
        SectionType,
        StyleConfig,
        TableColumn,
        TableConfig,
        TextConfig,
    )
except ImportError:
    from template_schema import (
        Alignment,
        DetailConfig,
        FieldDef,
        FieldFormat,
        HeaderConfig,
        HyperlinkDef,
        LayoutConfig,
        ListConfig,
        PageHeaderFooterConfig,
        PortableViewTemplate,
        Section,
        SectionType,
        StyleConfig,
        TableColumn,
        TableConfig,
        TextConfig,
    )


# =============================================================================
# Design Tokens - Professional Styling
# =============================================================================

class DesignTokens:
    """Professional design tokens for consistent styling."""
    
    # Typography
    FONT_HEADING = "Calibri Light"
    FONT_BODY = "Calibri"
    FONT_MONO = "Consolas"
    
    SIZE_TITLE = 24
    SIZE_HEADING_1 = 16
    SIZE_HEADING_2 = 13
    SIZE_BODY = 10
    SIZE_SMALL = 9
    SIZE_CAPTION = 8
    
    # Colors - Modern professional palette
    COLOR_PRIMARY = "#1a365d"      # Dark blue
    COLOR_SECONDARY = "#2c5282"    # Medium blue
    COLOR_ACCENT = "#3182ce"       # Bright blue
    COLOR_MUTED = "#718096"        # Gray
    COLOR_LIGHT = "#e2e8f0"        # Light gray
    COLOR_DIVIDER = "#cbd5e0"      # Border gray
    
    # Table colors
    TABLE_HEADER_BG = "#1a365d"
    TABLE_HEADER_FG = "#ffffff"
    TABLE_ROW_ALT = "#f7fafc"
    TABLE_BORDER = "#e2e8f0"
    
    # Spacing (in points)
    SPACE_SECTION = 18
    SPACE_PARAGRAPH = 6
    SPACE_TIGHT = 3
    
    # Table cell padding
    CELL_PADDING = Inches(0.08)


# =============================================================================
# Kahua Placeholder Builders (delegating to kahua_tokens module)
# =============================================================================

def _to_kahua_path(path: str, entity_prefix: Optional[str] = None) -> str:
    """Convert a template path to Kahua attribute path."""
    return to_kahua_path(path, entity_prefix)


def build_attribute(path: str, prefix: Optional[str] = None) -> str:
    """Build [Attribute(Path)] placeholder."""
    return build_attribute_token(path, prefix=prefix)


def build_currency(path: str, prefix: Optional[str] = None) -> str:
    """Build [Currency(...)] placeholder."""
    return build_currency_token(path, prefix=prefix, format=CurrencyFormat.DEFAULT)


def build_number(path: str, fmt: str = "N0", prefix: Optional[str] = None) -> str:
    """Build [Number(...)] placeholder."""
    # Map string format to NumberFormat enum
    format_map = {
        "N0": NumberFormat.INTEGER,
        "N1": NumberFormat.ONE_DECIMAL,
        "N2": NumberFormat.TWO_DECIMALS,
        "F0": NumberFormat.FIXED_0,
        "F1": NumberFormat.FIXED_1,
        "F2": NumberFormat.FIXED_2,
        "P0": NumberFormat.PERCENT_0,
        "P1": NumberFormat.PERCENT_1,
        "P2": NumberFormat.PERCENT_2,
    }
    number_format = format_map.get(fmt, NumberFormat.INTEGER)
    return build_number_token(path, prefix=prefix, format=number_format)


def build_date(path: str, fmt: str = "d", prefix: Optional[str] = None) -> str:
    """Build [Date(...)] placeholder."""
    # Map string format to DateFormat enum
    format_map = {
        "d": DateFormat.SHORT_DATE,
        "D": DateFormat.LONG_DATE,
        "t": DateFormat.SHORT_TIME,
        "T": DateFormat.LONG_TIME,
        "f": DateFormat.FULL_SHORT,
        "F": DateFormat.FULL_LONG,
        "g": DateFormat.GENERAL_SHORT,
        "G": DateFormat.GENERAL_LONG,
    }
    date_format = format_map.get(fmt, DateFormat.SHORT_DATE)
    return build_date_token(path, prefix=prefix, format=date_format)


def build_boolean(path: str, prefix: Optional[str] = None, true_value: str = "Yes", false_value: str = "No") -> str:
    """Build [Boolean(...)] placeholder."""
    return build_boolean_token(path, prefix=prefix, true_value=true_value, false_value=false_value)


def build_start_table(name: str, source: str, prefix: Optional[str] = None) -> str:
    """Build [StartTable(...)] marker."""
    return build_start_table_token(name, source, prefix=prefix)


def build_end_table() -> str:
    """Build [EndTable] closing marker."""
    return build_end_table_token()


def build_if(path: str, op: str = "IsNotEmpty", value: str = None, prefix: str = None) -> str:
    """Build [IF(...)] conditional."""
    # Map string operator to ConditionOperator enum
    op_map = {
        "IsNotEmpty": ConditionOperator.IS_NOT_EMPTY,
        "IsEmpty": ConditionOperator.IS_EMPTY,
        "IsNotNull": ConditionOperator.IS_NOT_NULL,
        "IsNull": ConditionOperator.IS_NULL,
        "Equals": ConditionOperator.EQUALS,
        "DoesNotEqual": ConditionOperator.DOES_NOT_EQUAL,
        "Contains": ConditionOperator.CONTAINS,
        "IsGreaterThan": ConditionOperator.IS_GREATER_THAN,
        "IsLessThan": ConditionOperator.IS_LESS_THAN,
    }
    operator = op_map.get(op, ConditionOperator.IS_NOT_EMPTY)
    return build_if_token(path, operator=operator, value=value, prefix=prefix)


def format_placeholder(field: FieldDef, prefix: str = None, in_table: bool = False) -> str:
    """
    Generate appropriate Kahua placeholder based on field format.
    
    Uses the kahua_tokens module for standards-compliant token generation.
    """
    p = None if in_table else prefix
    
    # Map FieldFormat to field_type string for build_field_token
    format_type_map = {
        FieldFormat.CURRENCY: "currency",
        FieldFormat.DECIMAL: "decimal",
        FieldFormat.NUMBER: "number",
        FieldFormat.PERCENT: "percent",
        FieldFormat.DATE: "date",
        FieldFormat.DATETIME: "datetime",
        FieldFormat.BOOLEAN: "boolean",
        FieldFormat.RICH_TEXT: "rich_text",
        FieldFormat.IMAGE: "image",
        FieldFormat.TEXT: "text",
        FieldFormat.PHONE: "text",
        FieldFormat.EMAIL: "text",
        FieldFormat.URL: "text",
    }
    
    field_type = format_type_map.get(field.format, "text")
    
    # Build format options dict
    format_options = {}
    if field.format_options:
        if field.format_options.decimals is not None:
            format_options["decimals"] = field.format_options.decimals
        if field.format_options.format:
            format_options["format"] = field.format_options.format
    
    return build_field_token(field.path, field_type, prefix=p, format_options=format_options)


# =============================================================================
# XML Helpers for Word Styling
# =============================================================================

def hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def set_cell_shading(cell: _Cell, color_hex: str) -> None:
    """Set cell background color."""
    color = color_hex.lstrip("#")
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(
    cell: _Cell,
    top: str = None,
    bottom: str = None,
    left: str = None,
    right: str = None,
    size: int = 4,
) -> None:
    """Set cell borders with color control."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    def border_xml(side: str, color: str) -> str:
        if color is None:
            return f'<w:{side} w:val="nil"/>'
        return f'<w:{side} w:val="single" w:sz="{size}" w:color="{color.lstrip("#")}"/>'
    
    borders = parse_xml(
        f'''<w:tcBorders {nsdecls("w")}>
            {border_xml("top", top)}
            {border_xml("bottom", bottom)}
            {border_xml("left", left)}
            {border_xml("right", right)}
        </w:tcBorders>'''
    )
    tcPr.append(borders)


def set_cell_margins(cell: _Cell, margin: Inches) -> None:
    """Set cell padding/margins."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margin_twips = int(margin.twips) if hasattr(margin, 'twips') else int(Inches(0.08).twips)
    
    tcMar = parse_xml(
        f'''<w:tcMar {nsdecls("w")}>
            <w:top w:w="{margin_twips}" w:type="dxa"/>
            <w:bottom w:w="{margin_twips}" w:type="dxa"/>
            <w:left w:w="{margin_twips}" w:type="dxa"/>
            <w:right w:w="{margin_twips}" w:type="dxa"/>
        </w:tcMar>'''
    )
    tcPr.append(tcMar)


def set_cell_vertical_alignment(cell: _Cell, align: str = "center") -> None:
    """Set vertical alignment: top, center, bottom."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{align}"/>')
    tcPr.append(vAlign)


def set_row_height(row, height: Pt, rule: str = "exact") -> None:
    """Set row height."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    rule_val = "exact" if rule == "exact" else "atLeast"
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(height.twips)}" w:hRule="{rule_val}"/>')
    trPr.append(trHeight)


def remove_table_borders(table: Table) -> None:
    """Remove all borders from table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    tblBorders = parse_xml(
        f'''<w:tblBorders {nsdecls("w")}>
            <w:top w:val="nil"/>
            <w:left w:val="nil"/>
            <w:bottom w:val="nil"/>
            <w:right w:val="nil"/>
            <w:insideH w:val="nil"/>
            <w:insideV w:val="nil"/>
        </w:tblBorders>'''
    )
    tblPr.append(tblBorders)


def set_table_width(table: Table, width_pct: int = 100) -> None:
    """Set table width as percentage."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{width_pct * 50}" w:type="pct"/>')
    tblPr.append(tblW)


# =============================================================================
# SOTA Document Renderer
# =============================================================================

class SOTADocxRenderer:
    """
    State-of-the-art DOCX renderer producing professional quality documents.
    
    Features:
    - Professional typography with proper hierarchy
    - Styled tables with headers, alternating rows
    - Flexible layout options
    - Proper spacing and visual rhythm
    - Full Kahua placeholder support
    """
    
    def __init__(self, template: PortableViewTemplate):
        self.template = template
        self.doc = Document()
        self.tokens = DesignTokens()
        
        # Extract entity prefix for Kahua paths
        entity_def = template.entity_def or ""
        self.entity_prefix = entity_def.split(".")[-1] if "." in entity_def else entity_def
        
        self._setup_document()
        self._setup_styles()
    
    def _setup_document(self) -> None:
        """Configure page layout."""
        section = self.doc.sections[0]
        
        layout = self.template.layout
        
        # Orientation
        if layout and layout.orientation == "landscape":
            section.page_width, section.page_height = section.page_height, section.page_width
        
        # Margins (professional defaults)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        # Configure page headers and footers if specified
        if layout:
            if layout.page_header:
                self._setup_page_header(section, layout.page_header)
            if layout.page_footer:
                self._setup_page_footer(section, layout.page_footer)
    
    def _setup_page_header(self, section, config: PageHeaderFooterConfig) -> None:
        """Setup page header with left/center/right text and optional page numbers."""
        header = section.header
        header.is_linked_to_previous = False
        
        # Clear default paragraph
        if header.paragraphs:
            header.paragraphs[0].clear()
        
        # Create a table for left/center/right alignment (invisible borders)
        table = header.add_table(rows=1, cols=3, width=Inches(7))
        table.autofit = False
        table.allow_autofit = False
        
        # Set column widths
        for cell in table.rows[0].cells:
            cell.width = Inches(7/3)
        
        # Remove table borders
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tblBorders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '<w:top w:val="nil"/>'
            '<w:left w:val="nil"/>'
            '<w:bottom w:val="nil"/>'
            '<w:right w:val="nil"/>'
            '<w:insideH w:val="nil"/>'
            '<w:insideV w:val="nil"/>'
            '</w:tblBorders>'
        )
        tblPr.append(tblBorders)
        
        cells = table.rows[0].cells
        
        # Left cell
        if config.left_text:
            p = cells[0].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(config.left_text)
            run.font.size = Pt(config.font_size)
        
        # Center cell
        if config.center_text:
            p = cells[1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(config.center_text)
            run.font.size = Pt(config.font_size)
        
        # Right cell - can include page numbers
        if config.right_text or config.include_page_number:
            p = cells[2].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if config.right_text:
                run = p.add_run(config.right_text)
                run.font.size = Pt(config.font_size)
            if config.include_page_number:
                if config.right_text:
                    p.add_run(" | ")
                self._add_page_number_field(p, config)
    
    def _setup_page_footer(self, section, config: PageHeaderFooterConfig) -> None:
        """Setup page footer with left/center/right text and optional page numbers."""
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # Clear default paragraph
        if footer.paragraphs:
            footer.paragraphs[0].clear()
        
        # Create a table for left/center/right alignment (invisible borders)
        table = footer.add_table(rows=1, cols=3, width=Inches(7))
        table.autofit = False
        table.allow_autofit = False
        
        # Set column widths
        for cell in table.rows[0].cells:
            cell.width = Inches(7/3)
        
        # Remove table borders
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tblBorders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '<w:top w:val="nil"/>'
            '<w:left w:val="nil"/>'
            '<w:bottom w:val="nil"/>'
            '<w:right w:val="nil"/>'
            '<w:insideH w:val="nil"/>'
            '<w:insideV w:val="nil"/>'
            '</w:tblBorders>'
        )
        tblPr.append(tblBorders)
        
        cells = table.rows[0].cells
        
        # Left cell
        if config.left_text:
            p = cells[0].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(config.left_text)
            run.font.size = Pt(config.font_size)
        
        # Center cell - typically page numbers go here in footers
        if config.center_text or config.include_page_number:
            p = cells[1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if config.include_page_number:
                self._add_page_number_field(p, config)
            elif config.center_text:
                run = p.add_run(config.center_text)
                run.font.size = Pt(config.font_size)
        
        # Right cell
        if config.right_text:
            p = cells[2].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(config.right_text)
            run.font.size = Pt(config.font_size)
    
    def _add_page_number_field(self, paragraph, config: PageHeaderFooterConfig) -> None:
        """Add PAGE and NUMPAGES fields to paragraph for dynamic page numbering."""
        # Parse the format string and insert fields
        # Format like "Page {page} of {total}" becomes "Page [PAGE] of [NUMPAGES]"
        format_str = config.page_number_format
        
        # Split by placeholders
        parts = re.split(r'(\{page\}|\{total\})', format_str, flags=re.IGNORECASE)
        
        for part in parts:
            if part.lower() == '{page}':
                # Add PAGE field using XML
                run = paragraph.add_run()
                run.font.size = Pt(config.font_size)
                fld_xml = (
                    f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
                )
                run._r.append(parse_xml(fld_xml))
            elif part.lower() == '{total}':
                # Add NUMPAGES field using XML
                run = paragraph.add_run()
                run.font.size = Pt(config.font_size)
                fld_xml = (
                    f'<w:fldSimple {nsdecls("w")} w:instr=" NUMPAGES "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
                )
                run._r.append(parse_xml(fld_xml))
            elif part:
                run = paragraph.add_run(part)
                run.font.size = Pt(config.font_size)
    
    def _setup_styles(self) -> None:
        """Configure document styles for professional output."""
        styles = self.doc.styles
        
        # Normal style - base for all text
        normal = styles["Normal"]
        normal.font.name = self.tokens.FONT_BODY
        normal.font.size = Pt(self.tokens.SIZE_BODY)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.line_spacing = 1.15
        
        # Title style - main document title
        if "PV Title" not in [s.name for s in styles]:
            title_style = styles.add_style("PV Title", WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = self.tokens.FONT_HEADING
            title_style.font.size = Pt(self.tokens.SIZE_TITLE)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(*hex_to_rgb(self.tokens.COLOR_PRIMARY))
            title_style.paragraph_format.space_after = Pt(4)
        
        # Heading 1 - Major section headers
        if "PV Heading 1" not in [s.name for s in styles]:
            h1_style = styles.add_style("PV Heading 1", WD_STYLE_TYPE.PARAGRAPH)
            h1_style.font.name = self.tokens.FONT_HEADING
            h1_style.font.size = Pt(self.tokens.SIZE_HEADING_1)
            h1_style.font.bold = True
            h1_style.font.color.rgb = RGBColor(*hex_to_rgb(self.tokens.COLOR_PRIMARY))
            h1_style.paragraph_format.space_before = Pt(self.tokens.SPACE_SECTION)
            h1_style.paragraph_format.space_after = Pt(self.tokens.SPACE_PARAGRAPH)
        
        # Heading 2 - Sub-section headers
        if "PV Heading 2" not in [s.name for s in styles]:
            h2_style = styles.add_style("PV Heading 2", WD_STYLE_TYPE.PARAGRAPH)
            h2_style.font.name = self.tokens.FONT_HEADING
            h2_style.font.size = Pt(self.tokens.SIZE_HEADING_2)
            h2_style.font.bold = True
            h2_style.font.color.rgb = RGBColor(*hex_to_rgb(self.tokens.COLOR_SECONDARY))
            h2_style.paragraph_format.space_before = Pt(12)
            h2_style.paragraph_format.space_after = Pt(4)
        
        # Subtitle - for header subtitles
        if "PV Subtitle" not in [s.name for s in styles]:
            sub_style = styles.add_style("PV Subtitle", WD_STYLE_TYPE.PARAGRAPH)
            sub_style.font.name = self.tokens.FONT_BODY
            sub_style.font.size = Pt(self.tokens.SIZE_HEADING_2)
            sub_style.font.color.rgb = RGBColor(*hex_to_rgb(self.tokens.COLOR_MUTED))
            sub_style.paragraph_format.space_after = Pt(2)
        
        # Label style - for field labels
        if "PV Label" not in [s.name for s in styles]:
            label_style = styles.add_style("PV Label", WD_STYLE_TYPE.CHARACTER)
            label_style.font.name = self.tokens.FONT_BODY
            label_style.font.size = Pt(self.tokens.SIZE_SMALL)
            label_style.font.bold = True
            label_style.font.color.rgb = RGBColor(*hex_to_rgb(self.tokens.COLOR_MUTED))
        
        # Value style - for field values
        if "PV Value" not in [s.name for s in styles]:
            value_style = styles.add_style("PV Value", WD_STYLE_TYPE.CHARACTER)
            value_style.font.name = self.tokens.FONT_BODY
            value_style.font.size = Pt(self.tokens.SIZE_BODY)
        
        # Caption style - for table captions, notes
        if "PV Caption" not in [s.name for s in styles]:
            caption_style = styles.add_style("PV Caption", WD_STYLE_TYPE.PARAGRAPH)
            caption_style.font.name = self.tokens.FONT_BODY
            caption_style.font.size = Pt(self.tokens.SIZE_CAPTION)
            caption_style.font.italic = True
            caption_style.font.color.rgb = RGBColor(*hex_to_rgb(self.tokens.COLOR_MUTED))
        
        # Backward compatibility - keep PV Heading alias
        if "PV Heading" not in [s.name for s in styles]:
            heading_style = styles.add_style("PV Heading", WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.name = self.tokens.FONT_HEADING
            heading_style.font.size = Pt(self.tokens.SIZE_HEADING_1)
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(*hex_to_rgb(self.tokens.COLOR_PRIMARY))
            heading_style.paragraph_format.space_before = Pt(self.tokens.SPACE_SECTION)
            heading_style.paragraph_format.space_after = Pt(self.tokens.SPACE_PARAGRAPH)
        
        # Hyperlink style - for clickable links
        if "Hyperlink" not in [s.name for s in styles]:
            try:
                link_style = styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
                link_style.font.color.rgb = RGBColor(0, 0, 255)  # Blue
                link_style.font.underline = True
            except:
                pass  # Style may already exist as built-in
    
    def render(self) -> Document:
        """Render the template to a Word document."""
        for section in self.template.get_sections_ordered():
            self._render_section(section)
        return self.doc
    
    def render_to_file(self, path: Union[str, Path]) -> Path:
        """Render and save to file."""
        self.render()
        path = Path(path)
        self.doc.save(str(path))
        return path
    
    def render_to_bytes(self) -> bytes:
        """Render and return as bytes."""
        self.render()
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _render_section(self, section: Section) -> None:
        """Route to appropriate section renderer."""
        renderers = {
            SectionType.HEADER: self._render_header,
            SectionType.DETAIL: self._render_detail,
            SectionType.TEXT: self._render_text,
            SectionType.TABLE: self._render_table,
            SectionType.IMAGE: self._render_image,
            SectionType.DIVIDER: self._render_divider,
            SectionType.SPACER: self._render_spacer,
            SectionType.LIST: self._render_list,
        }
        renderer = renderers.get(section.type)
        if renderer:
            renderer(section)
    
    # =========================================================================
    # Header Section - Professional styled header with logo positioning
    # =========================================================================
    
    def _render_header(self, section: Section) -> None:
        """Render professional header with optional logo positioning."""
        config = section.header_config
        if not config:
            return
        
        # Determine logo position (default: left)
        logo_position = getattr(config, 'logo_position', 'left')
        
        if config.show_logo:
            self._render_header_with_logo(config, logo_position)
        else:
            self._render_header_simple(config)
        
        # Additional header fields
        if config.fields:
            self._render_fields_grid(config.fields, columns=3, compact=True)
    
    def _render_header_with_logo(self, config: HeaderConfig, position: str) -> None:
        """Render header with logo on specified side."""
        # Use a 2-column table for layout
        table = self.doc.add_table(rows=1, cols=2)
        remove_table_borders(table)
        set_table_width(table, 100)
        
        # Determine which cell gets logo vs title
        if position == "right":
            title_cell, logo_cell = table.rows[0].cells[0], table.rows[0].cells[1]
        else:  # left is default
            logo_cell, title_cell = table.rows[0].cells[0], table.rows[0].cells[1]
        
        # Set column widths (logo column narrower)
        from docx.oxml.ns import qn as oxml_qn
        tblGrid = table._tbl.tblGrid
        gridCols = tblGrid.findall(oxml_qn('w:gridCol'))
        if position == "right":
            gridCols[0].set(oxml_qn('w:w'), '7000')  # Title column wider
            gridCols[1].set(oxml_qn('w:w'), '2000')  # Logo column
        else:
            gridCols[0].set(oxml_qn('w:w'), '2000')
            gridCols[1].set(oxml_qn('w:w'), '7000')
        
        # Logo cell - make the placeholder clearly visible
        logo_p = logo_cell.paragraphs[0]
        logo_token = "[CompanyLogo(Height=60,Width=60)]"
        run = logo_p.add_run(logo_token)
        run.font.size = Pt(12)  # Larger font for visibility
        run.font.bold = True
        run.font.color.rgb = RGBColor(*hex_to_rgb(self.tokens.COLOR_MUTED))  # Gray to indicate placeholder
        if position == "right":
            logo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_vertical_alignment(logo_cell, "center")
        
        # Static title first (if provided) - literal text, not a placeholder
        if config.static_title:
            static_p = title_cell.paragraphs[0]
            run = static_p.add_run(config.static_title)
            run.font.name = config.title_font or self.tokens.FONT_HEADING
            run.font.size = Pt(config.title_size or self.tokens.SIZE_TITLE)
            run.font.bold = config.title_bold
            if config.title_color:
                run.font.color.rgb = RGBColor(*hex_to_rgb(config.title_color))
            else:
                run.font.color.rgb = RGBColor(*hex_to_rgb(self.tokens.COLOR_PRIMARY))
            # Alignment
            align_map = {Alignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT, Alignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER, Alignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT}
            static_p.alignment = align_map.get(config.title_alignment, WD_ALIGN_PARAGRAPH.LEFT)
            # Add dynamic title below static title
            title_p = title_cell.add_paragraph()
        else:
            title_p = title_cell.paragraphs[0]
        
        # Dynamic title (placeholder)
        title_text = self._expand_template(config.title_template)
        run = title_p.add_run(title_text)
        run.font.name = config.title_font or self.tokens.FONT_HEADING
        run.font.size = Pt(config.title_size or self.tokens.SIZE_TITLE)
        run.font.bold = config.title_bold
        if config.title_color:
            run.font.color.rgb = RGBColor(*hex_to_rgb(config.title_color))
        else:
            run.font.color.rgb = RGBColor(*hex_to_rgb(self.tokens.COLOR_PRIMARY))
        if position == "right":
            title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_vertical_alignment(title_cell, "center")
        
        # Subtitle if present
        if config.subtitle_template:
            subtitle_text = self._expand_template(config.subtitle_template)
            sub_p = title_cell.add_paragraph()
            run = sub_p.add_run(subtitle_text)
            run.font.size = Pt(self.tokens.SIZE_HEADING_2)
            run.font.color.rgb = RGBColor(*hex_to_rgb(self.tokens.COLOR_MUTED))
            if position == "right":
                sub_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                sub_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add spacing after header
        self.doc.add_paragraph()
    
    def _render_header_simple(self, config: HeaderConfig) -> None:
        """Render simple header without logo."""
        # Static title first (if provided) - literal text, not a placeholder
        if config.static_title:
            p = self.doc.add_paragraph()
            run = p.add_run(config.static_title)
            run.font.name = config.title_font or self.tokens.FONT_HEADING
            run.font.size = Pt(config.title_size or self.tokens.SIZE_TITLE)
            run.font.bold = config.title_bold
            if config.title_color:
                run.font.color.rgb = RGBColor(*hex_to_rgb(config.title_color))
            else:
                run.font.color.rgb = RGBColor(*hex_to_rgb(self.tokens.COLOR_PRIMARY))
            # Alignment
            align_map = {Alignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT, Alignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER, Alignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT}
            p.alignment = align_map.get(config.title_alignment, WD_ALIGN_PARAGRAPH.LEFT)
        
        # Dynamic title (placeholder-based)
        title_text = self._expand_template(config.title_template)
        p = self.doc.add_paragraph()
        run = p.add_run(title_text)
        run.font.name = config.title_font or self.tokens.FONT_HEADING
        run.font.size = Pt(config.title_size or self.tokens.SIZE_TITLE)
        run.font.bold = config.title_bold
        if config.title_color:
            run.font.color.rgb = RGBColor(*hex_to_rgb(config.title_color))
        else:
            run.font.color.rgb = RGBColor(*hex_to_rgb(self.tokens.COLOR_PRIMARY))
        
        if config.subtitle_template:
            subtitle_text = self._expand_template(config.subtitle_template)
            p = self.doc.add_paragraph()
            run = p.add_run(subtitle_text)
            run.font.size = Pt(self.tokens.SIZE_HEADING_2)
            run.font.color.rgb = RGBColor(*hex_to_rgb(self.tokens.COLOR_MUTED))
    
    # =========================================================================
    # Detail Section - Professional field grid
    # =========================================================================
    
    def _render_detail(self, section: Section) -> None:
        """Render detail section as professional field grid."""
        config = section.detail_config
        if not config or not config.fields:
            return
        
        if section.title:
            self._add_section_heading(section.title)
        
        self._render_fields_grid(config.fields, columns=config.columns)
    
    def _render_fields_grid(
        self,
        fields: List[FieldDef],
        columns: int = 2,
        compact: bool = False,
    ) -> None:
        """Render fields as a professional label-value grid."""
        rows_needed = (len(fields) + columns - 1) // columns
        
        # Create table with label+value pairs per column
        table = self.doc.add_table(rows=rows_needed, cols=columns * 2)
        remove_table_borders(table)
        set_table_width(table, 100)
        
        for i, field in enumerate(fields):
            row_idx = i // columns
            col_idx = (i % columns) * 2
            
            # Label cell
            label_cell = table.rows[row_idx].cells[col_idx]
            label_p = label_cell.paragraphs[0]
            label_text = field.label or self._path_to_label(field.path)
            run = label_p.add_run(label_text)
            run.font.bold = True
            run.font.size = Pt(self.tokens.SIZE_SMALL)
            run.font.color.rgb = RGBColor(*hex_to_rgb(self.tokens.COLOR_MUTED))
            label_p.paragraph_format.space_before = Pt(4 if compact else 8)
            
            # Value cell
            value_cell = table.rows[row_idx].cells[col_idx + 1]
            value_p = value_cell.paragraphs[0]
            placeholder = format_placeholder(field, self.entity_prefix)
            run = value_p.add_run(placeholder)
            run.font.size = Pt(self.tokens.SIZE_BODY)
            
            # Apply emphasis if specified
            if getattr(field, 'emphasis', None) == 'bold':
                run.font.bold = True
            
            value_p.paragraph_format.space_before = Pt(4 if compact else 8)
            
            # Add subtle bottom border for visual rhythm
            if not compact:
                set_cell_borders(label_cell, bottom=self.tokens.COLOR_LIGHT)
                set_cell_borders(value_cell, bottom=self.tokens.COLOR_LIGHT)
    
    # =========================================================================
    # Table Section - Professional styled tables
    # =========================================================================
    
    def _render_table(self, section: Section) -> None:
        """Render professional data table with styling."""
        config = section.table_config
        if not config or not config.columns:
            return
        
        # Conditional wrapper
        if section.condition:
            self._add_condition_start(section.condition)
        
        if section.title:
            self._add_section_heading(section.title)
        
        # Table name for Kahua
        source_parts = config.source.split(".")
        table_name = source_parts[-1]
        
        # Start table marker
        start_marker = build_start_table(table_name, config.source, self.entity_prefix)
        p = self.doc.add_paragraph()
        run = p.add_run(start_marker)
        run.font.size = Pt(1)
        run.font.color.rgb = RGBColor(255, 255, 255)  # Hide marker
        
        # Determine rows needed (header + data + optional totals)
        num_rows = 2  # header + template row
        num_cols = len(config.columns)
        
        # Create table
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        set_table_width(table, 100)
        
        # Set column widths if specified
        self._set_column_widths(table, config.columns)
        
        # Style header row
        header_row = table.rows[0]
        for i, col in enumerate(config.columns):
            cell = header_row.cells[i]
            
            # Header background
            set_cell_shading(cell, self.tokens.TABLE_HEADER_BG)
            set_cell_margins(cell, Inches(0.08))
            set_cell_vertical_alignment(cell, "center")
            
            # Header text
            p = cell.paragraphs[0]
            p.clear()
            label = col.field.label or self._path_to_label(col.field.path)
            run = p.add_run(label)
            run.font.bold = True
            run.font.size = Pt(self.tokens.SIZE_SMALL)
            run.font.color.rgb = RGBColor(*hex_to_rgb(self.tokens.TABLE_HEADER_FG))
            
            # Alignment
            if col.alignment == Alignment.CENTER:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif col.alignment == Alignment.RIGHT:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Data row template
        data_row = table.rows[1]
        for i, col in enumerate(config.columns):
            cell = data_row.cells[i]
            set_cell_margins(cell, Inches(0.08))
            
            # Bottom border for rows
            set_cell_borders(cell, bottom=self.tokens.TABLE_BORDER)
            
            # Placeholder
            p = cell.paragraphs[0]
            p.clear()
            placeholder = format_placeholder(col.field, table_name, in_table=True)
            run = p.add_run(placeholder)
            run.font.size = Pt(self.tokens.SIZE_BODY)
            
            # Alignment
            if col.alignment == Alignment.CENTER:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif col.alignment == Alignment.RIGHT:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # End table marker
        p = self.doc.add_paragraph()
        run = p.add_run(build_end_table())
        run.font.size = Pt(1)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Subtotals section (outside table, after EndTable)
        if config.show_subtotals and config.subtotal_fields:
            self._render_table_subtotals(config, table_name)
        
        # Close conditional
        if section.condition:
            self._add_condition_end()
    
    def _set_column_widths(self, table: Table, columns: List[TableColumn]) -> None:
        """Set column widths based on specifications."""
        total_specified = sum(c.width or 0 for c in columns)
        unspecified_count = sum(1 for c in columns if not c.width)
        
        # Calculate remaining width for unspecified columns
        remaining = 100 - total_specified if total_specified < 100 else 0
        default_width = remaining / unspecified_count if unspecified_count else 0
        
        # Apply widths via XML
        tbl = table._tbl
        tblGrid = tbl.get_or_add_tblGrid()
        
        # Clear existing gridCol elements
        for gridCol in tblGrid.findall(qn('w:gridCol')):
            tblGrid.remove(gridCol)
        
        # Add new gridCol elements with widths
        for col in columns:
            pct = col.width if col.width else default_width
            # Convert percentage to twips (5000 twips per percentage point for pct type)
            width_val = int(pct * 50)  # Simplification for Word's pct calculation
            gridCol = parse_xml(f'<w:gridCol {nsdecls("w")} w:w="{width_val}"/>')
            tblGrid.append(gridCol)
    
    def _render_table_subtotals(self, config: TableConfig, table_name: str) -> None:
        """Render subtotals section after table."""
        # Create a summary table for totals
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        
        # Build subtotals as a right-aligned block
        totals_table = self.doc.add_table(rows=len(config.subtotal_fields), cols=2)
        remove_table_borders(totals_table)
        
        # Position on right side
        tbl = totals_table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            tbl.insert(0, tblPr)
        jc = parse_xml(f'<w:jc {nsdecls("w")} w:val="right"/>')
        tblPr.append(jc)
        
        # Set width to ~40%
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="2000" w:type="pct"/>')
        tblPr.append(tblW)
        
        for i, field_path in enumerate(config.subtotal_fields):
            row = totals_table.rows[i]
            
            # Label cell
            label_cell = row.cells[0]
            label_text = self._path_to_label(field_path) + ":"
            label_p = label_cell.paragraphs[0]
            run = label_p.add_run(label_text)
            run.font.bold = True
            run.font.size = Pt(self.tokens.SIZE_BODY)
            label_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Value cell - use aggregate placeholder
            value_cell = row.cells[1]
            # For subtotals, we use Sum or other aggregate functions
            kahua_path = _to_kahua_path(f"{config.source}.{field_path}", self.entity_prefix)
            placeholder = f'[Sum(Source=Attribute,Path={kahua_path},Format="C2")]'
            
            value_p = value_cell.paragraphs[0]
            run = value_p.add_run(placeholder)
            run.font.bold = True
            run.font.size = Pt(self.tokens.SIZE_BODY)
            value_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Add bottom border to last row
            if i == len(config.subtotal_fields) - 1:
                set_cell_borders(label_cell, top=self.tokens.COLOR_PRIMARY)
                set_cell_borders(value_cell, top=self.tokens.COLOR_PRIMARY)
    
    # =========================================================================
    # Text Section
    # =========================================================================
    
    def _render_text(self, section: Section) -> None:
        """Render text block section."""
        config = section.text_config
        if not config:
            return
        
        if section.condition:
            self._add_condition_start(section.condition)
        
        if section.title:
            self._add_section_heading(section.title)
        
        content = self._expand_template(config.content)
        p = self.doc.add_paragraph()
        run = p.add_run(content)
        run.font.size = Pt(self.tokens.SIZE_BODY)
        
        # Add hyperlinks if defined
        if config.hyperlinks:
            for link in config.hyperlinks:
                self._add_hyperlink_paragraph(link)
        
        if section.condition:
            self._add_condition_end()
    
    def _add_hyperlink_paragraph(self, hyperlink: HyperlinkDef) -> None:
        """Add a paragraph containing a hyperlink."""
        p = self.doc.add_paragraph()
        self._add_hyperlink(p, hyperlink.url, hyperlink.text, hyperlink.tooltip)
    
    def _add_hyperlink(self, paragraph, url: str, text: str, tooltip: str = None) -> None:
        """Add a hyperlink to a paragraph using XML manipulation.
        
        Args:
            paragraph: The paragraph to add the hyperlink to
            url: The URL (can contain Kahua placeholders like {WebUrl})
            text: The display text
            tooltip: Optional hover tooltip
        """
        # Expand any placeholders in the URL
        expanded_url = self._expand_template(url)
        
        # Get the paragraph's element
        part = paragraph.part
        r_id = part.relate_to(
            expanded_url,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            is_external=True
        )
        
        # Build the hyperlink XML
        hyperlink = parse_xml(
            f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" {nsdecls("r")}>'
            f'<w:r>'
            f'<w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr>'
            f'<w:t>{text}</w:t>'
            f'</w:r>'
            f'</w:hyperlink>'
        )
        
        # Add tooltip if specified
        if tooltip:
            hyperlink.set(qn('w:tooltip'), tooltip)
        
        paragraph._p.append(hyperlink)
    
    # =========================================================================
    # Divider & Spacer
    # =========================================================================
    
    def _render_divider(self, section: Section) -> None:
        """Render horizontal divider."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        
        # Add bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="6" w:color="{self.tokens.COLOR_DIVIDER.lstrip("#")}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    
    def _render_spacer(self, section: Section) -> None:
        """Render vertical spacing."""
        config = section.spacer_config
        height = config.height if config else 0.25
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Inches(height)
    
    # =========================================================================
    # List Section
    # =========================================================================
    
    def _render_list(self, section: Section) -> None:
        """Render bulleted or numbered list section."""
        config = section.list_config
        if not config:
            return
        
        if section.condition:
            self._add_condition_start(section.condition)
        
        if section.title:
            self._add_section_heading(section.title)
        
        # Determine list style
        list_style = 'List Bullet' if config.list_type == 'bullet' else 'List Number'
        
        # Render static items
        for item in config.items:
            p = self.doc.add_paragraph(style=list_style)
            # Check if item contains placeholders
            if '{' in item and '}' in item:
                # Item is a template with field references
                p.add_run(item)
            else:
                p.add_run(item)
            
            # Handle indent level
            if config.indent_level > 0:
                p.paragraph_format.left_indent = Inches(0.5 * config.indent_level)
        
        # Handle dynamic list from collection source
        if config.source and config.item_field:
            # Add Kahua loop syntax for dynamic lists
            start_p = self.doc.add_paragraph()
            start_p.add_run(f'[StartRow({config.source})]')
            
            p = self.doc.add_paragraph(style=list_style)
            p.add_run(f'[Attribute({config.item_field})]')
            
            end_p = self.doc.add_paragraph()
            end_p.add_run('[EndRow()]')
        
        if section.condition:
            self._add_condition_end()
    
    # =========================================================================
    # Image Section
    # =========================================================================
    
    def _render_image(self, section: Section) -> None:
        """Render image section with optional Kahua image placeholder."""
        config = section.image_config
        if not config:
            return
        
        if section.condition:
            self._add_condition_start(section.condition)
        
        if section.title:
            self._add_section_heading(section.title)
        
        p = self.doc.add_paragraph()
        
        # Handle different image source types
        source = getattr(config, 'source', None)
        image_type = getattr(config, 'type', 'static')  # static, company_logo, project_logo, attribute
        
        # Image dimensions
        width = getattr(config, 'width', 200)
        height = getattr(config, 'height', 150)
        
        if image_type == 'company_logo':
            # Kahua company logo placeholder
            run = p.add_run(f"[CompanyLogo(Height={height},Width={width})]")
        elif image_type == 'project_logo':
            # Kahua project logo placeholder  
            run = p.add_run(f"[ProjectLogo(Height={height},Width={width})]")
        elif image_type == 'attribute' and source:
            # Dynamic image from attribute
            kahua_path = _to_kahua_path(source, self.entity_prefix)
            run = p.add_run(f"[Image(Source=Attribute,Path={kahua_path},Height={height},Width={width})]")
        else:
            # Static image placeholder
            run = p.add_run(f"[Image(Height={height},Width={width})]")
        
        run.font.size = Pt(self.tokens.SIZE_BODY)
        
        # Alignment
        alignment = getattr(config, 'alignment', 'left')
        if alignment == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if section.condition:
            self._add_condition_end()
    
    # =========================================================================
    # Helpers
    # =========================================================================
    
    def _add_section_heading(self, title: str) -> None:
        """Add a section heading."""
        p = self.doc.add_paragraph()
        run = p.add_run(title)
        run.font.name = self.tokens.FONT_HEADING
        run.font.size = Pt(self.tokens.SIZE_HEADING_1)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*hex_to_rgb(self.tokens.COLOR_PRIMARY))
        p.paragraph_format.space_before = Pt(self.tokens.SPACE_SECTION)
        p.paragraph_format.space_after = Pt(self.tokens.SPACE_PARAGRAPH)
    
    def _expand_template(self, template_str: str) -> str:
        """Convert {field} to [Attribute()] placeholders."""
        def replace(match: re.Match) -> str:
            path = match.group(1)
            return build_attribute(path, self.entity_prefix)
        return re.sub(r"\{([^}]+)\}", replace, template_str)
    
    def _path_to_label(self, path: str) -> str:
        """Convert field path to readable label."""
        name = path.split(".")[-1]
        name = re.sub(r"([a-z])([A-Z])", r"\1 \2", name)
        name = name.replace("_", " ")
        return name.title()
    
    def _add_condition_start(self, condition) -> None:
        """Add IF condition markers."""
        op_map = {
            "not_empty": "IsNotEmpty",
            "empty": "IsEmpty",
            "equals": "Equals",
            "not_equals": "DoesNotEqual",
        }
        kahua_op = op_map.get(condition.op.value, "IsNotEmpty")
        
        p = self.doc.add_paragraph()
        run = p.add_run(build_if(condition.field, kahua_op, condition.value, self.entity_prefix))
        run.font.size = Pt(1)
        
        p = self.doc.add_paragraph()
        run = p.add_run("[THEN]")
        run.font.size = Pt(1)
    
    def _add_condition_end(self) -> None:
        """Add ENDIF markers."""
        p = self.doc.add_paragraph()
        run = p.add_run("[ENDTHEN]")
        run.font.size = Pt(1)
        
        p = self.doc.add_paragraph()
        run = p.add_run("[ENDIF]")
        run.font.size = Pt(1)


# =============================================================================
# Factory Functions
# =============================================================================

def render_sota_template(
    template: PortableViewTemplate,
    output_path: Optional[Union[str, Path]] = None,
) -> Union[Document, Path]:
    """Render template using SOTA renderer.
    
    Args:
        template: The template to render
        output_path: Optional file path
        
    Returns:
        Path if output_path provided, Document otherwise
    """
    renderer = SOTADocxRenderer(template)
    
    if output_path:
        return renderer.render_to_file(output_path)
    
    return renderer.render()


def render_sota_from_json(
    json_path: Union[str, Path],
    output_path: Optional[Union[str, Path]] = None,
) -> Union[Document, Path]:
    """Load template from JSON and render."""
    import json
    
    with open(json_path) as f:
        data = json.load(f)
    
    template = PortableViewTemplate(**data)
    return render_sota_template(template, output_path)
