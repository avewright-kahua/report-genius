"""
DOCX Renderer - Generates Kahua-compatible Word templates from PortableViewTemplate.

This module converts template definitions into Word documents with Kahua's
placeholder syntax: [Attribute()], [StartTable/EndTable], [StartList/EndList], etc.

The generated .docx files can be uploaded to Kahua and used as Portable View templates.
"""

from __future__ import annotations

import re
from datetime import date, datetime
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.table import Table, _Cell

from template_schema import (
    Alignment,
    DetailConfig,
    FieldDef,
    FieldFormat,
    HeaderConfig,
    LayoutConfig,
    PortableViewTemplate,
    Section,
    SectionType,
    StyleConfig,
    TableColumn,
    TableConfig,
    TextConfig,
)


# =============================================================================
# Kahua Placeholder Syntax Builders
# =============================================================================

def _to_kahua_path(path: str, entity_prefix: Optional[str] = None) -> str:
    """Convert a template path to Kahua attribute path.
    
    Example: "author.short_label" -> "RFI.Author.ShortLabel"
    """
    # Already PascalCase paths pass through
    if "." not in path and path[0].isupper():
        if entity_prefix:
            return f"{entity_prefix}.{path}"
        return path
    
    # Convert each segment to PascalCase
    parts = path.split(".")
    pascal_parts = []
    for part in parts:
        if "_" in part:
            # snake_case to PascalCase
            segments = part.split("_")
            pascal = "".join(seg.capitalize() for seg in segments if seg)
            pascal_parts.append(pascal)
        else:
            # Capitalize first letter
            pascal_parts.append(part[0].upper() + part[1:] if part else part)
    
    result = ".".join(pascal_parts)
    if entity_prefix and not result.startswith(entity_prefix):
        return f"{entity_prefix}.{result}"
    return result


def build_attribute_placeholder(path: str, entity_prefix: Optional[str] = None) -> str:
    """Build [Attribute(Path)] placeholder."""
    kahua_path = _to_kahua_path(path, entity_prefix)
    return f"[Attribute({kahua_path})]"


def build_number_placeholder(
    path: str,
    format_str: str = "F2",
    entity_prefix: Optional[str] = None,
) -> str:
    """Build [Number(Source=Attribute,Path=...,Format="...")] placeholder."""
    kahua_path = _to_kahua_path(path, entity_prefix)
    return f'[Number(Source=Attribute,Path={kahua_path},Format="{format_str}")]'


def build_currency_placeholder(
    path: str,
    entity_prefix: Optional[str] = None,
) -> str:
    """Build [Currency(...)] placeholder."""
    kahua_path = _to_kahua_path(path, entity_prefix)
    return f'[Currency(Source=Attribute,Path={kahua_path},Format="C2")]'


def build_date_placeholder(
    path: str,
    format_str: str = "d",
    entity_prefix: Optional[str] = None,
) -> str:
    """Build [Date(...)] placeholder for dates."""
    kahua_path = _to_kahua_path(path, entity_prefix)
    return f'[Date(Source=Attribute,Path={kahua_path},Format="{format_str}")]'


def build_start_table(
    name: str,
    source_path: str,
    rows_in_header: int = 1,
    entity_prefix: Optional[str] = None,
) -> str:
    """Build [StartTable(...)] marker."""
    kahua_path = _to_kahua_path(source_path, entity_prefix)
    return f"[StartTable(Name={name},Source=Attribute,Path={kahua_path},RowsInHeader={rows_in_header})]"


def build_end_table() -> str:
    """Build [EndTable] marker."""
    return "[EndTable]"


def build_start_list(
    name: str,
    source_path: str,
    entity_prefix: Optional[str] = None,
) -> str:
    """Build [StartList(...)] marker."""
    kahua_path = _to_kahua_path(source_path, entity_prefix)
    return f"[StartList(Name={name},Source=Attribute,Path={kahua_path})][StartContentTemplate]"


def build_end_list() -> str:
    """Build [EndContentTemplate][EndList] markers."""
    return "[EndContentTemplate][EndList]"


def build_if_condition(
    path: str,
    operator: str = "IsNotEmpty",
    value: Optional[str] = None,
    entity_prefix: Optional[str] = None,
) -> str:
    """Build [IF(...)] conditional start.
    
    Args:
        path: Attribute path to check
        operator: Kahua operator (IsNotEmpty, IsEmpty, Equals, etc.)
        value: Value for comparison operators
        entity_prefix: Entity name prefix
    """
    kahua_path = _to_kahua_path(path, entity_prefix)
    if value is not None:
        return f'[IF(Source=Attribute,Path={kahua_path},Operator={operator},Value="{value}")]'
    return f"[IF(Source=Attribute,Path={kahua_path},Operator={operator})]"


def build_then() -> str:
    """Build [THEN] marker."""
    return "[THEN]"


def build_end_then() -> str:
    """Build [ENDTHEN] marker."""
    return "[ENDTHEN]"


def build_else() -> str:
    """Build [ELSE] marker."""
    return "[ELSE]"


def build_end_else() -> str:
    """Build [ENDELSE] marker."""
    return "[ENDELSE]"


def build_endif() -> str:
    """Build [ENDIF] marker."""
    return "[ENDIF]"


def format_field_placeholder(
    field: FieldDef,
    entity_prefix: Optional[str] = None,
    in_table: bool = False,
) -> str:
    """Generate appropriate placeholder based on field format type."""
    # When inside a table, don't include entity prefix (path is relative to row)
    prefix = None if in_table else entity_prefix
    
    if field.format == FieldFormat.CURRENCY:
        return build_currency_placeholder(field.path, prefix)
    
    if field.format == FieldFormat.DECIMAL:
        decimals = 2
        if field.format_options and field.format_options.decimals:
            decimals = field.format_options.decimals
        return build_number_placeholder(field.path, f"F{decimals}", prefix)
    
    if field.format == FieldFormat.NUMBER:
        return build_number_placeholder(field.path, "N0", prefix)
    
    if field.format == FieldFormat.PERCENT:
        return build_number_placeholder(field.path, "P1", prefix)
    
    if field.format == FieldFormat.DATE:
        fmt = "d"  # Short date
        if field.format_options and field.format_options.format:
            fmt = field.format_options.format
        return build_date_placeholder(field.path, fmt, prefix)
    
    if field.format == FieldFormat.DATETIME:
        return build_date_placeholder(field.path, "g", prefix)
    
    # Default to simple attribute
    return build_attribute_placeholder(field.path, prefix)


# =============================================================================
# Color Utilities
# =============================================================================

def hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def set_cell_shading(cell: _Cell, color_hex: str) -> None:
    """Set cell background shading."""
    color = color_hex.lstrip("#")
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(
    cell: _Cell,
    top: bool = True,
    bottom: bool = True,
    left: bool = True,
    right: bool = True,
    color: str = "auto",
    size: int = 4,
) -> None:
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'''<w:tcBorders {nsdecls("w")}>
            <w:top w:val="{'single' if top else 'nil'}" w:sz="{size}" w:color="{color}"/>
            <w:bottom w:val="{'single' if bottom else 'nil'}" w:sz="{size}" w:color="{color}"/>
            <w:left w:val="{'single' if left else 'nil'}" w:sz="{size}" w:color="{color}"/>
            <w:right w:val="{'single' if right else 'nil'}" w:sz="{size}" w:color="{color}"/>
        </w:tcBorders>'''
    )
    tcPr.append(tcBorders)


# =============================================================================
# DOCX Renderer
# =============================================================================

class DocxRenderer:
    """Renders PortableViewTemplate to Kahua-compatible Word documents."""
    
    def __init__(self, template: PortableViewTemplate):
        self.template = template
        self.doc = Document()
        self.style = template.style
        self.layout = template.layout
        self.entity_prefix = self._extract_entity_prefix()
        
        self._setup_document()
    
    def _extract_entity_prefix(self) -> str:
        """Extract entity name for path prefixing."""
        # entity_def might be "RFI", "ExpenseContract", etc.
        entity_def = self.template.entity_def
        # Take last part if it's a path
        if "." in entity_def:
            return entity_def.split(".")[-1]
        return entity_def
    
    def _setup_document(self) -> None:
        """Configure document layout and styles."""
        # Page setup
        section = self.doc.sections[0]
        
        if self.layout.orientation == "landscape":
            section.page_width, section.page_height = section.page_height, section.page_width
        
        # Margins
        section.top_margin = Inches(self.layout.margin_top)
        section.bottom_margin = Inches(self.layout.margin_bottom)
        section.left_margin = Inches(self.layout.margin_left)
        section.right_margin = Inches(self.layout.margin_right)
        
        # Default font
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = self.style.font_family
        font.size = Pt(self.style.body_size)
        
        # Reduce paragraph spacing
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
    
    def render(self) -> Document:
        """Render the template to a Word document."""
        sections = self.template.get_sections_ordered()
        
        for section in sections:
            self._render_section(section)
        
        return self.doc
    
    def render_to_file(self, path: Union[str, Path]) -> None:
        """Render and save to file."""
        self.render()
        self.doc.save(str(path))
    
    def render_to_bytes(self) -> bytes:
        """Render and return as bytes."""
        self.render()
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _render_section(self, section: Section) -> None:
        """Route section to appropriate renderer."""
        renderers = {
            SectionType.HEADER: self._render_header,
            SectionType.DETAIL: self._render_detail,
            SectionType.TEXT: self._render_text,
            SectionType.TABLE: self._render_table,
            SectionType.DIVIDER: self._render_divider,
            SectionType.SPACER: self._render_spacer,
        }
        
        renderer = renderers.get(section.type)
        if renderer:
            renderer(section)
    
    def _add_section_title(self, title: str) -> None:
        """Add a section title paragraph."""
        p = self.doc.add_paragraph()
        run = p.add_run(title)
        run.font.name = self.style.heading_font
        run.font.size = Pt(self.style.heading_size)
        run.font.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    
    def _render_header(self, section: Section) -> None:
        """Render header section with title/subtitle and key fields."""
        config = section.header_config
        if not config:
            return
        
        # Logo if requested
        if config.show_logo:
            logo_token = "[CompanyLogo(Height=60,Width=60)]" if config.logo_type == "company" else "[ProjectLogo(Height=60,Width=60)]"
            p = self.doc.add_paragraph()
            p.add_run(logo_token)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Title from template
        title_text = self._convert_template_to_placeholders(config.title_template)
        p = self.doc.add_paragraph()
        run = p.add_run(title_text)
        run.font.name = self.style.heading_font
        run.font.size = Pt(self.style.title_size)
        run.font.bold = True
        
        # Subtitle if present
        if config.subtitle_template:
            subtitle_text = self._convert_template_to_placeholders(config.subtitle_template)
            p = self.doc.add_paragraph()
            run = p.add_run(subtitle_text)
            run.font.size = Pt(self.style.heading_size)
            run.font.color.rgb = RGBColor(*hex_to_rgb(self.style.muted_color))
        
        # Additional header fields in a grid
        if config.fields:
            self._render_fields_as_grid(config.fields, columns=3)
    
    def _render_detail(self, section: Section) -> None:
        """Render detail section as label-value grid."""
        config = section.detail_config
        if not config or not config.fields:
            return
        
        if section.title:
            self._add_section_title(section.title)
        
        self._render_fields_as_grid(config.fields, columns=config.columns)
    
    def _render_fields_as_grid(
        self,
        fields: List[FieldDef],
        columns: int = 2,
    ) -> None:
        """Render fields as a label-value table grid."""
        # Calculate rows needed
        rows_needed = (len(fields) + columns - 1) // columns
        
        # Create table
        table = self.doc.add_table(rows=rows_needed, cols=columns * 2)
        table.autofit = True
        
        # Calculate available width
        section = self.doc.sections[0]
        available_width = section.page_width - section.left_margin - section.right_margin
        col_width = available_width // (columns * 2)
        
        for i, field in enumerate(fields):
            row_idx = i // columns
            col_idx = (i % columns) * 2
            
            # Label cell
            label_cell = table.rows[row_idx].cells[col_idx]
            label_p = label_cell.paragraphs[0]
            label_run = label_p.add_run(field.label or _to_label(field.path))
            label_run.font.bold = True
            label_run.font.size = Pt(self.style.small_size)
            label_p.paragraph_format.space_before = Pt(6)
            
            # Value cell
            value_cell = table.rows[row_idx].cells[col_idx + 1]
            value_p = value_cell.paragraphs[0]
            placeholder = format_field_placeholder(field, self.entity_prefix)
            value_run = value_p.add_run(placeholder)
            value_run.font.size = Pt(self.style.body_size)
            value_p.paragraph_format.space_before = Pt(2)
            
            # Remove borders
            set_cell_border(label_cell, False, False, False, False)
            set_cell_border(value_cell, False, False, False, False)
    
    def _render_text(self, section: Section) -> None:
        """Render text block section with optional conditional wrapping."""
        config = section.text_config
        if not config:
            return
        
        # Map our condition operators to Kahua operators
        operator_map = {
            "not_empty": "IsNotEmpty",
            "empty": "IsEmpty",
            "equals": "Equals",
            "not_equals": "DoesNotEqual",
            "greater_than": "IsGreaterThan",
            "less_than": "IsLessThan",
            "contains": "Contains",
        }
        
        # Add IF condition if section has one
        if section.condition:
            kahua_op = operator_map.get(section.condition.op.value, "IsNotEmpty")
            if_marker = build_if_condition(
                section.condition.field,
                kahua_op,
                section.condition.value,
                self.entity_prefix,
            )
            p = self.doc.add_paragraph()
            p.add_run(if_marker)
            
            # [THEN]
            p = self.doc.add_paragraph()
            p.add_run(build_then())
        
        if section.title:
            self._add_section_title(section.title)
        
        content = self._convert_template_to_placeholders(config.content)
        p = self.doc.add_paragraph()
        run = p.add_run(content)
        run.font.size = Pt(self.style.body_size)
        
        # Close conditional if we opened one
        if section.condition:
            # [ENDTHEN]
            p = self.doc.add_paragraph()
            p.add_run(build_end_then())
            
            # [ENDIF]
            p = self.doc.add_paragraph()
            p.add_run(build_endif())
    
    def _render_table(self, section: Section) -> None:
        """Render table section with StartTable/EndTable markers and conditional wrapping."""
        config = section.table_config
        if not config or not config.columns:
            return
        
        # Map our condition operators to Kahua operators
        operator_map = {
            "not_empty": "IsNotEmpty",
            "empty": "IsEmpty",
            "equals": "Equals",
            "not_equals": "DoesNotEqual",
            "greater_than": "IsGreaterThan",
            "less_than": "IsLessThan",
            "contains": "Contains",
        }
        
        # Add IF condition if section has one
        if section.condition:
            kahua_op = operator_map.get(section.condition.op.value, "IsNotEmpty")
            if_marker = build_if_condition(
                section.condition.field,
                kahua_op,
                section.condition.value,
                self.entity_prefix,
            )
            p = self.doc.add_paragraph()
            p.add_run(if_marker)
            
            # [THEN]
            p = self.doc.add_paragraph()
            p.add_run(build_then())
        
        if section.title:
            self._add_section_title(section.title)
        
        # Extract collection name from source path
        source_parts = config.source.split(".")
        table_name = source_parts[-1]
        
        # StartTable marker (hidden/small font)
        start_marker = build_start_table(
            name=table_name,
            source_path=config.source,
            rows_in_header=1,
            entity_prefix=self.entity_prefix,
        )
        p = self.doc.add_paragraph()
        run = p.add_run(start_marker)
        run.font.size = Pt(2)  # Nearly invisible
        
        # Create table with header + template row
        num_cols = len(config.columns)
        table = self.doc.add_table(rows=2, cols=num_cols)
        table.autofit = True
        
        # Style header row
        header_row = table.rows[0]
        for i, col in enumerate(config.columns):
            cell = header_row.cells[i]
            p = cell.paragraphs[0]
            run = p.add_run(col.field.label or _to_label(col.field.path))
            run.font.bold = True
            run.font.size = Pt(self.style.body_size)
            
            # Header styling
            set_cell_border(cell, top=False, bottom=True, left=False, right=False)
            
            # Alignment
            if col.alignment == Alignment.CENTER:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif col.alignment == Alignment.RIGHT:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Template row with placeholders
        data_row = table.rows[1]
        for i, col in enumerate(config.columns):
            cell = data_row.cells[i]
            p = cell.paragraphs[0]
            
            # For table rows, path is relative to the collection item
            # e.g., "Items.Description" -> just "Description" in table context
            # But we need to handle nested paths like "Company.ShortLabel"
            field_path = col.field.path
            
            # Build placeholder with table context (relative path)
            placeholder = format_field_placeholder(col.field, table_name, in_table=True)
            run = p.add_run(placeholder)
            run.font.size = Pt(self.style.body_size)
            
            # Borders and alignment
            set_cell_border(cell, top=True, bottom=True, left=False, right=False)
            
            if col.alignment == Alignment.CENTER:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif col.alignment == Alignment.RIGHT:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # EndTable marker
        p = self.doc.add_paragraph()
        run = p.add_run(build_end_table())
        run.font.size = Pt(2)
        
        # Close conditional if we opened one
        if section.condition:
            # [ENDTHEN]
            p = self.doc.add_paragraph()
            p.add_run(build_end_then())
            
            # [ENDIF]
            p = self.doc.add_paragraph()
            p.add_run(build_endif())
    
    def _render_divider(self, section: Section) -> None:
        """Render a horizontal divider."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        # Add border bottom to paragraph
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:color="{self.style.table_border_color.lstrip("#")}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    
    def _render_spacer(self, section: Section) -> None:
        """Render vertical whitespace."""
        config = section.spacer_config
        height = config.height if config else 0.25
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Inches(height)
    
    def _convert_template_to_placeholders(self, template_str: str) -> str:
        """Convert {field} template strings to [Attribute()] placeholders."""
        def replace(match: re.Match) -> str:
            path = match.group(1)
            return build_attribute_placeholder(path, self.entity_prefix)
        
        return re.sub(r"\{([^}]+)\}", replace, template_str)


def _to_label(path: str) -> str:
    """Convert path to human-readable label."""
    # Take last segment
    name = path.split(".")[-1]
    # Convert camelCase/PascalCase to spaces
    name = re.sub(r"([a-z])([A-Z])", r"\1 \2", name)
    # Convert snake_case to spaces
    name = name.replace("_", " ")
    return name.title()


# =============================================================================
# Convenience Functions
# =============================================================================

def render_template_to_docx(
    template: PortableViewTemplate,
    output_path: Optional[Union[str, Path]] = None,
) -> Union[Document, None]:
    """Render a template to DOCX.
    
    Args:
        template: The PortableViewTemplate to render
        output_path: Optional file path to save. If None, returns Document object.
    
    Returns:
        Document object if no output_path, None otherwise
    """
    renderer = DocxRenderer(template)
    
    if output_path:
        renderer.render_to_file(output_path)
        return None
    
    return renderer.render()


def render_from_json(
    json_path: Union[str, Path],
    output_path: Optional[Union[str, Path]] = None,
) -> Union[Document, None]:
    """Load template from JSON and render to DOCX."""
    import json
    
    with open(json_path, "r") as f:
        data = json.load(f)
    
    template = PortableViewTemplate(**data)
    return render_template_to_docx(template, output_path)


# =============================================================================
# Demo
# =============================================================================

if __name__ == "__main__":
    from pathlib import Path
    from template_generator import generate_template
    from schema_introspector import get_available_schemas
    
    output_dir = Path("pv_templates/generated")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    schemas = get_available_schemas()
    
    for name, schema in schemas.items():
        print(f"Generating DOCX for: {schema.name}")
        
        # Generate template from schema
        template = generate_template(schema)
        
        # Render to DOCX
        output_path = output_dir / f"{name}_template.docx"
        render_template_to_docx(template, output_path)
        
        print(f"  Saved: {output_path}")
    
    print("\nDone! Templates saved to pv_templates/generated/")
