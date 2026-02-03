"""
DOCX Token Injector

Intelligently analyzes uploaded Word documents and injects Kahua tokens
based on detected patterns like "Label: " followed by whitespace.

Workflow:
1. Parse DOCX structure (paragraphs, tables, runs)
2. Detect blank/placeholder patterns: "ID: ", "Name: ___", "Date: [blank]"
3. Map labels to appropriate entity attributes using schema introspection
4. Inject Kahua tokens while preserving document formatting

This enables users to upload "blank" templates and have the agent
automatically add the correct tokens based on context.

Token Syntax Reference:
- [Attribute(FieldName)] - Text fields
- [Date(Source=Attribute,Path=DueDate,Format="d")] - Date fields
- [Currency(Source=Attribute,Path=Amount,Format="C2")] - Currency fields
- [Number(Source=Attribute,Path=Quantity,Format="N0")] - Number fields
- [Boolean(Source=Attribute,Path=IsComplete,TrueValue="Yes",FalseValue="No")] - Boolean fields
- [StartTable(...)]...[EndTable] - Tables/Collections
- [CompanyLogo(Width=100,Height=50)] - Logo placeholder
- [ReportModifiedTimeStamp] - Timestamp
"""

import io
import re
import logging
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional, List, Dict, Any, Tuple
from copy import deepcopy

from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table
from docx.oxml.ns import qn

# Import Kahua token builders for standards-compliant token generation
try:
    from report_genius.templates.kahua_tokens import (
        build_attribute_token,
        build_boolean_token,
        build_currency_token,
        build_date_token,
        build_number_token,
        build_field_token,
        to_kahua_path,
        DateFormat,
        NumberFormat,
        CurrencyFormat,
    )
except ImportError:
    try:
        from template_gen.kahua_tokens import (
            build_attribute_token,
            build_boolean_token,
            build_currency_token,
            build_date_token,
            build_number_token,
            build_field_token,
            to_kahua_path,
            DateFormat,
            NumberFormat,
            CurrencyFormat,
        )
    except ImportError:
        # Fallback for standalone usage
        from kahua_tokens import (
            build_attribute_token,
            build_boolean_token,
            build_currency_token,
            build_date_token,
            build_number_token,
            build_field_token,
            to_kahua_path,
            DateFormat,
            NumberFormat,
            CurrencyFormat,
        )

log = logging.getLogger("docx_token_injector")


# ============== Pattern Definitions ==============

# Patterns for detecting blank/placeholder areas (start of line)
BLANK_PATTERNS = [
    r'^(.+?):\s*$',                          # "ID: " with trailing whitespace
    r'^(.+?):\s*_{2,}\s*$',                  # "ID: ______"
    r'^(.+?):\s*\[.{0,15}\]\s*$',            # "ID: [blank]" or "ID: []"
    r'^(.+?):\s*\.{3,}\s*$',                 # "ID: ..."
    r'^(.+?):\s*<.+?>\s*$',                  # "ID: <value>"
    r'^(.+?):\s*\{.+?\}\s*$',                # "ID: {placeholder}"
]

# Currency row patterns - rows that end with just "$" or "$ " needing a currency value
CURRENCY_ROW_PATTERNS = [
    r'^(.+?)\s+\$\s*$',                       # "Original Contract Sum              $"
    r'^(.+?)\s{2,}\$\s*$',                    # "Label    $" (2+ spaces before $)
]

# Design hint patterns - italic/gray hints that should be removed or replaced
DESIGN_HINT_PATTERNS = [
    r'\[([A-Za-z][A-Za-z\s&]+)\]',            # "[Title & Site Address]"
    r'\[([A-Za-z#/]+)\]',                     # "[BP#/Category]"
]

# Design hint -> field path mapping (common document hints)
DESIGN_HINT_MAPPINGS = {
    'title & site address': 'DomainPartition.Name',
    'title and site address': 'DomainPartition.Name',
    'site address': 'SiteAddress',
    'project name': 'DomainPartition.Name',
    'project number': 'DomainPartition.Number',
    'project no': 'DomainPartition.Number',
    'name & address': 'Company.Name',
    'name and address': 'Company.Name',
    'contractor name & address': 'ContractorCompany.Name',
    'contractor name and address': 'ContractorCompany.Name',
    'owner name & address': 'Owner.Name',
    'owner name and address': 'Owner.Name',
    'architect name & address': 'ArchitectCompany.Name',
    'bp#/category': 'BPNumber',
    'bp#': 'BPNumber',
    'date': 'Date',
}

# Label normalization mapping (common variations → canonical)
LABEL_NORMALIZATIONS = {
    # Identity fields
    'id': 'Id',
    'number': 'Number', 
    'no': 'Number',
    'no.': 'Number',
    '#': 'Number',
    'ref': 'Number',
    'reference': 'Number',
    'name': 'Name',
    'title': 'Subject',
    'subject': 'Subject',
    
    # Status fields
    'status': 'Status.Name',
    'state': 'Status.Name',
    'priority': 'Priority.Name',
    'phase': 'Phase.Name',
    
    # Date fields
    'date': 'Date',
    'due': 'DueDate',
    'due date': 'DueDate',
    'start': 'StartDate',
    'start date': 'StartDate',
    'end': 'EndDate',
    'end date': 'EndDate',
    'created': 'CreatedDateTime',
    'created date': 'CreatedDateTime',
    'modified': 'ModifiedDateTime',
    'submitted': 'DateSubmitted',
    'date submitted': 'DateSubmitted',
    'required': 'DateRequired',
    'date required': 'DateRequired',
    
    # Contact fields
    'from': 'SubmittedBy.Name',
    'submitted by': 'SubmittedBy.Name',
    'by': 'Author.Name',
    'author': 'Author.Name',
    'to': 'AssignedTo.Name',
    'assigned to': 'AssignedTo.Name',
    'assignee': 'AssignedTo.Name',
    'contact': 'Contact.Name',
    'company': 'Company.Name',
    'contractor': 'ContractorCompany.ShortLabel',
    'vendor': 'Vendor.Name',
    
    # Financial fields
    'amount': 'Amount',
    'total': 'TotalAmount',
    'cost': 'Cost',
    'value': 'Value',
    'price': 'Price',
    'original amount': 'OriginalAmount',
    'revised amount': 'RevisedAmount',
    'contract amount': 'ContractAmount',
    'original contract sum': 'OriginalContractAmount',
    'original contract amount': 'OriginalContractAmount',
    'contract sum': 'CurrentContractAmount',
    'net change': 'CostItemsTotalTotalValue',
    'net change by previously authorized change orders': 'PreviouslyApprovedChangesAmount',
    'net change by previous change orders': 'PreviouslyApprovedChangesAmount',
    'previous change orders': 'PreviouslyApprovedChangesAmount',
    'net change by this change order': 'CostItemsTotalTotalValue',
    'previously authorized change orders': 'PreviouslyApprovedChangesAmount',
    'current contract sum': 'CurrentContractAmount',
    'new contract sum': 'NewContractAmount',
    'contract sum including this change order': 'NewContractAmount',
    'the contract sum including this change order': 'NewContractAmount',
    'the new contract sum including this change order will be': 'NewContractAmount',
    'change this order': 'CostItemsTotalTotalValue',
    'will be increased': 'CostItemsTotalTotalValue',
    'will be decreased': 'CostItemsTotalTotalValue',
    'will be unchanged': 'CostItemsTotalTotalValue',
    
    # Change Order specific
    'change order no': 'Number',
    'change order no.': 'Number',
    'change order number': 'Number',
    'co no': 'Number',
    'co number': 'Number',
    'co#': 'Number',
    'modification': 'Description',
    'modification to the contract': 'Description',
    'modification to the contract as follows': 'Description',
    'the contract is changed as follows': 'Description',
    'the owner and contractor agree': 'Description',
    
    # Contract reference fields (for Change Orders referencing parent Contract)
    'contractor': 'Contract.Company.ShortLabel',
    'contractor name': 'Contract.Company.Name',
    'contractor company': 'Contract.Company.Name',
    'contract for': 'Contract.Description',
    'contract number': 'Contract.Number',
    'contract no': 'Contract.Number',
    'contract date': 'Contract.Date',
    'dsa no': 'DSANumber',
    'dsa number': 'DSANumber',
    'bp#': 'BidPackageNo',
    'bp number': 'BidPackageNo',
    'bid package': 'BidPackageNo',
    
    # Contract/Agreement date fields
    'initiation date': 'InitiationDate',
    'effective date': 'EffectiveDate',
    'application date': 'ApplicationDate',
    'date executed': 'DateExecuted',
    'executed date': 'DateExecuted',
    'approval date': 'DateApproved',
    'date approved': 'DateApproved',
    
    # Text blocks
    'description': 'Description',
    'desc': 'Description',
    'notes': 'Notes',
    'scope': 'ScopeOfWork',
    'scope of work': 'ScopeOfWork',
    'question': 'Question',
    'answer': 'Answer',
    'response': 'Response',
    'comment': 'Comment',
    'comments': 'Comments',
    'remarks': 'Remarks',
    
    # Project fields
    'project': 'DomainPartition.Name',
    'project name': 'DomainPartition.Name',
    'project no': 'DomainPartition.Number',
    'project number': 'DomainPartition.Number',
    'location': 'Location.Name',
    'address': 'Address',
    
    # Spec/type fields
    'type': 'Type.Name',
    'category': 'Category.Name',
    'discipline': 'Discipline.Name',
    'specification': 'SpecSection',
    'spec section': 'SpecSection',
    'csi': 'SpecSection',
}

# Field type inference (determines token format)
DATE_FIELD_PATTERNS = {'date', 'datetime', 'due', 'start', 'end', 'created', 'modified', 'submitted', 'required', 'executed', 'approved', 'completed'}
CURRENCY_FIELD_PATTERNS = {'amount', 'total', 'cost', 'value', 'price', 'fee', 'budget', 'sum', 'balance'}
NUMBER_FIELD_PATTERNS = {'count', 'qty', 'quantity', 'days', 'hours', 'number', 'period', 'lead_time'}
BOOLEAN_FIELD_PATTERNS = {'is_', 'has_', 'can_', 'should_', 'allow', 'enable', 'active', 'complete'}


@dataclass
class DetectedPlaceholder:
    """A detected placeholder location in the document."""
    paragraph_idx: int
    run_idx: int
    label: str
    original_text: str
    suggested_path: str
    suggested_token: str
    confidence: float
    is_table_cell: bool = False
    table_idx: Optional[int] = None
    row_idx: Optional[int] = None
    cell_idx: Optional[int] = None


@dataclass
class InjectionPlan:
    """Plan for injecting tokens into a document."""
    placeholders: List[DetectedPlaceholder]
    document_summary: Dict[str, Any]
    entity_def: str
    warnings: List[str] = field(default_factory=list)
    suggestions: List[str] = field(default_factory=list)


@dataclass
class InjectionResult:
    """Result of token injection."""
    success: bool
    modified_document: Optional[bytes] = None
    changes_made: List[str] = field(default_factory=list)
    tokens_injected: int = 0
    warnings: List[str] = field(default_factory=list)
    error: Optional[str] = None


# ============== Label Matching ==============

def normalize_label(label: str) -> str:
    """Normalize a detected label to a canonical form."""
    label_clean = label.strip().lower()
    
    # Direct match
    if label_clean in LABEL_NORMALIZATIONS:
        return LABEL_NORMALIZATIONS[label_clean]
    
    # Fuzzy match (remove common suffixes/prefixes)
    label_stripped = re.sub(r'^(the|a|an)\s+', '', label_clean)
    label_stripped = re.sub(r'\s+(field|value|info|data)$', '', label_stripped)
    
    if label_stripped in LABEL_NORMALIZATIONS:
        return LABEL_NORMALIZATIONS[label_stripped]
    
    # Convert to PascalCase as fallback using kahua_tokens utility
    return to_kahua_path(label_clean)


def infer_field_format(field_path: str, label: str) -> str:
    """Infer the appropriate Kahua token format for a field."""
    label_lower = label.lower()
    path_lower = field_path.lower()
    
    # Check for boolean patterns first
    if any(p in label_lower or p in path_lower for p in BOOLEAN_FIELD_PATTERNS):
        return 'boolean'
    
    # Check for date patterns
    if any(p in label_lower or p in path_lower for p in DATE_FIELD_PATTERNS):
        if 'time' in label_lower or 'datetime' in path_lower:
            return 'datetime'
        return 'date'
    
    # Check for currency patterns
    if any(p in label_lower or p in path_lower for p in CURRENCY_FIELD_PATTERNS):
        return 'currency'
    
    # Check for number patterns (but not "Number" as ID)
    if any(p in label_lower for p in NUMBER_FIELD_PATTERNS):
        if label_lower != 'number' and path_lower != 'number':
            return 'number'
    
    return 'text'


def build_kahua_token(field_path: str, format_type: str) -> str:
    """
    Build the appropriate Kahua token syntax for a field.
    
    Uses the kahua_tokens module for standards-compliant token generation.
    """
    # Use the unified build_field_token function from kahua_tokens
    return build_field_token(field_path, format_type)


# ============== Document Analysis ==============

# Inline patterns for sentences with blanks (e.g., "submitted by _____")
INLINE_BLANK_PATTERNS = [
    # "submitted by _____" or "assigned to _____"
    (r'(submitted|sent|prepared|created)\s+by\s+[_\.]{2,}', 'SubmittedBy.Name'),
    (r'assigned\s+to\s+[_\.]{2,}', 'AssignedTo.Name'),
    (r'approved\s+by\s+[_\.]{2,}', 'ApprovedBy.Name'),
    (r'reviewed\s+by\s+[_\.]{2,}', 'ReviewedBy.Name'),
    # "on [date]" patterns
    (r'on\s+date\s*[_\.]{2,}', 'Date'),
    (r'dated?\s*[_\.]{2,}', 'Date'),
    # "for project _____"
    (r'for\s+project\s+[_\.]{2,}', 'DomainPartition.Name'),
    (r'project\s*(name|number)?\s*[_\.]{2,}', 'DomainPartition.Name'),
]


def detect_currency_row(text: str) -> Optional[Tuple[str, str]]:
    """
    Detect if text is a financial row ending with just '$'.
    
    Returns (label, field_path) if found.
    """
    for pattern in CURRENCY_ROW_PATTERNS:
        match = re.match(pattern, text.strip(), re.IGNORECASE)
        if match:
            label = match.group(1).strip()
            label_lower = label.lower()
            # Look up in normalizations
            field_path = LABEL_NORMALIZATIONS.get(label_lower)
            if not field_path:
                # Convert to PascalCase
                field_path = ''.join(word.capitalize() for word in label.split())
            return label, field_path
    return None


def detect_design_hint(text: str) -> Optional[Tuple[str, str, int, int]]:
    """
    Detect design hint text like [Title & Site Address].
    
    Returns (hint_text, field_path, start_pos, end_pos) if found.
    """
    for pattern in DESIGN_HINT_PATTERNS:
        match = re.search(pattern, text)
        if match:
            hint_text = match.group(1)
            hint_lower = hint_text.lower()
            # Look up in hint mappings first
            field_path = DESIGN_HINT_MAPPINGS.get(hint_lower)
            if not field_path:
                # Fallback to PascalCase
                field_path = ''.join(word.capitalize() for word in hint_text.replace('&', 'And').split())
            return hint_text, field_path, match.start(), match.end()
    return None


def detect_placeholder_in_text(text: str) -> Optional[Tuple[str, str]]:
    """
    Detect if text contains a placeholder pattern.
    
    Returns (label, remaining_text) if found, None otherwise.
    """
    for pattern in BLANK_PATTERNS:
        match = re.match(pattern, text, re.IGNORECASE)
        if match:
            label = match.group(1).strip()
            return label, text[match.end():]
    
    # Also check for inline patterns: "The ID is: " or "Subject: [value]"
    # Pattern: text ending with ": " and nothing substantial after
    inline_match = re.search(r'([A-Za-z][A-Za-z\s]{1,25}):\s*$', text)
    if inline_match:
        label = inline_match.group(1).strip()
        return label, ''
    
    return None


def detect_inline_blanks(text: str) -> List[Tuple[str, str, int, int]]:
    """
    Detect inline blanks in a sentence.
    
    Returns list of (match_text, field_path, start_pos, end_pos)
    """
    results = []
    for pattern, field_path in INLINE_BLANK_PATTERNS:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            results.append((match.group(0), field_path, match.start(), match.end()))
    return results


def analyze_paragraph(para: Paragraph, para_idx: int) -> List[DetectedPlaceholder]:
    """Analyze a paragraph for placeholder patterns."""
    placeholders = []
    
    # Get full paragraph text
    full_text = para.text.strip()
    if not full_text:
        return placeholders
    
    # Check for currency row pattern (e.g., "Original Contract Sum    $")
    currency_detection = detect_currency_row(full_text)
    if currency_detection:
        label, field_path = currency_detection
        token = build_kahua_token(field_path, 'currency')
        
        placeholders.append(DetectedPlaceholder(
            paragraph_idx=para_idx,
            run_idx=0,
            label=label,
            original_text=full_text,
            suggested_path=field_path,
            suggested_token=token,
            confidence=0.85 if label.lower() in LABEL_NORMALIZATIONS else 0.65
        ))
        return placeholders  # Don't process further if it's a currency row
    
    # Check for design hints like [Title & Site Address]
    hint_detection = detect_design_hint(full_text)
    if hint_detection:
        hint_text, field_path, start, end = hint_detection
        token = build_kahua_token(field_path, 'text')
        
        # Increase confidence if we have a known mapping
        confidence = 0.8 if hint_text.lower() in DESIGN_HINT_MAPPINGS else 0.5
        
        placeholders.append(DetectedPlaceholder(
            paragraph_idx=para_idx,
            run_idx=0,
            label=hint_text,
            original_text=full_text,
            suggested_path=field_path,
            suggested_token=token,
            confidence=confidence
        ))
        # Continue to check for other patterns too
    
    # Check for whole-paragraph pattern first
    detection = detect_placeholder_in_text(full_text)
    if detection:
        label, _ = detection
        field_path = normalize_label(label)
        format_type = infer_field_format(field_path, label)
        token = build_kahua_token(field_path, format_type)
        
        # Find which run contains the label
        run_idx = 0
        for i, run in enumerate(para.runs):
            if label.lower() in run.text.lower():
                run_idx = i
                break
        
        placeholders.append(DetectedPlaceholder(
            paragraph_idx=para_idx,
            run_idx=run_idx,
            label=label,
            original_text=full_text,
            suggested_path=field_path,
            suggested_token=token,
            confidence=0.9 if label.lower() in LABEL_NORMALIZATIONS else 0.6
        ))
    
    return placeholders


def is_currency_placeholder_cell(text: str) -> bool:
    """Check if a cell contains just '$' or similar currency placeholder."""
    stripped = text.strip()
    return stripped in ('$', '$ ', '$_', '$0.00', '$-', '$—', '$ -', '$ —')


def analyze_table(table: Table, table_idx: int) -> List[DetectedPlaceholder]:
    """Analyze a table for placeholder patterns, including row-level context."""
    placeholders = []
    
    for row_idx, row in enumerate(table.rows):
        cells = row.cells
        num_cells = len(cells)
        
        # FIRST PASS: Check for two-cell financial rows (Label | $)
        # Common pattern: cell 0 = "Original Contract Sum", cell 1 = "$"
        if num_cells >= 2:
            cell0_text = cells[0].paragraphs[0].text.strip() if cells[0].paragraphs else ""
            cell1_text = cells[1].paragraphs[0].text.strip() if cells[1].paragraphs else ""
            
            # Check if this is a label + currency placeholder row
            if cell0_text and is_currency_placeholder_cell(cell1_text):
                label = cell0_text
                label_lower = label.lower()
                # Look up field path from normalizations
                field_path = LABEL_NORMALIZATIONS.get(label_lower)
                if not field_path:
                    # Convert label to PascalCase 
                    field_path = ''.join(word.capitalize() for word in label.split())
                
                token = build_kahua_token(field_path, 'currency')
                
                placeholders.append(DetectedPlaceholder(
                    paragraph_idx=0,
                    run_idx=0,
                    label=label,
                    original_text=f"{cell0_text} | {cell1_text}",
                    suggested_path=field_path,
                    suggested_token=token,
                    confidence=0.85 if label_lower in LABEL_NORMALIZATIONS else 0.65,
                    is_table_cell=True,
                    table_idx=table_idx,
                    row_idx=row_idx,
                    cell_idx=1  # Token goes in cell 1 (the $ cell)
                ))
                continue  # Skip further analysis of this row
        
        # SECOND PASS: Analyze each cell individually
        for cell_idx, cell in enumerate(cells):
            for para_idx, para in enumerate(cell.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                
                # Skip standalone $ cells - handled at row level above
                if is_currency_placeholder_cell(text):
                    continue
                
                # Check for currency row pattern (single cell with "Label   $")
                currency_detection = detect_currency_row(text)
                if currency_detection:
                    label, field_path = currency_detection
                    token = build_kahua_token(field_path, 'currency')
                    
                    placeholders.append(DetectedPlaceholder(
                        paragraph_idx=para_idx,
                        run_idx=0,
                        label=label,
                        original_text=text,
                        suggested_path=field_path,
                        suggested_token=token,
                        confidence=0.85 if label.lower() in LABEL_NORMALIZATIONS else 0.65,
                        is_table_cell=True,
                        table_idx=table_idx,
                        row_idx=row_idx,
                        cell_idx=cell_idx
                    ))
                    continue
                
                # Check for design hints like [Project Number]
                hint_detection = detect_design_hint(text)
                if hint_detection:
                    hint_text, field_path, start, end = hint_detection
                    token = build_kahua_token(field_path, 'text')
                    confidence = 0.8 if hint_text.lower() in DESIGN_HINT_MAPPINGS else 0.5
                    
                    placeholders.append(DetectedPlaceholder(
                        paragraph_idx=para_idx,
                        run_idx=0,
                        label=hint_text,
                        original_text=text,
                        suggested_path=field_path,
                        suggested_token=token,
                        confidence=confidence,
                        is_table_cell=True,
                        table_idx=table_idx,
                        row_idx=row_idx,
                        cell_idx=cell_idx
                    ))
                    continue  # Hints are full cell replacement
                
                # Check for standard label:value pattern
                detection = detect_placeholder_in_text(text)
                if detection:
                    label, _ = detection
                    field_path = normalize_label(label)
                    format_type = infer_field_format(field_path, label)
                    token = build_kahua_token(field_path, format_type)
                    
                    placeholders.append(DetectedPlaceholder(
                        paragraph_idx=para_idx,
                        run_idx=0,
                        label=label,
                        original_text=text,
                        suggested_path=field_path,
                        suggested_token=token,
                        confidence=0.8 if label.lower() in LABEL_NORMALIZATIONS else 0.5,
                        is_table_cell=True,
                        table_idx=table_idx,
                        row_idx=row_idx,
                        cell_idx=cell_idx
                    ))
    
    return placeholders


def analyze_document(doc_bytes: bytes, entity_def: str = "") -> InjectionPlan:
    """
    Analyze a DOCX document and create an injection plan.
    
    Args:
        doc_bytes: The raw bytes of the DOCX file
        entity_def: Optional entity definition for context
        
    Returns:
        InjectionPlan with detected placeholders and suggestions
    """
    doc = Document(io.BytesIO(doc_bytes))
    placeholders = []
    warnings = []
    suggestions = []
    
    # Track document structure
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    
    log.info(f"Analyzing document: {para_count} paragraphs, {table_count} tables")
    
    # Analyze paragraphs
    for idx, para in enumerate(doc.paragraphs):
        detected = analyze_paragraph(para, idx)
        placeholders.extend(detected)
    
    # Analyze tables
    for idx, table in enumerate(doc.tables):
        detected = analyze_table(table, idx)
        placeholders.extend(detected)
    
    # Generate suggestions based on what was found
    if not placeholders:
        suggestions.append("No placeholder patterns detected. Consider adding labels like 'ID: ' or 'Status: ' followed by whitespace.")
    elif len(placeholders) < 3:
        suggestions.append(f"Found {len(placeholders)} placeholders. Consider adding more field labels for a complete template.")
    
    # Check for missing common fields
    found_labels = {p.label.lower() for p in placeholders}
    common_missing = []
    for label in ['id', 'number', 'status', 'date', 'description']:
        if label not in found_labels:
            common_missing.append(label.title())
    
    if common_missing:
        suggestions.append(f"Consider adding placeholders for: {', '.join(common_missing)}")
    
    # Check for low-confidence matches
    low_conf = [p for p in placeholders if p.confidence < 0.7]
    if low_conf:
        labels = [p.label for p in low_conf]
        warnings.append(f"Low-confidence field mappings: {', '.join(labels)}. Please verify these mappings.")
    
    return InjectionPlan(
        placeholders=placeholders,
        document_summary={
            'paragraph_count': para_count,
            'table_count': table_count,
            'placeholder_count': len(placeholders),
            'high_confidence': len([p for p in placeholders if p.confidence >= 0.7]),
            'low_confidence': len([p for p in placeholders if p.confidence < 0.7])
        },
        entity_def=entity_def,
        warnings=warnings,
        suggestions=suggestions
    )


# ============== Token Injection ==============

def inject_token_in_paragraph(para: Paragraph, placeholder: DetectedPlaceholder) -> bool:
    """
    Inject a token into a paragraph, replacing the blank area after the label.
    
    Preserves the label and its formatting, adds the token after.
    """
    try:
        full_text = para.text
        
        # Special case: Currency rows ending with just "$"
        # Pattern: "Original Contract Sum    $" -> "Original Contract Sum    [Currency...]"
        if full_text.rstrip().endswith('$'):
            for i, run in enumerate(para.runs):
                if '$' in run.text:
                    # Replace $ with currency token
                    run.text = run.text.replace('$', placeholder.suggested_token)
                    return True
            # Fallback: replace in last run
            if para.runs:
                para.runs[-1].text = para.runs[-1].text.rstrip().rstrip('$') + placeholder.suggested_token
                return True
        
        # Special case: Design hint text like [Title & Site Address]
        # Replace the bracketed hint with the token
        hint_match = re.search(r'\[([A-Za-z][A-Za-z\s&#/]+)\]', full_text)
        if hint_match:
            hint_text = hint_match.group(0)
            for i, run in enumerate(para.runs):
                if hint_text in run.text or hint_match.group(1) in run.text:
                    run.text = run.text.replace(hint_text, placeholder.suggested_token)
                    run.text = run.text.replace(f'[{hint_match.group(1)}]', placeholder.suggested_token)
                    return True
        
        # Standard case: Label with colon pattern
        label_with_colon = f"{placeholder.label}:"
        found_label_run = False
        clear_subsequent_runs = False
        
        for i, run in enumerate(para.runs):
            if clear_subsequent_runs:
                # Clear runs that come after the label (these are the blanks)
                run.text = ""
                continue
                
            if label_with_colon.lower() in run.text.lower():
                # Find where the label ends
                idx = run.text.lower().find(label_with_colon.lower())
                label_end = idx + len(label_with_colon)
                
                # Keep label, replace rest with token
                new_text = run.text[:label_end] + " " + placeholder.suggested_token
                run.text = new_text
                found_label_run = True
                clear_subsequent_runs = True  # Clear any remaining blank runs
        
        if found_label_run:
            return True
        
        # Try matching just the label (colon might be in different run)
        for i, run in enumerate(para.runs):
            if placeholder.label.lower() in run.text.lower() and ':' in para.text:
                # Found the label, inject token after colon
                colon_idx = para.text.find(':')
                if colon_idx >= 0:
                    # Rebuild paragraph: everything up to and including colon, then token
                    label_end_text = para.text[:colon_idx + 1]
                    
                    # Clear all runs and rebuild
                    for j, r in enumerate(para.runs):
                        if j == 0:
                            r.text = label_end_text + " " + placeholder.suggested_token
                        else:
                            r.text = ""
                    return True
        
        # Fallback: Just append token to last run, stripping any blanks
        if para.runs:
            last_run = para.runs[-1]
            # Strip common blank patterns
            cleaned = re.sub(r'[_\.]{2,}|\[\s*\w*\s*\]|<\s*\w*\s*>', '', last_run.text)
            last_run.text = cleaned.rstrip() + " " + placeholder.suggested_token
            return True
        
        # Last resort: add new run
        para.add_run(" " + placeholder.suggested_token)
        return True
        
    except Exception as e:
        log.error(f"Failed to inject token: {e}")
        return False


def inject_token_in_table_cell(
    doc: Document, 
    placeholder: DetectedPlaceholder
) -> bool:
    """Inject a token into a table cell."""
    try:
        table = doc.tables[placeholder.table_idx]
        cell = table.rows[placeholder.row_idx].cells[placeholder.cell_idx]
        para = cell.paragraphs[placeholder.paragraph_idx]
        
        return inject_token_in_paragraph(para, placeholder)
        
    except Exception as e:
        log.error(f"Failed to inject token in table: {e}")
        return False


def inject_tokens(doc_bytes: bytes, plan: InjectionPlan) -> InjectionResult:
    """
    Execute an injection plan, inserting tokens into the document.
    
    Args:
        doc_bytes: Original document bytes
        plan: The injection plan from analyze_document
        
    Returns:
        InjectionResult with modified document and summary
    """
    try:
        doc = Document(io.BytesIO(doc_bytes))
        changes_made = []
        tokens_injected = 0
        warnings = list(plan.warnings)  # Copy existing warnings
        
        # Sort placeholders by location (tables first, then paragraphs in reverse order)
        # Processing in reverse avoids index shifting issues
        para_placeholders = [p for p in plan.placeholders if not p.is_table_cell]
        table_placeholders = [p for p in plan.placeholders if p.is_table_cell]
        
        # Process table placeholders
        for ph in table_placeholders:
            success = inject_token_in_table_cell(doc, ph)
            if success:
                tokens_injected += 1
                changes_made.append(f"Injected {ph.suggested_token} for '{ph.label}' in table")
            else:
                warnings.append(f"Failed to inject token for '{ph.label}' in table")
        
        # Process paragraph placeholders (in reverse order)
        for ph in sorted(para_placeholders, key=lambda x: x.paragraph_idx, reverse=True):
            para = doc.paragraphs[ph.paragraph_idx]
            success = inject_token_in_paragraph(para, ph)
            if success:
                tokens_injected += 1
                changes_made.append(f"Injected {ph.suggested_token} for '{ph.label}'")
            else:
                warnings.append(f"Failed to inject token for '{ph.label}'")
        
        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return InjectionResult(
            success=True,
            modified_document=output.getvalue(),
            changes_made=changes_made,
            tokens_injected=tokens_injected,
            warnings=warnings
        )
        
    except Exception as e:
        log.error(f"Token injection failed: {e}")
        return InjectionResult(
            success=False,
            error=str(e),
            warnings=[f"Injection failed: {e}"]
        )


# ============== High-Level API ==============

def analyze_and_inject(
    doc_bytes: bytes,
    entity_def: str = "",
    auto_inject: bool = False,
    schema_fields: Optional[List[Dict[str, str]]] = None
) -> Dict[str, Any]:
    """
    Analyze a document and optionally inject tokens.
    
    Args:
        doc_bytes: The uploaded DOCX file bytes
        entity_def: The target Kahua entity definition
        auto_inject: If True, automatically inject tokens
        schema_fields: Optional list of available schema fields for better matching
        
    Returns:
        Analysis results with injection plan and optionally modified document
    """
    # Enhance label mappings with schema fields if provided
    if schema_fields:
        _enhance_mappings_from_schema(schema_fields)
    
    # Analyze the document
    plan = analyze_document(doc_bytes, entity_def)
    
    result = {
        'analysis': {
            'document_summary': plan.document_summary,
            'entity_def': plan.entity_def,
            'placeholders': [
                {
                    'label': p.label,
                    'original_text': p.original_text,
                    'suggested_path': p.suggested_path,
                    'suggested_token': p.suggested_token,
                    'confidence': p.confidence,
                    'location': 'table' if p.is_table_cell else 'paragraph'
                }
                for p in plan.placeholders
            ],
            'warnings': plan.warnings,
            'suggestions': plan.suggestions
        }
    }
    
    if auto_inject and plan.placeholders:
        injection = inject_tokens(doc_bytes, plan)
        result['injection'] = {
            'success': injection.success,
            'tokens_injected': injection.tokens_injected,
            'changes_made': injection.changes_made,
            'warnings': injection.warnings,
            'error': injection.error
        }
        if injection.success and injection.modified_document:
            result['modified_document'] = injection.modified_document
    
    return result


def _enhance_mappings_from_schema(schema_fields: List[Dict[str, str]]) -> None:
    """Add schema field labels to the normalization mappings."""
    for field in schema_fields:
        label = field.get('label', '').lower()
        path = field.get('path', field.get('name', ''))
        
        if label and path and label not in LABEL_NORMALIZATIONS:
            LABEL_NORMALIZATIONS[label] = path


# ============== Additional Aesthetic Enhancements ==============

def add_logo_placeholder(doc_bytes: bytes, position: str = 'header') -> bytes:
    """
    Add a logo placeholder token to the document.
    
    Args:
        doc_bytes: Original document bytes
        position: Where to add logo ('header', 'top', 'footer')
        
    Returns:
        Modified document bytes with logo token
    """
    doc = Document(io.BytesIO(doc_bytes))
    logo_token = '[CompanyLogo(Height=60,Width=60)]'
    
    if position == 'header':
        # Add to document header
        section = doc.sections[0]
        header = section.header
        if not header.paragraphs:
            header.add_paragraph()
        header.paragraphs[0].insert_paragraph_before(logo_token)
        
    elif position == 'top':
        # Add as first paragraph
        if doc.paragraphs:
            doc.paragraphs[0].insert_paragraph_before(logo_token)
        else:
            doc.add_paragraph(logo_token)
            
    elif position == 'footer':
        section = doc.sections[0]
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        footer.paragraphs[0].text = logo_token
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def add_timestamp_token(doc_bytes: bytes, position: str = 'footer') -> bytes:
    """
    Add a timestamp token to the document.
    
    Args:
        doc_bytes: Original document bytes
        position: Where to add timestamp ('header', 'footer', 'bottom')
        
    Returns:
        Modified document bytes with timestamp token
    """
    doc = Document(io.BytesIO(doc_bytes))
    timestamp_token = '[ReportModifiedTimeStamp]'
    
    if position == 'footer':
        section = doc.sections[0]
        footer = section.footer
        para = footer.add_paragraph()
        para.text = f"Generated: {timestamp_token}"
        
    elif position == 'header':
        section = doc.sections[0]
        header = section.header
        para = header.add_paragraph()
        para.text = timestamp_token
        
    elif position == 'bottom':
        para = doc.add_paragraph()
        para.text = f"\nGenerated: {timestamp_token}"
    
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


# ============== CLI / Debug ==============

def main(argv: Optional[List[str]] = None) -> int:
    import sys

    args = list(sys.argv[1:] if argv is None else argv)
    if len(args) < 1:
        print("Usage: python -m report_genius.injection.docx_token_injector <path_to_docx>")
        return 1

    doc_path = Path(args[0])
    if not doc_path.exists():
        print(f"File not found: {doc_path}")
        return 1

    logging.basicConfig(level=logging.INFO)

    with open(doc_path, 'rb') as f:
        doc_bytes = f.read()

    result = analyze_and_inject(doc_bytes, auto_inject=False)

    print("\n=== Document Analysis ===")
    print(f"Paragraphs: {result['analysis']['document_summary']['paragraph_count']}")
    print(f"Tables: {result['analysis']['document_summary']['table_count']}")
    print(f"Placeholders found: {result['analysis']['document_summary']['placeholder_count']}")

    print("\n=== Detected Placeholders ===")
    for p in result['analysis']['placeholders']:
        print(f"  {p['label']:20} → {p['suggested_token']}")
        print(f"    Confidence: {p['confidence']:.0%}, Location: {p['location']}")

    if result['analysis']['suggestions']:
        print("\n=== Suggestions ===")
        for s in result['analysis']['suggestions']:
            print(f"  • {s}")

    if result['analysis']['warnings']:
        print("\n=== Warnings ===")
        for w in result['analysis']['warnings']:
            print(f"  ⚠ {w}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
