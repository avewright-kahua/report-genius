"""
Portable View Template Renderer
Generates Word documents from PortableTemplate definitions and entity data.
"""

import io
import re
import logging
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pv_template_schema import (
    PortableTemplate, Section, SectionType, FieldMapping, FieldFormat,
    HeaderSection, DetailSection, TableSection, TextSection, ChartSection,
    ImageSection, SignatureSection, ColumnDef, Alignment, PageLayout, 
    Orientation, StyleConfig
)

log = logging.getLogger("pv_template_renderer")

# Default output directory
DEFAULT_OUTPUT_DIR = Path(__file__).parent / "reports"


def hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


class TemplateRenderer:
    """Renders PortableTemplate + data into Word documents."""
    
    def __init__(self, output_dir: Optional[Path] = None):
        self.output_dir = output_dir or DEFAULT_OUTPUT_DIR
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def render(
        self,
        template: PortableTemplate,
        data: Dict[str, Any],
        filename: Optional[str] = None
    ) -> Tuple[Path, bytes]:
        """
        Render a template with data to a Word document.
        
        Args:
            template: The PortableTemplate to render
            data: Entity data dictionary
            filename: Optional output filename (will be generated if not provided)
        
        Returns:
            Tuple of (file path, file bytes)
        """
        doc = Document()
        
        # Apply page layout
        self._apply_layout(doc, template.layout)
        
        # Apply base styles
        self._setup_styles(doc, template.style)
        
        # Sort sections by order
        sections = sorted(template.sections, key=lambda s: s.order)
        
        # Render each section
        for section in sections:
            self._render_section(doc, section, data, template.style)
        
        # Generate filename
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_name = re.sub(r'[^\w\-]', '_', template.name)[:30]
            filename = f"{safe_name}_{timestamp}.docx"
        elif not filename.endswith('.docx'):
            filename = f"{filename}.docx"
        
        # Save
        output_path = self.output_dir / filename
        doc.save(output_path)
        
        # Also return bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return output_path, buffer.read()
    
    def render_to_bytes(
        self,
        template: PortableTemplate,
        data: Dict[str, Any]
    ) -> bytes:
        """Render template to bytes without saving to disk."""
        doc = Document()
        self._apply_layout(doc, template.layout)
        self._setup_styles(doc, template.style)
        
        sections = sorted(template.sections, key=lambda s: s.order)
        for section in sections:
            self._render_section(doc, section, data, template.style)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _apply_layout(self, doc: Document, layout: PageLayout):
        """Apply page layout settings."""
        section = doc.sections[0]
        
        # Orientation
        if layout.orientation == Orientation.LANDSCAPE:
            section.orientation = WD_ORIENT.LANDSCAPE
            # Swap width and height for landscape
            section.page_width, section.page_height = section.page_height, section.page_width
        
        # Margins
        section.top_margin = Inches(layout.margin_top)
        section.bottom_margin = Inches(layout.margin_bottom)
        section.left_margin = Inches(layout.margin_left)
        section.right_margin = Inches(layout.margin_right)
    
    def _setup_styles(self, doc: Document, style: StyleConfig):
        """Set up document styles for professional appearance."""
        # Heading 1 style - Document title
        h1_style = doc.styles['Heading 1']
        h1_style.font.name = style.heading_font
        h1_style.font.size = Pt(style.title_size)
        h1_style.font.color.rgb = RGBColor(*hex_to_rgb(style.primary_color))
        h1_style.font.bold = True
        h1_style.paragraph_format.space_before = Pt(0)
        h1_style.paragraph_format.space_after = Pt(12)
        
        # Heading 2 style - Section headers
        h2_style = doc.styles['Heading 2']
        h2_style.font.name = style.heading_font
        h2_style.font.size = Pt(style.heading_size)
        h2_style.font.color.rgb = RGBColor(*hex_to_rgb(style.secondary_color))
        h2_style.font.bold = True
        h2_style.paragraph_format.space_before = Pt(18)
        h2_style.paragraph_format.space_after = Pt(6)
        
        # Normal style
        normal = doc.styles['Normal']
        normal.font.name = style.font_family
        normal.font.size = Pt(style.body_size)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15
    
    def _render_section(
        self,
        doc: Document,
        section: Section,
        data: Dict[str, Any],
        style: StyleConfig
    ):
        """Render a single section."""
        
        # Check condition if specified
        if section.condition and not self._evaluate_condition(section.condition, data):
            return
        
        # Add section title if present
        if section.title and section.type not in [SectionType.HEADER, SectionType.PAGE_BREAK]:
            heading = doc.add_heading(section.title, level=2)
        
        # Render based on type
        if section.type == SectionType.HEADER:
            self._render_header(doc, section.header_config, data, style)
        
        elif section.type == SectionType.DETAIL:
            self._render_detail(doc, section.detail_config, data, style)
        
        elif section.type == SectionType.TABLE:
            self._render_table(doc, section.table_config, data, style)
        
        elif section.type == SectionType.TEXT:
            self._render_text(doc, section.text_config, data, style)
        
        elif section.type == SectionType.CHART:
            self._render_chart(doc, section.chart_config, data, style)
        
        elif section.type == SectionType.IMAGE:
            self._render_image(doc, section.image_config, data)
        
        elif section.type == SectionType.SIGNATURE:
            self._render_signature(doc, section.signature_config, data)
        
        elif section.type == SectionType.PAGE_BREAK:
            doc.add_page_break()
        
        elif section.type == SectionType.FOOTER:
            # Footer handled separately in page setup
            pass
    
    def _render_header(
        self,
        doc: Document,
        config: HeaderSection,
        data: Dict[str, Any],
        style: StyleConfig
    ):
        """Render header section with professional styling."""
        if not config:
            return
        
        # Title with colored accent bar
        if config.title_template:
            title_text = self._resolve_template_string(config.title_template, data)
            
            # Create a table for the title block with accent bar
            title_table = doc.add_table(rows=1, cols=1)
            title_table.autofit = False
            title_table.allow_autofit = False
            
            cell = title_table.rows[0].cells[0]
            cell.width = Inches(6.5)
            
            # Add accent bar using cell shading
            self._set_cell_shading(cell, style.primary_color)
            
            # Add title text
            title_para = cell.paragraphs[0]
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title_text)
            title_run.bold = True
            title_run.font.size = Pt(style.title_size)
            title_run.font.color.rgb = RGBColor(255, 255, 255)  # White text
            title_run.font.name = style.heading_font
            
            # Add padding
            self._set_cell_padding(cell, top=12, bottom=12, left=12, right=12)
            
            doc.add_paragraph()  # Spacing after title
        
        # Fields in grid layout with improved styling
        if config.fields:
            if config.layout == "horizontal":
                self._render_fields_horizontal(doc, config.fields, data, style)
            else:
                self._render_fields_grid_styled(doc, config.fields, data, config.columns, config.show_labels, style)
    
    def _render_detail(
        self,
        doc: Document,
        config: DetailSection,
        data: Dict[str, Any],
        style: StyleConfig
    ):
        """Render detail/form section."""
        if not config or not config.fields:
            return
        
        if config.layout == "horizontal":
            self._render_fields_horizontal(doc, config.fields, data, style)
        else:
            self._render_fields_grid(doc, config.fields, data, config.columns, config.show_labels, style)
    
    def _render_fields_horizontal(
        self,
        doc: Document,
        fields: List[FieldMapping],
        data: Dict[str, Any],
        style: StyleConfig
    ):
        """Render fields in a horizontal single-line layout."""
        para = doc.add_paragraph()
        for i, field in enumerate(fields):
            if i > 0:
                para.add_run("    |    ")
            
            # Ensure field is a FieldMapping object
            field = self._ensure_field_mapping(field)
            
            label = field.label or field.path.split('.')[-1]
            value = self._get_field_value(data, field)
            
            run = para.add_run(f"{label}: ")
            run.bold = True
            para.add_run(str(value))
    
    def _render_fields_grid(
        self,
        doc: Document,
        fields: List[FieldMapping],
        data: Dict[str, Any],
        columns: int,
        show_labels: bool,
        style: StyleConfig
    ):
        """Render fields in a grid table layout (legacy - calls styled version)."""
        self._render_fields_grid_styled(doc, fields, data, columns, show_labels, style)
    
    def _render_fields_grid_styled(
        self,
        doc: Document,
        fields: List[FieldMapping],
        data: Dict[str, Any],
        columns: int,
        show_labels: bool,
        style: StyleConfig
    ):
        """Render fields in an aesthetic grid layout with subtle styling."""
        if not fields:
            return
        
        # Create table for grid - use clean borderless design
        rows_needed = (len(fields) + columns - 1) // columns
        table = doc.add_table(rows=rows_needed, cols=columns * 2)  # label + value pairs
        table.autofit = False
        table.allow_autofit = False
        
        # Set column widths (label narrower than value)
        total_width = 6.5  # inches
        label_width = total_width / (columns * 3)  # labels get 1/3
        value_width = total_width / (columns * 3) * 2  # values get 2/3
        
        for row in table.rows:
            for col_idx, cell in enumerate(row.cells):
                if col_idx % 2 == 0:  # Label column
                    cell.width = Inches(label_width)
                else:  # Value column
                    cell.width = Inches(value_width)
        
        # Remove table borders for clean look
        self._set_table_borders(table, False)
        
        # Apply alternating row shading for readability
        for row_idx, row in enumerate(table.rows):
            if row_idx % 2 == 0:
                for cell in row.cells:
                    self._set_cell_shading(cell, "#f8fafc")  # Very light gray
        
        for idx, field in enumerate(fields):
            # Ensure field is a FieldMapping object
            field = self._ensure_field_mapping(field)
            
            row_idx = idx // columns
            col_idx = (idx % columns) * 2
            
            row = table.rows[row_idx]
            
            # Label cell - styled
            label_cell = row.cells[col_idx]
            label = field.label or self._format_label(field.path)
            label_para = label_cell.paragraphs[0]
            label_para.paragraph_format.space_before = Pt(4)
            label_para.paragraph_format.space_after = Pt(4)
            label_run = label_para.add_run(label)
            label_run.bold = True
            label_run.font.size = Pt(style.body_size)
            label_run.font.color.rgb = RGBColor(*hex_to_rgb(style.primary_color))
            label_run.font.name = style.font_family
            
            # Value cell - styled
            value_cell = row.cells[col_idx + 1]
            value = self._get_field_value(data, field)
            value_para = value_cell.paragraphs[0]
            value_para.paragraph_format.space_before = Pt(4)
            value_para.paragraph_format.space_after = Pt(4)
            value_run = value_para.add_run(str(value) if value else "—")
            value_run.font.size = Pt(style.body_size)
            value_run.font.name = style.font_family
            if not value:
                value_run.font.color.rgb = RGBColor(156, 163, 175)  # Gray for empty
        
        doc.add_paragraph()  # Spacing
    
    def _format_label(self, path: str) -> str:
        """Convert camelCase/PascalCase field path to human-readable label."""
        # Get last part of path
        name = path.split('.')[-1]
        # Insert space before capitals
        import re
        label = re.sub(r'(?<!^)(?=[A-Z])', ' ', name)
        return label
    
    def _render_table(
        self,
        doc: Document,
        config: TableSection,
        data: Dict[str, Any],
        style: StyleConfig
    ):
        """Render a data table section."""
        if not config:
            return
        
        # Get data for table
        table_data = self._resolve_path(data, config.source)
        if not isinstance(table_data, list):
            table_data = [table_data] if table_data else []
        
        if not table_data:
            doc.add_paragraph(config.empty_message or "No items")
            return
        
        # Apply sorting
        if config.sort_by:
            reverse = config.sort_direction.lower() == "desc"
            table_data = sorted(
                table_data,
                key=lambda x: self._resolve_path(x, config.sort_by) or "",
                reverse=reverse
            )
        
        # Apply row limit
        if config.max_rows:
            table_data = table_data[:config.max_rows]
        
        # Create table
        num_cols = len(config.columns)
        if config.show_row_numbers:
            num_cols += 1
        
        table = doc.add_table(rows=1 + len(table_data), cols=num_cols)
        table.style = 'Table Grid'
        
        # Header row
        header_row = table.rows[0]
        col_offset = 0
        
        if config.show_row_numbers:
            header_row.cells[0].text = "#"
            self._style_header_cell(header_row.cells[0], style)
            col_offset = 1
        
        for i, col in enumerate(config.columns):
            cell = header_row.cells[i + col_offset]
            # Ensure col.field is a FieldMapping object
            col_field = self._ensure_field_mapping(col.field) if hasattr(col, 'field') else self._ensure_field_mapping(col)
            cell.text = col_field.label or col_field.path
            self._style_header_cell(cell, style)
            self._set_cell_alignment(cell, col.alignment if hasattr(col, 'alignment') else None)
        
        # Data rows
        subtotals = {f: 0.0 for f in config.subtotal_fields}
        
        for row_idx, row_data in enumerate(table_data):
            row = table.rows[row_idx + 1]
            
            # Alternating row colors
            if row_idx % 2 == 1:
                self._set_row_shading(row, style.table_alt_row_bg)
            
            if config.show_row_numbers:
                row.cells[0].text = str(row_idx + 1)
            
            for col_idx, col in enumerate(config.columns):
                cell = row.cells[col_idx + col_offset]
                col_field = self._ensure_field_mapping(col.field) if hasattr(col, 'field') else self._ensure_field_mapping(col)
                value = self._get_field_value(row_data, col_field)
                cell.text = str(value)
                self._set_cell_alignment(cell, col.alignment if hasattr(col, 'alignment') else None)
                
                # Track subtotals
                if col_field.path in config.subtotal_fields:
                    try:
                        subtotals[col_field.path] += float(self._resolve_path(row_data, col_field.path) or 0)
                    except (ValueError, TypeError):
                        pass
        
        # Subtotal row
        if config.show_subtotals and subtotals:
            subtotal_row = table.add_row()
            self._set_row_shading(subtotal_row, style.table_header_bg)
            
            first_cell = subtotal_row.cells[col_offset] if config.show_row_numbers else subtotal_row.cells[0]
            first_cell.text = "Total"
            run = first_cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(*hex_to_rgb(style.table_header_fg))
            
            for col_idx, col in enumerate(config.columns):
                col_field = self._ensure_field_mapping(col.field) if hasattr(col, 'field') else self._ensure_field_mapping(col)
                if col_field.path in subtotals:
                    cell = subtotal_row.cells[col_idx + col_offset]
                    value = subtotals[col_field.path]
                    cell.text = self._format_value(value, col_field.format)
                    self._set_cell_alignment(cell, col.alignment if hasattr(col, 'alignment') else None)
                    run = cell.paragraphs[0].runs[0]
                    run.bold = True
                    run.font.color.rgb = RGBColor(*hex_to_rgb(style.table_header_fg))
        
        doc.add_paragraph()  # Spacing
    
    def _render_text(
        self,
        doc: Document,
        config: TextSection,
        data: Dict[str, Any],
        style: StyleConfig
    ):
        """Render a text section with field substitutions."""
        if not config:
            return
        
        content = self._resolve_template_string(config.content, data)
        
        # Handle multiple paragraphs
        for para_text in content.split('\n'):
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
    
    def _render_chart(
        self,
        doc: Document,
        config: ChartSection,
        data: Dict[str, Any],
        style: StyleConfig
    ):
        """Render a chart section."""
        if not config:
            return
        
        try:
            # Import chart generation from report_generator
            from report_generator import ReportGenerator, ChartSpec
            
            # Get chart data
            chart_data = self._resolve_path(data, config.data_source)
            if isinstance(chart_data, list):
                labels = [self._resolve_path(item, config.label_field) for item in chart_data]
                values = [self._resolve_path(item, config.value_field) or 0 for item in chart_data]
            else:
                labels = []
                values = []
            
            if not labels or not values:
                doc.add_paragraph(f"[Chart: {config.title} - No data]")
                return
            
            # Generate chart
            spec = ChartSpec(
                chart_type=config.chart_type,
                title=config.title,
                data={"labels": labels, "values": values},
                colors=config.colors,
                width=config.width,
                height=config.height
            )
            
            generator = ReportGenerator()
            chart_bytes = generator.generate_chart_image(spec)
            chart_bytes.seek(0)
            
            # Add to document
            doc.add_picture(chart_bytes, width=Inches(config.width))
            
        except ImportError:
            doc.add_paragraph(f"[Chart: {config.title} - chart generation unavailable]")
        except Exception as e:
            log.error(f"Chart rendering failed: {e}")
            doc.add_paragraph(f"[Chart: {config.title} - error: {e}]")
    
    def _render_image(
        self,
        doc: Document,
        config: ImageSection,
        data: Dict[str, Any]
    ):
        """Render an image section."""
        if not config:
            return
        
        try:
            import base64
            
            image_data = None
            
            if config.source == "field" and config.path:
                # Get image from data field (base64 encoded)
                img_value = self._resolve_path(data, config.path)
                if img_value and isinstance(img_value, str):
                    image_data = io.BytesIO(base64.b64decode(img_value))
            
            elif config.source == "static" and config.url:
                # Load from URL/path
                if config.url.startswith(('http://', 'https://')):
                    import httpx
                    response = httpx.get(config.url)
                    image_data = io.BytesIO(response.content)
                else:
                    image_data = config.url  # Local file path
            
            if image_data:
                width = Inches(config.width) if config.width else None
                height = Inches(config.height) if config.height else None
                doc.add_picture(image_data, width=width, height=height)
                
                if config.caption:
                    caption = doc.add_paragraph(config.caption)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Skip silently if no image - don't show placeholder text
                
        except Exception as e:
            log.error(f"Image rendering failed: {e}")
            doc.add_paragraph(f"[Image error: {e}]")
    
    def _render_signature(
        self,
        doc: Document,
        config: SignatureSection,
        data: Dict[str, Any]
    ):
        """Render signature blocks."""
        if not config:
            return
        
        table = doc.add_table(rows=2, cols=config.columns)
        
        for i, line in enumerate(config.lines[:config.columns]):
            col_idx = i % config.columns
            
            # Signature line
            sig_cell = table.rows[0].cells[col_idx]
            sig_para = sig_cell.paragraphs[0]
            sig_para.add_run("_" * 30)
            
            # Label
            label_cell = table.rows[1].cells[col_idx]
            label = line.get("label", "Signature")
            if "path" in line:
                value = self._resolve_path(data, line["path"])
                label = f"{label}: {value}" if value else label
            label_cell.paragraphs[0].add_run(label)
        
        doc.add_paragraph()
    
    # ============== Helper Methods ==============
    
    def _get_field_value(self, data: Dict[str, Any], field: FieldMapping) -> str:
        """Get formatted field value from data."""
        # Ensure field is a FieldMapping object
        field = self._ensure_field_mapping(field)
        
        raw_value = self._resolve_path(data, field.path)
        
        if raw_value is None:
            return field.default_value or ""
        
        return self._format_value(raw_value, field.format, field.format_options)
    
    def _format_value(
        self,
        value: Any,
        fmt: FieldFormat,
        options: Dict[str, Any] = None
    ) -> str:
        """Format a value according to field format."""
        if value is None:
            return ""
        
        options = options or {}
        
        if fmt == FieldFormat.CURRENCY:
            try:
                num = float(value)
                prefix = options.get("prefix", "$")
                decimals = options.get("decimals", 2)
                return f"{prefix}{num:,.{decimals}f}"
            except (ValueError, TypeError):
                return str(value)
        
        elif fmt == FieldFormat.NUMBER:
            try:
                num = float(value)
                decimals = options.get("decimals", 0)
                return f"{num:,.{decimals}f}"
            except (ValueError, TypeError):
                return str(value)
        
        elif fmt == FieldFormat.PERCENT:
            try:
                num = float(value)
                return f"{num:.1f}%"
            except (ValueError, TypeError):
                return str(value)
        
        elif fmt == FieldFormat.DATE:
            if isinstance(value, str):
                # Try to parse and format
                try:
                    from dateutil import parser
                    dt = parser.parse(value)
                    return dt.strftime(options.get("format", "%m/%d/%Y"))
                except:
                    return value[:10] if len(value) >= 10 else value
            return str(value)
        
        elif fmt == FieldFormat.DATETIME:
            if isinstance(value, str):
                try:
                    from dateutil import parser
                    dt = parser.parse(value)
                    # Convert format string from user-friendly to strftime
                    fmt_str = options.get("format", "%m/%d/%Y %I:%M %p")
                    fmt_str = fmt_str.replace("MMM", "%b").replace("MM", "%m").replace("DD", "%d").replace("D", "%-d").replace("YYYY", "%Y").replace("YY", "%y")
                    return dt.strftime(fmt_str)
                except:
                    return value
            return str(value)
        
        elif fmt == FieldFormat.BOOLEAN:
            return "Yes" if value else "No"
        
        return str(value)
    
    def _resolve_path(self, data: Dict[str, Any], path: str) -> Any:
        """Resolve a dot-notation path in data dictionary (case-insensitive)."""
        if not path or not data:
            return None
        
        parts = path.split('.')
        current = data
        
        for part in parts:
            if current is None:
                return None
            
            # Handle array indexing: Items[0]
            match = re.match(r'(\w+)\[(\d+)\]', part)
            if match:
                key, idx = match.groups()
                current = self._get_case_insensitive(current, key)
                if isinstance(current, list) and int(idx) < len(current):
                    current = current[int(idx)]
                else:
                    return None
            elif isinstance(current, dict):
                current = self._get_case_insensitive(current, part)
            else:
                return None
        
        return current
    
    def _get_case_insensitive(self, data: Dict[str, Any], key: str) -> Any:
        """Get value from dict with case-insensitive key matching."""
        if not isinstance(data, dict):
            return None
        
        # Try exact match first
        if key in data:
            return data[key]
        
        # Try case-insensitive match
        key_lower = key.lower()
        for k, v in data.items():
            if k.lower() == key_lower:
                return v
        
        return None
    
    def _ensure_field_mapping(self, field: Any) -> FieldMapping:
        """Ensure field is a FieldMapping object, converting from dict if needed."""
        if isinstance(field, FieldMapping):
            return field
        if isinstance(field, dict):
            return FieldMapping(
                path=field.get('path', ''),
                label=field.get('label'),
                format=FieldFormat(field['format']) if field.get('format') else None,
                default=field.get('default'),
                width=field.get('width')
            )
        # Fallback: create empty field mapping
        return FieldMapping(path=str(field), label=None)
    
    def _resolve_template_string(self, template: str, data: Dict[str, Any]) -> str:
        """Replace {field_path} placeholders with values."""
        def replace_match(match):
            path = match.group(1)
            value = self._resolve_path(data, path)
            return str(value) if value is not None else ""
        
        return re.sub(r'\{([^}]+)\}', replace_match, template)
    
    def _evaluate_condition(self, condition: Dict[str, Any], data: Dict[str, Any]) -> bool:
        """Evaluate a condition against data."""
        path = condition.get("path")
        op = condition.get("op", "exists")
        value = condition.get("value")
        
        actual = self._resolve_path(data, path)
        
        if op == "exists":
            return actual is not None
        elif op == "eq":
            return actual == value
        elif op == "ne":
            return actual != value
        elif op == "gt":
            return actual is not None and actual > value
        elif op == "lt":
            return actual is not None and actual < value
        elif op == "contains":
            return value in str(actual) if actual else False
        
        return True
    
    def _style_header_cell(self, cell, style: StyleConfig):
        """Apply header styling to a table cell."""
        # Background color
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), style.table_header_bg.lstrip('#'))
        cell._tc.get_or_add_tcPr().append(shading)
        
        # Text styling
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(*hex_to_rgb(style.table_header_fg))
    
    def _set_cell_alignment(self, cell, alignment: Alignment):
        """Set cell text alignment."""
        align_map = {
            Alignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
            Alignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
            Alignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
        }
        for para in cell.paragraphs:
            para.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
    
    def _set_row_shading(self, row, color: str):
        """Set background color for entire row."""
        for cell in row.cells:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), color.lstrip('#'))
            cell._tc.get_or_add_tcPr().append(shading)
    
    def _set_table_borders(self, table, show: bool):
        """Set or remove table borders."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single' if show else 'nil')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
    
    def _set_cell_shading(self, cell, color: str):
        """Set background color for a single cell."""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color.lstrip('#'))
        cell._tc.get_or_add_tcPr().append(shading)
    
    def _set_cell_padding(self, cell, top: int = 0, bottom: int = 0, left: int = 0, right: int = 0):
        """Set cell padding in points."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        tcMar = OxmlElement('w:tcMar')
        for side, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            margin = OxmlElement(f'w:{side}')
            margin.set(qn('w:w'), str(int(value * 20)))  # Points to twips
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        
        tcPr.append(tcMar)
    
    def _add_horizontal_line(self, doc: Document, color: str = "#e2e8f0", thickness: float = 0.5):
        """Add a horizontal line/divider."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        
        # Add bottom border to paragraph
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(int(thickness * 8)))
        bottom.set(qn('w:color'), color.lstrip('#'))
        pBdr.append(bottom)
        pPr.append(pBdr)


# ============== Convenience Function ==============

def render_template(
    template: PortableTemplate,
    data: Dict[str, Any],
    output_dir: Optional[Path] = None,
    filename: Optional[str] = None
) -> Tuple[Path, bytes]:
    """
    Convenience function to render a template.
    
    Args:
        template: PortableTemplate to render
        data: Entity data dictionary
        output_dir: Optional output directory
        filename: Optional output filename
    
    Returns:
        Tuple of (output path, file bytes)
    """
    renderer = TemplateRenderer(output_dir)
    return renderer.render(template, data, filename)
