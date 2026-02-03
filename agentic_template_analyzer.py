"""
Agentic Template Analyzer

AI-powered analysis of uploaded Word documents to:
1. Extract semantic structure (sections, lists, tables, headers/footers)
2. Identify blank/placeholder patterns
3. Generate complete PortableViewTemplate with proper configs

This goes beyond simple token injection - it understands document intent
and creates a rich template specification that can be rendered with all features.
"""

import io
import re
import logging
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional, List, Dict, Any, Tuple, Literal
from copy import deepcopy

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Pt, Twips
from docx.enum.style import WD_STYLE_TYPE

# Template schema imports
from report_genius.templates import (
    PortableViewTemplate,
    Section,
    SectionType,
    HeaderConfig,
    DetailConfig,
    TextConfig,
    TableConfig,
    TableColumn,
    ListConfig,
    ImageConfig,
    DividerConfig,
    LayoutConfig,
    StyleConfig,
    PageHeaderFooterConfig,
    FieldDef,
    FieldFormat,
    Alignment,
    LayoutType,
)

log = logging.getLogger("agentic_template_analyzer")


# ============== Document Structure Analysis ==============

@dataclass
class ParsedParagraph:
    """Parsed information about a paragraph."""
    index: int
    text: str
    style_name: str
    is_heading: bool
    heading_level: int  # 0 if not heading
    is_list_item: bool
    list_type: Optional[Literal["bullet", "number"]] = None
    has_blank_pattern: bool = False
    detected_label: Optional[str] = None
    suggested_field: Optional[str] = None
    font_name: Optional[str] = None
    font_size: Optional[float] = None
    font_color: Optional[str] = None
    is_bold: bool = False
    alignment: str = "left"


@dataclass 
class ParsedTable:
    """Parsed information about a table."""
    index: int
    row_count: int
    col_count: int
    headers: List[str]
    sample_data: List[List[str]]  # First few rows
    has_blank_cells: bool = False
    looks_like_form: bool = False  # Label-value pairs
    looks_like_data_table: bool = False  # Repeating data rows


@dataclass
class ParsedSection:
    """A semantic section detected in the document."""
    type: SectionType
    title: Optional[str] = None
    paragraphs: List[ParsedParagraph] = field(default_factory=list)
    tables: List[ParsedTable] = field(default_factory=list)
    fields: List[Dict[str, Any]] = field(default_factory=list)
    list_items: List[str] = field(default_factory=list)
    list_type: Optional[str] = None
    order: int = 0


@dataclass
class DocumentAnalysis:
    """Complete analysis of an uploaded document."""
    sections: List[ParsedSection]
    page_header: Optional[Dict[str, str]] = None
    page_footer: Optional[Dict[str, str]] = None
    detected_entity_type: Optional[str] = None
    detected_fields: List[Dict[str, Any]] = field(default_factory=list)
    style_info: Dict[str, Any] = field(default_factory=dict)
    warnings: List[str] = field(default_factory=list)
    suggestions: List[str] = field(default_factory=list)


# ============== Pattern Detection ==============

# Label patterns for field detection
LABEL_TO_FIELD = {
    # Identity
    'id': 'Number', 'number': 'Number', 'no': 'Number', 'ref': 'Number',
    'name': 'Name', 'title': 'Subject', 'subject': 'Subject',
    
    # Status
    'status': 'Status.Name', 'state': 'Status.Name', 'priority': 'Priority.Name',
    
    # Dates
    'date': 'Date', 'due': 'DueDate', 'due date': 'DueDate',
    'start': 'StartDate', 'start date': 'StartDate',
    'end': 'EndDate', 'end date': 'EndDate',
    'created': 'CreatedDateTime', 'submitted': 'DateSubmitted',
    'required': 'DateRequired',
    
    # People
    'from': 'SubmittedBy.Name', 'submitted by': 'SubmittedBy.Name',
    'to': 'AssignedTo.Name', 'assigned to': 'AssignedTo.Name',
    'by': 'Author.Name', 'author': 'Author.Name',
    'contact': 'Contact.Name', 'company': 'Company.Name',
    
    # Financial
    'amount': 'Amount', 'total': 'TotalAmount', 'cost': 'Cost',
    'value': 'Value', 'price': 'Price',
    
    # Text
    'description': 'Description', 'notes': 'Notes', 'scope': 'Scope',
    'question': 'Question', 'answer': 'Answer', 'response': 'Response',
    'comment': 'Comment', 'remarks': 'Remarks',
}

# Field format inference
DATE_INDICATORS = {'date', 'due', 'start', 'end', 'created', 'submitted', 'required'}
CURRENCY_INDICATORS = {'amount', 'total', 'cost', 'value', 'price', 'fee', 'budget'}


def detect_blank_pattern(text: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Detect if text contains a blank/placeholder pattern.
    
    Returns (label, field_path) if found.
    """
    # Patterns: "Label: ", "Label: ______", "Label: [blank]", etc.
    patterns = [
        r'^(.+?):\s*$',                    # "ID: "
        r'^(.+?):\s*_{2,}\s*$',            # "ID: ______"
        r'^(.+?):\s*\[.{0,15}\]\s*$',      # "ID: [blank]"
        r'^(.+?):\s*\.{3,}\s*$',           # "ID: ..."
        r'^(.+?):\s*<.+?>\s*$',            # "ID: <value>"
    ]
    
    for pattern in patterns:
        match = re.match(pattern, text.strip(), re.IGNORECASE)
        if match:
            label = match.group(1).strip()
            label_lower = label.lower()
            field_path = LABEL_TO_FIELD.get(label_lower, label.title().replace(' ', ''))
            return label, field_path
    
    # Check for trailing "Label: " with nothing after
    match = re.search(r'([A-Za-z][A-Za-z\s]{1,25}):\s*$', text)
    if match:
        label = match.group(1).strip()
        label_lower = label.lower()
        field_path = LABEL_TO_FIELD.get(label_lower, label.title().replace(' ', ''))
        return label, field_path
    
    return None, None


def infer_field_format(field_path: str, label: str = "") -> FieldFormat:
    """Infer the appropriate FieldFormat for a field."""
    combined = (field_path + " " + label).lower()
    
    if any(ind in combined for ind in DATE_INDICATORS):
        return FieldFormat.DATE
    if any(ind in combined for ind in CURRENCY_INDICATORS):
        return FieldFormat.CURRENCY
    
    return FieldFormat.TEXT


# ============== Document Parsing ==============

def parse_paragraph(para: Paragraph, index: int) -> ParsedParagraph:
    """Parse a single paragraph's properties."""
    text = para.text.strip()
    style_name = para.style.name if para.style else ""
    
    # Detect heading
    is_heading = style_name.lower().startswith('heading')
    heading_level = 0
    if is_heading:
        try:
            heading_level = int(style_name.split()[-1])
        except (ValueError, IndexError):
            heading_level = 1
    
    # Detect list item
    is_list = False
    list_type = None
    
    # Check for list formatting via XML
    if para._p is not None:
        numPr = para._p.find(qn('w:numPr'))
        if numPr is not None:
            is_list = True
            # Try to determine bullet vs number
            ilvl = numPr.find(qn('w:ilvl'))
            numId = numPr.find(qn('w:numId'))
            # Simple heuristic: assume bullet unless we detect numbering
            list_type = "bullet"  # Default
    
    # Check style name for list types
    if 'bullet' in style_name.lower() or 'list bullet' in style_name.lower():
        is_list = True
        list_type = "bullet"
    elif 'number' in style_name.lower() or 'list number' in style_name.lower():
        is_list = True
        list_type = "number"
    
    # Detect blank pattern
    label, field = detect_blank_pattern(text)
    has_blank = label is not None
    
    # Extract font info from first run
    font_name = None
    font_size = None
    font_color = None
    is_bold = False
    
    if para.runs:
        run = para.runs[0]
        if run.font:
            font_name = run.font.name
            if run.font.size:
                font_size = run.font.size.pt
            if run.font.color and run.font.color.rgb:
                font_color = f"#{run.font.color.rgb}"
            is_bold = run.font.bold or False
    
    # Alignment
    alignment = "left"
    if para.alignment:
        align_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
        alignment = align_map.get(para.alignment, "left")
    
    return ParsedParagraph(
        index=index,
        text=text,
        style_name=style_name,
        is_heading=is_heading,
        heading_level=heading_level,
        is_list_item=is_list,
        list_type=list_type,
        has_blank_pattern=has_blank,
        detected_label=label,
        suggested_field=field,
        font_name=font_name,
        font_size=font_size,
        font_color=font_color,
        is_bold=is_bold,
        alignment=alignment,
    )


def parse_table(table: Table, index: int) -> ParsedTable:
    """Parse a table's structure."""
    rows = list(table.rows)
    if not rows:
        return ParsedTable(index=index, row_count=0, col_count=0, headers=[], sample_data=[])
    
    # Get dimensions
    row_count = len(rows)
    col_count = len(rows[0].cells) if rows else 0
    
    # Extract headers (first row)
    headers = [cell.text.strip() for cell in rows[0].cells]
    
    # Extract sample data (next few rows)
    sample_data = []
    for row in rows[1:4]:  # Up to 3 rows
        row_data = [cell.text.strip() for cell in row.cells]
        sample_data.append(row_data)
    
    # Detect if this looks like a form (2 columns, label-value pattern)
    looks_like_form = False
    if col_count == 2:
        # Check if first column contains labels
        label_count = sum(1 for h in headers + [r[0] for r in sample_data if r] if h.endswith(':') or h.endswith('：'))
        looks_like_form = label_count > len(sample_data) / 2
    
    # Detect if this looks like a data table (consistent structure, many rows)
    looks_like_data_table = row_count > 3 and col_count >= 3
    
    # Check for blank cells
    has_blank = any(
        not cell.strip() or cell.strip() in ['___', '...', '[blank]', '<value>']
        for row in sample_data
        for cell in row
    )
    
    return ParsedTable(
        index=index,
        row_count=row_count,
        col_count=col_count,
        headers=headers,
        sample_data=sample_data,
        has_blank_cells=has_blank,
        looks_like_form=looks_like_form,
        looks_like_data_table=looks_like_data_table,
    )


def extract_header_footer(doc: DocxDocument) -> Tuple[Optional[Dict], Optional[Dict]]:
    """Extract page header and footer content."""
    header_info = None
    footer_info = None
    
    try:
        for section in doc.sections:
            # Header
            if section.header and section.header.paragraphs:
                header_text = " ".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
                if header_text:
                    header_info = {"text": header_text}
            
            # Footer
            if section.footer and section.footer.paragraphs:
                footer_text = " ".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
                if footer_text:
                    footer_info = {"text": footer_text}
            
            break  # Just check first section
    except Exception as e:
        log.warning(f"Could not extract header/footer: {e}")
    
    return header_info, footer_info


# ============== Semantic Section Detection ==============

def group_into_sections(paragraphs: List[ParsedParagraph], tables: List[ParsedTable]) -> List[ParsedSection]:
    """
    Group parsed elements into semantic sections based on headings and structure.
    """
    sections = []
    current_section = None
    current_list_items = []
    current_list_type = None
    
    for para in paragraphs:
        # Skip empty paragraphs
        if not para.text:
            continue
        
        # New heading = new section
        if para.is_heading and para.heading_level <= 2:
            # Flush any pending list
            if current_list_items and current_section:
                current_section.list_items = current_list_items
                current_section.list_type = current_list_type
                if current_section.type == SectionType.TEXT:
                    current_section.type = SectionType.LIST
            
            # Save previous section
            if current_section:
                sections.append(current_section)
            
            # Start new section
            section_type = SectionType.HEADER if para.heading_level == 0 or (para.heading_level == 1 and len(sections) == 0) else SectionType.TEXT
            current_section = ParsedSection(
                type=section_type,
                title=para.text,
                order=len(sections),
            )
            current_list_items = []
            current_list_type = None
            continue
        
        # Handle list items
        if para.is_list_item:
            current_list_items.append(para.text)
            current_list_type = para.list_type
            continue
        
        # Handle blank patterns (detail fields)
        if para.has_blank_pattern:
            if current_section is None:
                current_section = ParsedSection(type=SectionType.DETAIL, order=len(sections))
            
            current_section.fields.append({
                'label': para.detected_label,
                'path': para.suggested_field,
                'format': infer_field_format(para.suggested_field or '', para.detected_label or '').value,
            })
            
            # Convert to detail section if we have multiple fields
            if len(current_section.fields) >= 2:
                current_section.type = SectionType.DETAIL
            continue
        
        # Regular paragraph
        if current_section is None:
            # First content - check if it looks like a header/title
            if para.is_bold or para.font_size and para.font_size > 14:
                current_section = ParsedSection(
                    type=SectionType.HEADER,
                    title=para.text,
                    order=len(sections),
                )
                # Extract styling for header
                current_section.paragraphs.append(para)
            else:
                current_section = ParsedSection(type=SectionType.TEXT, order=len(sections))
                current_section.paragraphs.append(para)
        else:
            current_section.paragraphs.append(para)
    
    # Flush final section
    if current_list_items and current_section:
        current_section.list_items = current_list_items
        current_section.list_type = current_list_type
        if current_section.type == SectionType.TEXT:
            current_section.type = SectionType.LIST
    
    if current_section:
        sections.append(current_section)
    
    # Add tables to appropriate sections (or create new ones)
    for table in tables:
        # For now, add each table as its own section
        table_section = ParsedSection(
            type=SectionType.TABLE,
            tables=[table],
            order=len(sections),
        )
        sections.append(table_section)
    
    return sections


# ============== Main Analysis Function ==============

def analyze_document(doc_bytes: bytes, entity_def: str = "") -> DocumentAnalysis:
    """
    Perform complete semantic analysis of an uploaded document.
    
    Args:
        doc_bytes: Raw bytes of the DOCX file
        entity_def: Optional entity definition for context
        
    Returns:
        DocumentAnalysis with extracted structure and suggestions
    """
    doc = Document(io.BytesIO(doc_bytes))
    warnings = []
    suggestions = []
    
    # Parse all paragraphs
    parsed_paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        parsed = parse_paragraph(para, i)
        parsed_paragraphs.append(parsed)
    
    # Parse all tables
    parsed_tables = []
    for i, table in enumerate(doc.tables):
        parsed = parse_table(table, i)
        parsed_tables.append(parsed)
    
    # Extract header/footer
    page_header, page_footer = extract_header_footer(doc)
    
    # Group into semantic sections
    sections = group_into_sections(parsed_paragraphs, parsed_tables)
    
    # Extract all detected fields
    all_fields = []
    for section in sections:
        all_fields.extend(section.fields)
    
    # Try to detect entity type from content
    detected_entity = None
    content_lower = " ".join(p.text.lower() for p in parsed_paragraphs)
    if 'rfi' in content_lower or 'request for information' in content_lower:
        detected_entity = "RFI"
    elif 'invoice' in content_lower:
        detected_entity = "Invoice"
    elif 'contract' in content_lower or 'agreement' in content_lower:
        detected_entity = "Contract"
    elif 'submittal' in content_lower:
        detected_entity = "Submittal"
    elif 'punch' in content_lower or 'deficiency' in content_lower:
        detected_entity = "PunchList"
    
    # Generate suggestions
    if not all_fields:
        suggestions.append("No field placeholders detected. Add labels like 'ID: ' or 'Status: ' followed by whitespace.")
    
    if not any(s.type == SectionType.HEADER for s in sections):
        suggestions.append("Consider adding a header section with title and key identifiers.")
    
    if page_header is None:
        suggestions.append("Consider adding a page header with company name or document title.")
    
    if page_footer is None:
        suggestions.append("Consider adding a page footer with page numbers.")
    
    # Extract overall style info
    style_info = {}
    for para in parsed_paragraphs:
        if para.font_name and 'font_family' not in style_info:
            style_info['font_family'] = para.font_name
        if para.font_color and 'primary_color' not in style_info:
            style_info['primary_color'] = para.font_color
    
    return DocumentAnalysis(
        sections=sections,
        page_header=page_header,
        page_footer=page_footer,
        detected_entity_type=detected_entity or (entity_def.split('.')[-1] if entity_def else None),
        detected_fields=all_fields,
        style_info=style_info,
        warnings=warnings,
        suggestions=suggestions,
    )


# ============== Template Generation ==============

def analysis_to_template(
    analysis: DocumentAnalysis,
    entity_def: str,
    template_name: str = "Imported Template"
) -> PortableViewTemplate:
    """
    Convert document analysis into a PortableViewTemplate.
    
    Args:
        analysis: Result from analyze_document
        entity_def: Kahua entity definition
        template_name: Name for the generated template
        
    Returns:
        Complete PortableViewTemplate ready for rendering
    """
    sections = []
    
    for parsed_section in analysis.sections:
        if parsed_section.type == SectionType.HEADER:
            # Create header config
            title_para = parsed_section.paragraphs[0] if parsed_section.paragraphs else None
            config = HeaderConfig(
                static_title=parsed_section.title,
                title_font=title_para.font_name if title_para else None,
                title_size=int(title_para.font_size) if title_para and title_para.font_size else None,
                title_color=title_para.font_color if title_para else None,
                title_bold=title_para.is_bold if title_para else True,
                title_alignment=Alignment(title_para.alignment) if title_para else Alignment.LEFT,
            )
            sections.append(Section(
                type=SectionType.HEADER,
                order=parsed_section.order,
                header_config=config,
            ))
        
        elif parsed_section.type == SectionType.DETAIL:
            # Create detail config from detected fields
            fields = [
                FieldDef(
                    path=f['path'],
                    label=f['label'],
                    format=FieldFormat(f.get('format', 'text')),
                )
                for f in parsed_section.fields
            ]
            config = DetailConfig(
                fields=fields,
                columns=2 if len(fields) > 2 else 1,
            )
            sections.append(Section(
                type=SectionType.DETAIL,
                title=parsed_section.title,
                order=parsed_section.order,
                detail_config=config,
            ))
        
        elif parsed_section.type == SectionType.LIST:
            # Create list config
            config = ListConfig(
                list_type=parsed_section.list_type or "bullet",
                items=parsed_section.list_items,
            )
            sections.append(Section(
                type=SectionType.LIST,
                title=parsed_section.title,
                order=parsed_section.order,
                list_config=config,
            ))
        
        elif parsed_section.type == SectionType.TABLE:
            # Create table config from parsed tables
            for table in parsed_section.tables:
                columns = [
                    TableColumn(
                        field=FieldDef(path=header, label=header),
                        alignment=Alignment.LEFT,
                    )
                    for header in table.headers
                ]
                config = TableConfig(
                    source="Items",  # Default collection name
                    columns=columns,
                )
                sections.append(Section(
                    type=SectionType.TABLE,
                    order=parsed_section.order,
                    table_config=config,
                ))
        
        elif parsed_section.type == SectionType.TEXT:
            # Create text config
            content = "\n".join(p.text for p in parsed_section.paragraphs)
            config = TextConfig(content=content)
            sections.append(Section(
                type=SectionType.TEXT,
                title=parsed_section.title,
                order=parsed_section.order,
                text_config=config,
            ))
    
    # Build layout config
    layout = LayoutConfig()
    
    if analysis.page_header:
        layout.page_header = PageHeaderFooterConfig(
            center_text=analysis.page_header.get('text'),
            font_size=10,
        )
    
    if analysis.page_footer:
        layout.page_footer = PageHeaderFooterConfig(
            center_text=analysis.page_footer.get('text'),
            include_page_number=True,
            font_size=9,
        )
    elif not analysis.page_footer:
        # Add default footer with page numbers
        layout.page_footer = PageHeaderFooterConfig(
            include_page_number=True,
            page_number_format="Page {page} of {total}",
            font_size=9,
        )
    
    # Build style config
    style = StyleConfig()
    if analysis.style_info.get('font_family'):
        style.font_family = analysis.style_info['font_family']
    if analysis.style_info.get('primary_color'):
        style.primary_color = analysis.style_info['primary_color']
    
    return PortableViewTemplate(
        name=template_name,
        entity_def=entity_def,
        layout=layout,
        style=style,
        sections=sections,
        source_type="imported",
    )


# ============== High-Level API ==============

def analyze_and_convert(
    doc_bytes: bytes,
    entity_def: str,
    template_name: str = "Imported Template"
) -> Dict[str, Any]:
    """
    Analyze a document and convert it to a PortableViewTemplate.
    
    Args:
        doc_bytes: Raw DOCX bytes
        entity_def: Target Kahua entity
        template_name: Name for the template
        
    Returns:
        Dict with analysis results and generated template
    """
    analysis = analyze_document(doc_bytes, entity_def)
    template = analysis_to_template(analysis, entity_def, template_name)
    
    # Build summary
    section_summary = []
    for s in template.sections:
        config = s.get_config()
        field_count = 0
        if hasattr(config, 'fields'):
            field_count = len(config.fields)
        elif hasattr(config, 'columns'):
            field_count = len(config.columns)
        elif hasattr(config, 'items'):
            field_count = len(config.items)
        
        section_summary.append({
            'type': s.type.value,
            'title': s.title,
            'field_count': field_count,
        })
    
    return {
        'status': 'ok',
        'analysis': {
            'detected_entity': analysis.detected_entity_type,
            'section_count': len(analysis.sections),
            'field_count': len(analysis.detected_fields),
            'has_page_header': analysis.page_header is not None,
            'has_page_footer': analysis.page_footer is not None,
            'warnings': analysis.warnings,
            'suggestions': analysis.suggestions,
        },
        'template': template.model_dump(),
        'template_id': template.id,
        'sections_summary': section_summary,
    }


if __name__ == "__main__":
    import sys
    import json
    
    if len(sys.argv) < 2:
        print("Usage: python agentic_template_analyzer.py <path_to_docx> [entity_def]")
        sys.exit(1)
    
    doc_path = Path(sys.argv[1])
    entity_def = sys.argv[2] if len(sys.argv) > 2 else "kahua_AEC_RFI.RFI"
    
    if not doc_path.exists():
        print(f"File not found: {doc_path}")
        sys.exit(1)
    
    logging.basicConfig(level=logging.INFO)
    
    with open(doc_path, 'rb') as f:
        doc_bytes = f.read()
    
    result = analyze_and_convert(doc_bytes, entity_def, doc_path.stem)
    
    print("\n=== Document Analysis ===")
    print(f"Detected Entity: {result['analysis']['detected_entity']}")
    print(f"Sections: {result['analysis']['section_count']}")
    print(f"Fields: {result['analysis']['field_count']}")
    
    print("\n=== Generated Template ===")
    print(f"ID: {result['template_id']}")
    print(f"Sections:")
    for s in result['sections_summary']:
        print(f"  - {s['type']}: {s['title'] or '(untitled)'} ({s['field_count']} fields)")
    
    if result['analysis']['suggestions']:
        print("\n=== Suggestions ===")
        for s in result['analysis']['suggestions']:
            print(f"  - {s}")
